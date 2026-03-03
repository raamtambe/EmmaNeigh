from __future__ import annotations

import copy
import json
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple
import xml.etree.ElementTree as ET

from docx import Document
from docx.document import Document as _Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


@dataclass
class BluePlaceholder:
    start: int
    end: int
    token: str
    inner: str


@dataclass
class ClauseMatch:
    clause_id: int
    paragraph: Paragraph
    start: int
    end: int
    text: str
    blue_placeholders: List[BluePlaceholder] = field(default_factory=list)
    footnote_id: Optional[str] = None


@dataclass
class ParagraphLayout:
    runs: List[Run]
    run_ranges: List[Tuple[int, int]]
    text: str
    per_char_highlight: List[Optional[str]]


class DeterministicAgreementTransformer:
    """
    Deterministic rule-based DOCX transformer.
    No LLM calls, no probabilistic behavior.
    """

    def __init__(self) -> None:
        self.change_log: Dict[str, Any] = {
            "deleted_first_yellow_block": False,
            "clauses_deleted": [],
            "clauses_preserved_for_missing_data": [],
            "blue_placeholders_filled": {},
            "blue_placeholders_missing": [],
            "global_variables": {},
            "footnotes_removed": False,
        }
        self._clause_counter = 0
        self._footnote_map: Dict[str, str] = {}
        self._currency_selection: Optional[str] = None

    # -----------------------------
    # Public API
    # -----------------------------
    def transform(
        self,
        template_path: str | Path,
        extracted_json: Dict[str, Any] | str,
        output_path: str | Path,
        change_log_path: Optional[str | Path] = None,
    ) -> Dict[str, Any]:
        template_path = str(template_path)
        output_path = str(output_path)

        data = self._normalize_data(extracted_json)
        self._footnote_map = self._read_footnotes_text(template_path)

        doc = self.load_document(template_path)
        self.merge_runs_by_format(doc)
        all_clauses = self.detect_yellow_blocks(doc)
        first_clause_id = self.delete_first_yellow_block(all_clauses)

        clauses_by_paragraph: Dict[int, List[ClauseMatch]] = {}
        for clause in all_clauses:
            if clause.clause_id == first_clause_id:
                continue
            key = id(clause.paragraph)
            clauses_by_paragraph.setdefault(key, []).append(clause)

        for paragraph in self._iter_all_paragraphs(doc):
            key = id(paragraph)
            para_clauses = clauses_by_paragraph.get(key, [])
            if not para_clauses:
                continue

            # Process in reverse textual order so offsets stay valid.
            para_clauses.sort(key=lambda c: c.start, reverse=True)
            for clause in para_clauses:
                self._process_clause(paragraph, clause, data)

        self.apply_global_variables(doc)
        self.remove_all_footnotes(doc)
        self.cleanup_formatting(doc)
        self.export_document(doc, output_path)
        self._remove_footnotes_part_content(output_path)
        self.change_log["footnotes_removed"] = True

        if change_log_path:
            Path(change_log_path).write_text(
                json.dumps(self.change_log, indent=2), encoding="utf-8"
            )

        return self.change_log

    # -----------------------------
    # Pipeline stage methods
    # -----------------------------
    def load_document(self, template_path: str) -> _Document:
        return Document(template_path)

    def merge_runs_by_format(self, doc: _Document) -> None:
        for paragraph in self._iter_all_paragraphs(doc):
            i = 0
            while i < len(paragraph.runs) - 1:
                left = paragraph.runs[i]
                right = paragraph.runs[i + 1]
                if self._same_run_format(left, right):
                    left.text = (left.text or "") + (right.text or "")
                    right._element.getparent().remove(right._element)
                else:
                    i += 1

    def detect_yellow_blocks(self, doc: _Document) -> List[ClauseMatch]:
        clauses: List[ClauseMatch] = []
        for paragraph in self._iter_all_paragraphs(doc):
            layout = self._build_layout(paragraph)
            if not layout.text:
                continue

            top_level_brackets, nested_brackets = self._parse_brackets(layout.text)
            for start, end in top_level_brackets:
                segment_hl = layout.per_char_highlight[start : end + 1]
                if "yellow" not in segment_hl:
                    continue

                inner_blues: List[BluePlaceholder] = []
                for n_start, n_end in nested_brackets:
                    if not (start < n_start and n_end < end):
                        continue
                    nested_hl = layout.per_char_highlight[n_start : n_end + 1]
                    if nested_hl and all(h == "blue" for h in nested_hl):
                        token = layout.text[n_start : n_end + 1]
                        inner = token[1:-1].strip()
                        inner_blues.append(
                            BluePlaceholder(
                                start=n_start,
                                end=n_end,
                                token=token,
                                inner=inner,
                            )
                        )

                self._clause_counter += 1
                clause = ClauseMatch(
                    clause_id=self._clause_counter,
                    paragraph=paragraph,
                    start=start,
                    end=end,
                    text=layout.text[start : end + 1],
                    blue_placeholders=inner_blues,
                    footnote_id=self._detect_attached_footnote_id(paragraph, end),
                )
                clauses.append(clause)

        return clauses

    def delete_first_yellow_block(self, clauses: List[ClauseMatch]) -> Optional[int]:
        if not clauses:
            return None
        first = clauses[0]
        self._replace_text_range_in_paragraph(first.paragraph, first.start, first.end, "")
        self.change_log["deleted_first_yellow_block"] = True
        self.change_log["clauses_deleted"].append(
            {
                "clause_id": first.clause_id,
                "reason": "RULE_1_ALWAYS_DELETE_FIRST_YELLOW_BLOCK",
                "text": first.text,
            }
        )
        return first.clause_id

    def detect_blue_placeholders(self, clause: ClauseMatch) -> List[BluePlaceholder]:
        return clause.blue_placeholders

    def detect_footnotes(self) -> Dict[str, str]:
        return self._footnote_map

    def evaluate_conditional_rules(self, footnote_text: str, data: Dict[str, Any]) -> bool:
        text = (footnote_text or "").lower()
        if not text.strip():
            return True

        predicate = self._evaluate_predicate(text, data)
        if predicate is None:
            # Preserve template over guessing.
            return True

        if "remove if" in text or "delete if" in text:
            return not predicate
        if "only include if" in text or "include if" in text:
            return predicate
        # Default: interpret as applicability condition.
        return predicate

    def populate_blue_placeholders(
        self,
        clause_text: str,
        blue_placeholders: List[BluePlaceholder],
        data: Dict[str, Any],
    ) -> Tuple[str, List[str]]:
        missing: List[str] = []
        updated = clause_text

        for ph in sorted(blue_placeholders, key=lambda p: len(p.token), reverse=True):
            value = self._resolve_placeholder_value(ph.inner, data)
            if value is None or value == "":
                missing.append(ph.inner)
                continue
            updated = updated.replace(ph.token, str(value))
            self.change_log["blue_placeholders_filled"][ph.inner] = str(value)

        return updated, missing

    def apply_global_variables(self, doc: _Document) -> None:
        # Currency variable is already applied at replacement time.
        if self._currency_selection is not None:
            self.change_log["global_variables"]["currency_selection"] = self._currency_selection

    def remove_all_footnotes(self, doc: _Document) -> None:
        for paragraph in self._iter_all_paragraphs(doc):
            for run in paragraph.runs:
                refs = run._r.xpath(".//w:footnoteReference")
                for ref in refs:
                    ref.getparent().remove(ref)
                ref_marks = run._r.xpath(".//w:footnoteRef")
                for mark in ref_marks:
                    mark.getparent().remove(mark)

    def cleanup_formatting(self, doc: _Document) -> None:
        # Remove empty runs that were left behind.
        for paragraph in self._iter_all_paragraphs(doc):
            for run in list(paragraph.runs):
                if run.text == "":
                    run._element.getparent().remove(run._element)

    def export_document(self, doc: _Document, output_path: str) -> None:
        doc.save(output_path)

    def generate_change_log(self) -> Dict[str, Any]:
        return self.change_log

    # -----------------------------
    # Clause processing
    # -----------------------------
    def _process_clause(
        self,
        paragraph: Paragraph,
        clause: ClauseMatch,
        data: Dict[str, Any],
    ) -> None:
        footnote_text = ""
        keep_by_condition = True
        if clause.footnote_id:
            footnote_text = self._footnote_map.get(clause.footnote_id, "")
            keep_by_condition = self.evaluate_conditional_rules(footnote_text, data)

        if not keep_by_condition:
            self._replace_text_range_in_paragraph(paragraph, clause.start, clause.end, "")
            self.change_log["clauses_deleted"].append(
                {
                    "clause_id": clause.clause_id,
                    "reason": "CONDITIONAL_RULE_FALSE",
                    "footnote_id": clause.footnote_id,
                    "footnote_text": footnote_text,
                    "text": clause.text,
                }
            )
            return

        blue_placeholders = self.detect_blue_placeholders(clause)
        updated_text, missing = self.populate_blue_placeholders(clause.text, blue_placeholders, data)
        if missing:
            self.change_log["clauses_preserved_for_missing_data"].append(
                {
                    "clause_id": clause.clause_id,
                    "missing": sorted(set(missing)),
                    "text": clause.text,
                }
            )
            for key in sorted(set(missing)):
                if key not in self.change_log["blue_placeholders_missing"]:
                    self.change_log["blue_placeholders_missing"].append(key)
            # RULE 2: preserve unchanged clause with brackets/highlight intact.
            return

        # Finalize clause: remove outer wrapper brackets and remove highlight.
        finalized = self._remove_outer_brackets(updated_text)
        self._replace_text_range_in_paragraph(
            paragraph,
            clause.start,
            clause.end,
            finalized,
            clear_highlights=True,
        )

    # -----------------------------
    # Low-level helpers
    # -----------------------------
    def _normalize_data(self, extracted_json: Dict[str, Any] | str) -> Dict[str, Any]:
        if isinstance(extracted_json, str):
            extracted_json = json.loads(extracted_json)

        out: Dict[str, Any] = {}
        for k, v in extracted_json.items():
            out[k] = v
            out[self._norm(k)] = v

        # Canonical aliases.
        aliases = {
            "advisor_name": ["name", "advisor"],
            "advisor_email": ["email"],
            "advisor_mailing_address": ["mailing_address", "address"],
            "is_us_taxpayer": ["us_taxpayer", "taxpayer"],
            "engagement_mode": ["individual_or_entity", "mode"],
            "daily_rate": ["rate", "daily_fee"],
            "currency": ["payment_currency"],
            "term": ["term_length"],
            "geography": ["region"],
            "advisor_type": ["type"],
            "fee_mode": ["fee"],
        }
        for canonical, keys in aliases.items():
            if canonical in out:
                continue
            for key in keys:
                if key in out:
                    out[canonical] = out[key]
                    break
                nk = self._norm(key)
                if nk in out:
                    out[canonical] = out[nk]
                    break

        # Normalize useful values.
        geo = str(out.get("geography", "")).strip().upper().replace("/", "_").replace(" ", "_")
        if geo == "USUK":
            geo = "US_UK"
        out["geography"] = geo

        cur = str(out.get("currency", "")).strip().upper()
        out["currency"] = cur

        out["is_us_taxpayer"] = str(out.get("is_us_taxpayer", "")).strip().lower()
        out["engagement_mode"] = str(out.get("engagement_mode", "")).strip().lower()

        return out

    def _resolve_placeholder_value(self, placeholder_inner: str, data: Dict[str, Any]) -> Optional[str]:
        stripped = placeholder_inner.strip()

        # Currency multi-option placeholder e.g. $/EUROS/GBP
        if "/" in stripped and self._looks_like_currency_options(stripped):
            selected = self._select_currency_option(stripped, data.get("currency", ""))
            if selected is not None:
                self._currency_selection = selected
                return selected

        nk = self._norm(stripped)
        candidates = [
            stripped,
            nk,
            nk.replace("advisor", "advisor_"),
        ]

        # Heuristic mapping by token keywords
        keyword_map = [
            ("advisorname", "advisor_name"),
            ("name", "advisor_name"),
            ("email", "advisor_email"),
            ("mailingaddress", "advisor_mailing_address"),
            ("address", "advisor_mailing_address"),
            ("taxpayer", "is_us_taxpayer"),
            ("entity", "engagement_mode"),
            ("individual", "engagement_mode"),
            ("rate", "daily_rate"),
            ("fee", "daily_rate"),
            ("currency", "currency"),
            ("term", "term"),
            ("geography", "geography"),
            ("region", "geography"),
        ]
        for token, key in keyword_map:
            if token in nk:
                candidates.append(key)

        for key in candidates:
            if key in data and data[key] not in (None, ""):
                return str(data[key])

        return None

    def _build_layout(self, paragraph: Paragraph) -> ParagraphLayout:
        runs = list(paragraph.runs)
        text_parts: List[str] = []
        run_ranges: List[Tuple[int, int]] = []
        per_char_hl: List[Optional[str]] = []

        cursor = 0
        for run in runs:
            txt = run.text or ""
            text_parts.append(txt)
            start = cursor
            end = cursor + len(txt) - 1
            run_ranges.append((start, end))
            hl = self._run_highlight_name(run)
            for _ in txt:
                per_char_hl.append(hl)
            cursor += len(txt)

        return ParagraphLayout(
            runs=runs,
            run_ranges=run_ranges,
            text="".join(text_parts),
            per_char_highlight=per_char_hl,
        )

    def _parse_brackets(self, text: str) -> Tuple[List[Tuple[int, int]], List[Tuple[int, int]]]:
        stack: List[int] = []
        top_level: List[Tuple[int, int]] = []
        nested: List[Tuple[int, int]] = []

        for idx, ch in enumerate(text):
            if ch == "[":
                stack.append(idx)
            elif ch == "]" and stack:
                start = stack.pop()
                pair = (start, idx)
                nested.append(pair)
                if not stack:
                    top_level.append(pair)

        # nested includes top-level as well; caller can filter.
        return top_level, nested

    def _detect_attached_footnote_id(self, paragraph: Paragraph, clause_end_char: int) -> Optional[str]:
        layout = self._build_layout(paragraph)
        end_run_idx = self._find_run_index_for_char(layout.run_ranges, clause_end_char)
        if end_run_idx is None:
            return None

        for run in layout.runs[end_run_idx + 1 :]:
            refs = run._r.xpath(".//w:footnoteReference")
            if refs:
                rid = refs[0].get(qn("w:id"))
                return str(rid) if rid is not None else None
            if (run.text or "").strip():
                break
        return None

    def _replace_text_range_in_paragraph(
        self,
        paragraph: Paragraph,
        start_char: int,
        end_char: int,
        new_text: str,
        clear_highlights: bool = False,
    ) -> None:
        if start_char > end_char:
            return

        self._split_paragraph_runs_at_char(paragraph, end_char + 1)
        self._split_paragraph_runs_at_char(paragraph, start_char)

        layout = self._build_layout(paragraph)
        start_run = self._find_run_index_for_char(layout.run_ranges, start_char)
        end_run = self._find_run_index_for_char(layout.run_ranges, end_char)
        if start_run is None or end_run is None:
            return

        target_runs = paragraph.runs[start_run : end_run + 1]
        if not target_runs:
            return

        first = target_runs[0]
        first.text = new_text
        if clear_highlights:
            first.font.highlight_color = None

        # Remove remaining runs in range.
        for run in target_runs[1:]:
            run._element.getparent().remove(run._element)

        # Ensure first run highlight cleared for finalized clause content.
        if clear_highlights and new_text:
            first.font.highlight_color = None

    def _split_paragraph_runs_at_char(self, paragraph: Paragraph, absolute_char: int) -> None:
        if absolute_char <= 0:
            return
        layout = self._build_layout(paragraph)
        if absolute_char >= len(layout.text):
            return

        idx = self._find_run_index_for_char(layout.run_ranges, absolute_char)
        if idx is None:
            return

        run = paragraph.runs[idx]
        start, _ = layout.run_ranges[idx]
        offset = absolute_char - start
        txt = run.text or ""
        if offset <= 0 or offset >= len(txt):
            return

        left = txt[:offset]
        right = txt[offset:]
        run.text = left

        new_r = copy.deepcopy(run._element)
        new_run = Run(new_r, paragraph)
        new_run.text = right
        run._element.addnext(new_run._element)

    def _find_run_index_for_char(
        self, run_ranges: List[Tuple[int, int]], absolute_char: int
    ) -> Optional[int]:
        for i, (start, end) in enumerate(run_ranges):
            if start <= absolute_char <= end:
                return i
        return None

    def _remove_outer_brackets(self, text: str) -> str:
        if text.startswith("[") and text.endswith("]"):
            return text[1:-1]
        return text

    def _same_run_format(self, left: Run, right: Run) -> bool:
        return (
            left.bold == right.bold
            and left.italic == right.italic
            and left.underline == right.underline
            and left.font.name == right.font.name
            and left.font.size == right.font.size
            and left.font.color.rgb == right.font.color.rgb
            and self._run_highlight_name(left) == self._run_highlight_name(right)
        )

    def _run_highlight_name(self, run: Run) -> Optional[str]:
        direct = run.font.highlight_color
        if direct == WD_COLOR_INDEX.YELLOW:
            return "yellow"
        if direct == WD_COLOR_INDEX.BLUE:
            return "blue"

        # Fallback to raw XML value
        hl_nodes = run._r.xpath("./w:rPr/w:highlight")
        if hl_nodes:
            val = hl_nodes[0].get(qn("w:val"))
            if val:
                return str(val).lower()
        return None

    def _iter_all_paragraphs(self, doc: _Document) -> Iterable[Paragraph]:
        seen: set[int] = set()

        def emit(paragraphs: Iterable[Paragraph]) -> Iterable[Paragraph]:
            for p in paragraphs:
                marker = id(p._p)
                if marker in seen:
                    continue
                seen.add(marker)
                yield p

        for p in emit(doc.paragraphs):
            yield p

        for table in doc.tables:
            for p in self._iter_table_paragraphs(table, seen):
                yield p

        for section in doc.sections:
            parts = [
                section.header,
                section.first_page_header,
                section.even_page_header,
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            ]
            for part in parts:
                for p in emit(part.paragraphs):
                    yield p
                for table in part.tables:
                    for p in self._iter_table_paragraphs(table, seen):
                        yield p

    def _iter_table_paragraphs(self, table: Table, seen: set[int]) -> Iterable[Paragraph]:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    marker = id(p._p)
                    if marker in seen:
                        continue
                    seen.add(marker)
                    yield p
                for nested in cell.tables:
                    for p in self._iter_table_paragraphs(nested, seen):
                        yield p

    def _read_footnotes_text(self, template_path: str) -> Dict[str, str]:
        out: Dict[str, str] = {}
        with zipfile.ZipFile(template_path, "r") as zf:
            if "word/footnotes.xml" not in zf.namelist():
                return out
            raw = zf.read("word/footnotes.xml")

        root = ET.fromstring(raw)
        for fn in root.findall(f".//{{{W_NS}}}footnote"):
            fid = fn.attrib.get(f"{{{W_NS}}}id")
            if fid is None:
                continue
            parts: List[str] = []
            for t in fn.findall(f".//{{{W_NS}}}t"):
                parts.append(t.text or "")
            text = "".join(parts).strip()
            if text:
                out[str(fid)] = text
        return out

    def _remove_footnotes_part_content(self, output_path: str) -> None:
        with zipfile.ZipFile(output_path, "a") as zf:
            if "word/footnotes.xml" not in zf.namelist():
                return
            raw = zf.read("word/footnotes.xml")
            root = ET.fromstring(raw)
            for fn in list(root.findall(f".//{{{W_NS}}}footnote")):
                fid = fn.attrib.get(f"{{{W_NS}}}id")
                if fid in ("0", "-1"):
                    continue
                root.remove(fn)
            zf.writestr("word/footnotes.xml", ET.tostring(root, encoding="utf-8", xml_declaration=True))

    def _evaluate_predicate(self, text: str, data: Dict[str, Any]) -> Optional[bool]:
        # Geography predicates
        if "latam" in text or "latin america" in text:
            val = str(data.get("geography", "")).upper() == "LATAM"
            if re.search(r"not\s+(in\s+)?(latin america|latam)", text):
                val = not val
            return val

        if re.search(r"\beu\b|europe", text):
            val = str(data.get("geography", "")).upper() == "EU"
            if re.search(r"not\s+(in\s+)?(eu|europe)", text):
                val = not val
            return val

        if "us/uk" in text or "united states" in text or "united kingdom" in text:
            val = str(data.get("geography", "")).upper() == "US_UK"
            if re.search(r"not\s+(in\s+)?(us/uk|united states|united kingdom)", text):
                val = not val
            return val

        # Taxpayer predicates
        if "taxpayer" in text:
            is_tax = str(data.get("is_us_taxpayer", "")).lower() in ("yes", "true", "1")
            if re.search(r"not\s+(a\s+)?(u\.?s\.?\s*)?taxpayer", text):
                is_tax = not is_tax
            return is_tax

        # Engagement mode predicates
        if "individual" in text:
            is_individual = str(data.get("engagement_mode", "")).lower() == "individual"
            if re.search(r"not\s+an?\s+individual", text):
                is_individual = not is_individual
            return is_individual

        if "entity" in text or "llc" in text or "corporation" in text:
            is_entity = str(data.get("engagement_mode", "")).lower() == "entity"
            if re.search(r"not\s+an?\s+entity", text):
                is_entity = not is_entity
            return is_entity

        # Fee mode predicates
        if "without additional fee" in text or "without fee" in text:
            return str(data.get("fee_mode", "")).upper() == "WITHOUT_FEE"
        if "with additional fee" in text or "with fee" in text:
            return str(data.get("fee_mode", "")).upper() == "WITH_FEE"

        return None

    def _looks_like_currency_options(self, token: str) -> bool:
        lower = token.lower()
        return any(x in lower for x in ("$", "usd", "eur", "euro", "gbp", "pound"))

    def _select_currency_option(self, token: str, currency: str) -> Optional[str]:
        options = [p.strip() for p in token.split("/") if p.strip()]
        cur = (currency or "").upper()

        option_map = {
            "USD": ["$", "USD", "US DOLLAR", "US DOLLARS"],
            "EUR": ["EUR", "EURO", "EUROS", "€"],
            "GBP": ["GBP", "POUND", "POUNDS", "£"],
        }

        if cur in option_map:
            wanted = option_map[cur]
            for op in options:
                up = op.upper()
                if any(w in up for w in wanted):
                    return op

        # Fallback if currency unknown: no inference.
        return None

    def _norm(self, text: str) -> str:
        return re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_")


def transform_agreement(
    template_path: str | Path,
    extracted_json: Dict[str, Any] | str,
    output_path: str | Path,
    change_log_path: Optional[str | Path] = None,
) -> Dict[str, Any]:
    transformer = DeterministicAgreementTransformer()
    return transformer.transform(template_path, extracted_json, output_path, change_log_path)
