#!/usr/bin/env python3
"""
Checklist document-name extractor for Agent Mode batch redlines.

Usage:
    python checklist_docname_extractor.py <checklist_path>
"""

import json
import re
import sys
from docx import Document


DOCUMENT_COLUMN_PATTERNS = [
    r"^document",
    r"^doc\s*name",
    r"^agreement",
    r"^instrument",
    r"^deliverable",
    r"^item",
    r"^description",
    r"^closing\s*document",
]

STATUS_COLUMN_PATTERNS = [
    r"^status",
    r"^state",
    r"^progress",
    r"^current\s*status",
]

SKIP_DOC_PATTERNS = (
    r"^n/?a$",
    r"^none$",
    r"^tbd$",
    r"^to be determined$",
    r"^document$",
    r"^documents$",
)


def normalize_cell_text(value):
    return re.sub(r"\s+", " ", str(value or "")).strip()


def find_column_index(headers, patterns):
    for i, header in enumerate(headers):
        header_lower = normalize_cell_text(header).lower()
        for pattern in patterns:
            if re.match(pattern, header_lower):
                return i
    return -1


def infer_document_column(headers, rows):
    if not headers:
        return -1

    max_cols = max(len(headers), max((len(r) for r in rows), default=0))
    best_col = -1
    best_score = -1
    blocked_terms = ("status", "state", "progress", "date", "sent", "received")

    for col_idx in range(max_cols):
        header_value = normalize_cell_text(headers[col_idx] if col_idx < len(headers) else "")
        header_lower = header_value.lower()
        if any(term in header_lower for term in blocked_terms):
            continue

        non_empty = 0
        alpha_like = 0
        penalties = 0
        for row in rows[:200]:
            cell = normalize_cell_text(row[col_idx] if col_idx < len(row) else "")
            if not cell:
                continue
            non_empty += 1
            if re.search(r"[A-Za-z]", cell):
                alpha_like += 1
            if len(cell) > 260 or re.fullmatch(r"[\d\W_]+", cell):
                penalties += 1

        score = (non_empty * 2) + (alpha_like * 3) - (penalties * 2)
        if score > best_score:
            best_score = score
            best_col = col_idx

    return best_col


def score_header_candidate(all_rows, header_idx):
    headers = all_rows[header_idx]
    non_empty_headers = [normalize_cell_text(h) for h in headers if normalize_cell_text(h)]
    if len(non_empty_headers) < 2:
        return None

    doc_col = find_column_index(headers, DOCUMENT_COLUMN_PATTERNS)
    status_col = find_column_index(headers, STATUS_COLUMN_PATTERNS)
    avg_header_len = (
        sum(len(x) for x in non_empty_headers) / len(non_empty_headers)
        if non_empty_headers else 0
    )
    looks_like_header = 1 if avg_header_len <= 48 else -5

    data_rows = all_rows[header_idx + 1: header_idx + 11]
    non_empty_data_rows = sum(1 for row in data_rows if any(normalize_cell_text(c) for c in row))

    score = len(non_empty_headers) + (non_empty_data_rows * 3) + looks_like_header
    if doc_col != -1:
        score += 40
    if status_col != -1:
        score += 15

    return {
        "score": score,
        "header_idx": header_idx,
        "doc_col": doc_col,
        "status_col": status_col,
    }


def parse_best_checklist_table(doc):
    if not doc.tables:
        return None

    best_candidate = None
    for table in doc.tables:
        all_rows = []
        for row in table.rows:
            all_rows.append([normalize_cell_text(cell.text) for cell in row.cells])

        if not all_rows:
            continue

        max_probe = min(6, len(all_rows) - 1)
        for header_idx in range(max_probe + 1):
            candidate = score_header_candidate(all_rows, header_idx)
            if not candidate:
                continue
            candidate["all_rows"] = all_rows
            if best_candidate is None or candidate["score"] > best_candidate["score"]:
                best_candidate = candidate

    if not best_candidate:
        return None

    all_rows = best_candidate["all_rows"]
    header_idx = best_candidate["header_idx"]
    headers = all_rows[header_idx]
    rows = [row for row in all_rows[header_idx + 1:] if any(normalize_cell_text(c) for c in row)]
    doc_col_idx = best_candidate["doc_col"]
    if doc_col_idx == -1:
        doc_col_idx = infer_document_column(headers, rows)

    return {
        "headers": headers,
        "rows": rows,
        "doc_col_idx": doc_col_idx,
        "data_row_start_idx": header_idx + 1,
    }


def canonical_doc_key(value):
    return re.sub(r"\s+", " ", str(value or "")).strip().lower()


def should_skip_document_name(value):
    text = normalize_cell_text(value)
    if not text:
        return True
    lower = text.lower()
    for pattern in SKIP_DOC_PATTERNS:
        if re.match(pattern, lower):
            return True
    if len(text) < 3:
        return True
    if re.fullmatch(r"[\d\W_]+", text):
        return True
    return False


def extract_checklist_document_names(checklist_path):
    doc = Document(checklist_path)
    parsed = parse_best_checklist_table(doc)
    if not parsed:
        return []

    headers = parsed["headers"]
    rows = parsed["rows"]
    doc_col_idx = parsed["doc_col_idx"]
    data_row_start_idx = parsed["data_row_start_idx"]

    if doc_col_idx < 0:
        return []

    deduped = {}
    ordered_keys = []

    for row_idx, row in enumerate(rows):
        if doc_col_idx >= len(row):
            continue
        doc_name = normalize_cell_text(row[doc_col_idx])
        if should_skip_document_name(doc_name):
            continue

        row_context_parts = []
        for col_idx, cell_value in enumerate(row):
            if col_idx == doc_col_idx:
                continue
            cleaned = normalize_cell_text(cell_value)
            if not cleaned:
                continue
            header_label = normalize_cell_text(headers[col_idx] if col_idx < len(headers) else "")
            if header_label:
                row_context_parts.append(f"{header_label}: {cleaned}")
            else:
                row_context_parts.append(cleaned)

        row_id = data_row_start_idx + row_idx
        key = canonical_doc_key(doc_name)
        candidate = {
            "row_id": row_id,
            "document_name": doc_name,
            "row_context": " | ".join(row_context_parts[:8]),
        }
        if key not in deduped:
            deduped[key] = candidate
            ordered_keys.append(key)
        elif len(candidate["document_name"]) > len(deduped[key]["document_name"]):
            deduped[key] = candidate

    return [deduped[key] for key in ordered_keys]


def main():
    if len(sys.argv) < 2:
        print(json.dumps({
            "success": False,
            "error": "Usage: checklist_docname_extractor.py <checklist_path>",
        }))
        sys.exit(1)

    checklist_path = sys.argv[1]
    try:
        items = extract_checklist_document_names(checklist_path)
        print(json.dumps({
            "success": True,
            "count": len(items),
            "items": items,
        }))
    except Exception as exc:
        print(json.dumps({
            "success": False,
            "error": str(exc),
        }))
        sys.exit(1)


if __name__ == "__main__":
    main()
