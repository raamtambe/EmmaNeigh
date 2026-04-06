#!/usr/bin/env python3
"""
Checklist Updater - Update transaction checklist based on email activity

This module parses a transaction checklist (Word document) and an email activity dataset,
then updates the checklist status column based on detected email activity.

Usage:
    python checklist_updater.py <checklist_path> <email_input_path> <output_folder> <api_key> [provider] [model] [provider_base_url]
"""

import sys
import os
import json
import re
import socket
import time
import urllib.error
import urllib.parse
import urllib.request
import uuid
from datetime import datetime
import pandas as pd
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:  # pragma: no cover - environment dependent
    Document = None
    Pt = None
    Inches = None
    WD_ALIGN_PARAGRAPH = None

DEFAULT_CLAUDE_MODEL = "claude-sonnet-4-20250514"
DEFAULT_OPENAI_MODEL = "gpt-4.1-mini"
DEFAULT_HARVEY_BASE_URL = "https://api.harvey.ai"
DEFAULT_LOCALAI_BASE_URL = "http://127.0.0.1:11435"
DEFAULT_OLLAMA_BASE_URL = "http://127.0.0.1:11434"
DEFAULT_LMSTUDIO_BASE_URL = "http://127.0.0.1:1234"
DEFAULT_LOCALAI_MODEL = "qwen2.5:1.5b"
DEFAULT_OLLAMA_MODEL = "llama3.1:8b"
DEFAULT_LMSTUDIO_MODEL = "local-model"
MAX_HTTP_ATTEMPTS = 3
RETRYABLE_HTTP_STATUS = {408, 425, 429, 500, 502, 503, 504}
RETRYABLE_ERROR_MARKERS = (
    "timeout",
    "temporarily unavailable",
    "connection reset",
    "connection aborted",
    "network is unreachable",
    "name or service not known",
    "temporary failure in name resolution",
)

LLM_STATUS_VALUES = {
    "Pending Draft",
    "Draft Circulated",
    "With Opposing Counsel",
    "Agreed Form",
    "Execution Version",
    "Executed",
}

STATUS_ALIASES = {
    "pending": "Pending Draft",
    "pending draft": "Pending Draft",
    "to be drafted": "Pending Draft",
    "draft circulated": "Draft Circulated",
    "with opposing counsel": "With Opposing Counsel",
    "under review": "With Opposing Counsel",
    "agreed form": "Agreed Form",
    "execution version": "Execution Version",
    "executed": "Executed",
    "fully executed": "Executed",
}

# Status detection patterns - order matters (more specific first)
STATUS_PATTERNS = [
    # Executed/Signed (highest priority)
    {
        'patterns': [
            r'\bfully\s+executed\b',
            r'\bexecuted\s+(?:copy|version|document)\b',
            r'\bsigned\s+(?:by\s+all|copy|version)\b',
            r'\b(?:all\s+)?signatures?\s+(?:received|obtained|collected)\b',
        ],
        'status': 'Executed',
        'priority': 5
    },
    # Execution Version Circulated
    {
        'patterns': [
            r'\bexecution\s+(?:version|copy)\b',
            r'\bfor\s+(?:signature|execution)\b',
            r'\b(?:ready|circulated?)\s+for\s+signature\b',
            r'\bsignature\s+pages?\s+(?:attached|circulated)\b',
        ],
        'status': 'Execution Version',
        'priority': 4
    },
    # Agreed Form
    {
        'patterns': [
            r'\bagreed\s+form\b',
            r'\bfinal\s+(?:form|version)\b',
            r'\bfinalized\b',
            r'\bno\s+(?:further\s+)?comments\b',
        ],
        'status': 'Agreed Form',
        'priority': 3
    },
    # Sent to Opposing Counsel / Under Review
    {
        'patterns': [
            r'\bsent\s+to\s+(?:opposing\s+)?counsel\b',
            r'\bsent\s+to\s+(?:counterparty|other\s+side|buyer|seller|lender|borrower)\b',
            r'\b(?:circulated|distributed)\s+(?:to|for)\s+(?:review|comment)\b',
            r'\bfor\s+(?:your\s+)?review\b',
            r'\bplease\s+(?:review|comment)\b',
            r'\bawaiting\s+(?:comments?|review|feedback)\b',
            r'\bunder\s+review\b',
        ],
        'status': 'With Opposing Counsel',
        'priority': 2
    },
    # Draft Circulated (internal)
    {
        'patterns': [
            r'\b(?:initial\s+)?draft\s+(?:attached|circulated)\b',
            r'\battached\s+(?:is\s+)?(?:a\s+)?draft\b',
            r'\bfirst\s+draft\b',
        ],
        'status': 'Draft Circulated',
        'priority': 1
    },
]

# Column name patterns for detecting checklist columns
DOCUMENT_COLUMN_PATTERNS = [
    r'^document',
    r'^doc\s*name',
    r'^agreement',
    r'^instrument',
    r'^deliverable',
    r'^item',
    r'^description',
    r'^closing\s*document',
]

STATUS_COLUMN_PATTERNS = [
    r'^status',
    r'^state',
    r'^progress',
    r'^current\s*status',
]

PARTY_COLUMN_PATTERNS = [
    r'^party',
    r'^responsible',
    r'^owner',
    r'^assignee',
    r'^who',
    r'^counsel',
]

NOTES_COLUMN_PATTERNS = [
    r'^notes?',
    r'^comments?',
    r'^remarks?',
    r'^details?',
]

CHECKLIST_LLM_BATCH_SIZE = 10
MAX_CANDIDATE_EMAILS_PER_ROW = 8
MAX_CANDIDATE_THREADS_PER_ROW = 3
MAX_THREAD_EVENTS_PER_SUMMARY = 4
MAX_THREAD_ISSUES_PER_SUMMARY = 3
HIGH_CONFIDENCE_THREAD_SCORE = 8.5
SEARCH_TOKEN_STOPWORDS = {
    'agreement', 'agreements', 'amendment', 'amended', 'and', 'annex', 'appendix',
    'certificate', 'consent', 'copy', 'counsel', 'credit', 'date', 'document',
    'documents', 'draft', 'email', 'exhibit', 'execution', 'final', 'for', 'from',
    'guaranty', 'guarantee', 'hereof', 'incumbency', 'indenture', 'intercreditor',
    'joinder', 'letter', 'loan', 'name', 'note', 'officer', 'page', 'pages',
    'party', 'pledge', 'promissory', 'review', 'schedule', 'security', 'sent',
    'signature', 'signed', 'status', 'the', 'title', 'version', 'with'
}
DEAL_TOKEN_STOPWORDS = SEARCH_TOKEN_STOPWORDS | {
    'buyer', 'seller', 'company', 'holdings', 'parent', 'merger', 'acquisition',
    'closing', 'deal', 'transaction', 'group', 'corp', 'corporation', 'inc',
    'llc', 'ltd', 'limited', 'lp', 'plc', 'co', 'drafts', 'checklist'
}
THREAD_SUBJECT_PREFIX_RE = re.compile(r'^\s*(?:(?:re|fw|fwd|aw|wg)\s*:\s*)+', re.IGNORECASE)
THREAD_SUBJECT_TAG_RE = re.compile(r'^\s*\[[^\]]+\]\s*')
ISSUE_KEYWORDS = (
    'issue', 'issues', 'comment', 'comments', 'open point', 'outstanding', 'concern',
    'concerns', 'requested', 'request', 'revise', 'revision', 'revisions', 'markup',
    'redline', 'blackline', 'tbd', 'confirm', 'missing', 'need', 'needs', 'pending',
    'question', 'questions', 'bracket', 'fix', 'update', 'change'
)
DOCUMENT_ACTIVITY_KEYWORDS = (
    'attached', 'attachment', 'draft', 'revised draft', 'markup', 'redline',
    'blackline', 'comments', 'execution version', 'executed', 'signed', 'for review',
    'for signature', 'circulated'
)
DOCUMENT_ROW_KEYWORDS = (
    'agreement', 'amendment', 'assignment', 'bill of sale', 'bringdown', 'certificate',
    'consent', 'contract', 'disclosure schedule', 'escrow', 'exhibit', 'funds flow',
    'guaranty', 'guarantee', 'incumbency', 'instruction', 'joinder', 'lease', 'letter',
    'management', 'minutes', 'note', 'notice', 'officer', 'opinion', 'pledge',
    'power of attorney', 'proxy', 'questionnaire', 'release', 'resolution',
    'schedule', 'security', 'services', 'settlement', 'side letter', 'statement',
    'subscription', 'support', 'tax', 'termination', 'transition', 'voting'
)
SECTION_ROW_KEYWORDS = (
    'article', 'section', 'recitals', 'definitions', 'representations', 'warranties',
    'covenants', 'conditions', 'indemnification', 'miscellaneous', 'signatures',
    'signature pages', 'closing deliverables', 'general', 'other provisions'
)
DOCUMENT_FILENAME_NOISE_RE = re.compile(
    r'\b(?:draft|redline|blackline|markup|comments?|clean|execution|executed|signed|'
    r'final|revised|version|copy|copies|dated|sig(?:nature)?(?:\s+pages?)?|'
    r'v\d+|rev(?:ision)?\s*\d+)\b',
    re.IGNORECASE,
)
ROW_NUMBER_PREFIX_RE = re.compile(
    r'^\s*(?:\(?\d+(?:\.\d+)*\)?[.)-]?|\(?[a-z]\)?[.)-]?|\(?[ivxlcdm]+\)?[.)-]?)\s+',
    re.IGNORECASE,
)
INLINE_DATE_PATTERNS = (
    re.compile(r'\b\d{1,2}/\d{1,2}/\d{2,4}\b'),
    re.compile(r'\b\d{1,2}-\d{1,2}-\d{2,4}\b'),
    re.compile(r'\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\s+\d{1,2},\s+\d{4}\b', re.IGNORECASE),
)


def normalize_email_record(raw_email):
    """Normalize one email record from CSV or JSON input."""
    if not isinstance(raw_email, dict):
        return None

    attachments = raw_email.get('attachments', '')
    if isinstance(attachments, list):
        attachment_text = '; '.join(str(item).strip() for item in attachments if str(item).strip())
    else:
        attachment_text = str(attachments or '').strip()

    has_attachments = raw_email.get('has_attachments')
    if isinstance(has_attachments, str):
        has_attachments = has_attachments.strip().lower() not in ('', '0', 'false', 'no', 'none', 'nan')
    else:
        has_attachments = bool(has_attachments)
    if not has_attachments and attachment_text:
        has_attachments = attachment_text.lower() not in ('none', 'false', 'no', 'n/a')

    is_sent_folder = raw_email.get('is_sent_folder')
    if isinstance(is_sent_folder, str):
        is_sent_folder = is_sent_folder.strip().lower() not in ('', '0', 'false', 'no', 'none', 'nan')
    else:
        is_sent_folder = bool(is_sent_folder)

    body_text = str(raw_email.get('body', '') or '').strip()
    body_current = clean_email_body_for_analysis(body_text)

    email = {
        'subject': str(raw_email.get('subject', '') or '').strip(),
        'body': body_text,
        'body_current': body_current,
        'from': str(raw_email.get('from', raw_email.get('sender', raw_email.get('sender_email', ''))) or '').strip(),
        'to': str(raw_email.get('to', '') or '').strip(),
        'cc': str(raw_email.get('cc', '') or '').strip(),
        'attachments': attachment_text,
        'has_attachments': has_attachments,
        'date': str(raw_email.get('date', '') or '').strip(),
        'folder': str(raw_email.get('folder', '') or '').strip(),
        'root_folder': str(raw_email.get('root_folder', '') or '').strip(),
        'store': str(raw_email.get('store', '') or '').strip(),
        'is_sent_folder': is_sent_folder,
        'entry_id': str(raw_email.get('entry_id', '') or '').strip(),
        'conversation_id': str(raw_email.get('conversation_id', '') or '').strip(),
        'conversation_topic': str(raw_email.get('conversation_topic', '') or '').strip(),
    }

    sent_value = str(raw_email.get('date_sent', raw_email.get('sent', '')) or '').strip()
    received_value = str(raw_email.get('date_received', raw_email.get('received', '')) or '').strip()
    if sent_value:
        email['date_sent'] = sent_value
    else:
        email['date_sent'] = ''
    if received_value:
        email['date_received'] = received_value
    else:
        email['date_received'] = ''
    if not email['date']:
        email['date'] = email['date_received'] or email['date_sent']

    email['searchable'] = " ".join([
        email['subject'],
        email['body_current'] or email['body'],
        email['from'],
        email['to'],
        email['cc'],
        email['attachments'],
        email['folder'],
        email['root_folder'],
        email['store'],
        email['conversation_topic'],
    ]).lower()
    return email


def parse_email_csv(csv_path):
    """
    Parse Outlook email CSV export.

    Returns list of email dicts with normalized metadata for LLM analysis:
    subject, body, from, to, cc, date, attachments, has_attachments, searchable
    """
    try:
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                df = pd.read_csv(csv_path, encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError("Could not decode CSV file")

        # Normalize column names
        df.columns = df.columns.str.lower().str.strip()

        # Map common Outlook/Graph export column names.
        column_mapping = {
            'subject': ['subject', 'email subject', 'title'],
            'body': ['body', 'content', 'message', 'email body', 'notes'],
            'from': ['from', 'sender', 'from email', 'from address'],
            'to': ['to', 'recipient', 'to email', 'to address', 'recipients'],
            'cc': ['cc', 'cc recipients', 'cc list'],
            'attachments': ['attachments', 'attachment', 'files', 'file names', 'attachment names'],
            'has_attachments': ['has attachments', 'has attachment', 'attachments?', 'hasattachments'],
            'date': ['date', 'sent', 'received', 'date sent', 'date received', 'sent date', 'received time', 'sent time'],
        }

        emails = []
        for _, row in df.iterrows():
            raw_email = {}
            for field, possible_cols in column_mapping.items():
                for col in possible_cols:
                    if col in df.columns and pd.notna(row.get(col)):
                        raw_value = str(row[col]).strip()
                        if field == 'has_attachments':
                            raw_email[field] = raw_value.lower() not in ('', '0', 'false', 'no', 'none', 'nan')
                        else:
                            raw_email[field] = raw_value
                        break
                else:
                    raw_email[field] = False if field == 'has_attachments' else ''

            email = normalize_email_record(raw_email)
            if email:
                emails.append(email)

        return emails

    except Exception as e:
        print(f"Error parsing email CSV: {e}", file=sys.stderr)
        return []


def parse_email_json(json_path):
    """Parse JSON email activity written by the Electron app."""
    try:
        with open(json_path, 'r', encoding='utf-8') as handle:
            payload = json.load(handle)
    except Exception as e:
        print(f"Error parsing email JSON: {e}", file=sys.stderr)
        return []

    if isinstance(payload, dict):
        raw_emails = payload.get('emails', [])
    elif isinstance(payload, list):
        raw_emails = payload
    else:
        raw_emails = []

    emails = []
    for raw_email in raw_emails:
        email = normalize_email_record(raw_email)
        if email:
            emails.append(email)
    return emails


def parse_email_input(email_input_path):
    """Load email activity from either CSV export or JSON dataset."""
    ext = os.path.splitext(str(email_input_path or ''))[1].lower()
    if ext == '.json':
        return parse_email_json(email_input_path)
    return parse_email_csv(email_input_path)


def find_column_index(headers, patterns):
    """Find column index matching any of the patterns."""
    for i, header in enumerate(headers):
        header_lower = header.lower().strip()
        for pattern in patterns:
            if re.match(pattern, header_lower):
                return i
    return -1


def parse_json_response_text(response_text):
    """
    Parse JSON from plain text or markdown fenced code block output.
    """
    text = (response_text or "").strip()
    if not text:
        raise ValueError("LLM returned empty response")

    if text.startswith("```"):
        lines = text.split("\n")
        json_lines = []
        in_block = False
        for line in lines:
            marker = line.strip()
            if marker.startswith("```"):
                if not in_block:
                    in_block = True
                    continue
                break
            if in_block:
                json_lines.append(line)
        if json_lines:
            text = "\n".join(json_lines).strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        fragment = extract_json_object_fragment(text)
        if fragment:
            return json.loads(fragment)
        raise ValueError("LLM response was not valid JSON.")


def extract_json_object_fragment(text):
    """Extract the first complete JSON object from a larger text blob."""
    source = str(text or "")
    start = source.find("{")
    if start == -1:
        return None

    depth = 0
    in_string = False
    escaped = False

    for idx in range(start, len(source)):
        ch = source[idx]

        if in_string:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == '"':
                in_string = False
            continue

        if ch == '"':
            in_string = True
            continue
        if ch == "{":
            depth += 1
            continue
        if ch == "}":
            depth -= 1
            if depth == 0:
                return source[start:idx + 1]

    return None


def normalize_llm_status(raw_status):
    """
    Normalize model status text to one of the supported checklist statuses.
    """
    if not raw_status:
        return None

    status_text = str(raw_status).strip()
    if status_text in LLM_STATUS_VALUES:
        return status_text

    alias_key = status_text.lower()
    normalized = STATUS_ALIASES.get(alias_key)
    if normalized in LLM_STATUS_VALUES:
        return normalized

    return None


def canonical_doc_key(value):
    """Normalize document names for key matching."""
    text = re.sub(r'\s+', ' ', str(value or '')).strip().lower()
    return text


def normalize_api_key(api_key):
    return str(api_key or "").strip()


def normalize_provider(provider):
    value = str(provider or "").strip().lower()
    if value == "claude":
        return "anthropic"
    if value in ("managed", "localai", "anthropic", "openai", "harvey", "ollama", "lmstudio"):
        return value
    if value in ("managed ai", "managed-ai", "managed remote", "managed-remote", "firm", "firm-managed", "firmmanaged"):
        return "managed"
    if value in ("local", "local ai", "local-ai", "local pack", "localpack", "built-in", "builtin"):
        return "localai"
    if value in ("lm studio", "lm-studio"):
        return "lmstudio"
    return "anthropic"


def infer_provider_from_api_key(api_key):
    value = normalize_api_key(api_key).lower()
    if not value:
        return None
    if value.startswith("sk-ant-"):
        return "anthropic"
    if value.startswith("harvey_") or value.startswith("hv_"):
        return "harvey"
    if value.startswith("sk-proj-") or value.startswith("sk-") or value.startswith("sess-"):
        return "openai"
    return None


def resolve_provider(provider, api_key):
    preferred = normalize_provider(provider)
    if preferred in ("managed", "localai", "ollama", "lmstudio"):
        return preferred
    inferred = infer_provider_from_api_key(api_key)
    return inferred or preferred


def provider_requires_api_key(provider):
    return normalize_provider(provider) in ("anthropic", "openai", "harvey")


def extract_openai_text(content):
    if isinstance(content, str):
        return content.strip()
    if not isinstance(content, list):
        return ""
    parts = []
    for item in content:
        if isinstance(item, str):
            parts.append(item)
            continue
        if isinstance(item, dict):
            text = item.get("text")
            if isinstance(text, str):
                parts.append(text)
    return "\n".join(parts).strip()


def parse_error_detail(raw_text, network_error=None):
    if network_error:
        return str(network_error)

    try:
        payload = json.loads(raw_text or "{}")
    except Exception:
        payload = {}

    if isinstance(payload, dict):
        err = payload.get("error")
        if isinstance(err, dict) and err.get("message"):
            return str(err.get("message"))
        if payload.get("message"):
            return str(payload.get("message"))

    return (raw_text or "").strip()[:300] or "Unknown error"


def should_retry_request(status_code, detail):
    detail_text = str(detail or "").lower()
    if status_code in RETRYABLE_HTTP_STATUS:
        return True
    return any(marker in detail_text for marker in RETRYABLE_ERROR_MARKERS)


def backoff_sleep(attempt_idx):
    # Short incremental backoff for transient API/network failures.
    time.sleep(min(0.5 * (attempt_idx + 1), 2.0))


def format_provider_error(provider_label, status_code, detail):
    if status_code in (401, 403):
        return (
            f"{provider_label} authentication failed ({status_code}). "
            f"Check API key and account/model access. Details: {detail}"
        )
    if status_code == 429:
        return f"{provider_label} rate limit reached (429). Please retry. Details: {detail}"
    if status_code == 0:
        return f"{provider_label} network error: {detail}"
    return f"{provider_label} error ({status_code}): {detail}"


def is_likely_model_error(status_code, detail):
    detail_text = str(detail or "").lower()
    if status_code == 404:
        return True
    if status_code == 400 and "model" in detail_text:
        return True
    return "model" in detail_text and any(
        part in detail_text
        for part in ("not found", "invalid", "unsupported", "available", "access")
    )


def perform_http_request(url, headers, body_bytes, timeout=90):
    req = urllib.request.Request(url, data=body_bytes, headers=headers, method="POST")
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            status_code = resp.getcode()
            raw_text = resp.read().decode("utf-8", errors="replace")
            return status_code, raw_text, None
    except urllib.error.HTTPError as e:
        raw_text = e.read().decode("utf-8", errors="replace")
        return e.code, raw_text, None
    except urllib.error.URLError as e:
        reason = getattr(e, "reason", e)
        return 0, "", f"Network error: {reason}"
    except (socket.timeout, TimeoutError):
        return 0, "", "Network error: request timed out"
    except Exception as e:
        return 0, "", f"Network error: {e}"


def perform_json_request(url, headers, payload, timeout=90):
    request_headers = dict(headers or {})
    request_headers["Content-Type"] = "application/json"
    body_bytes = json.dumps(payload).encode("utf-8")
    status_code, raw_text, network_error = perform_http_request(
        url, request_headers, body_bytes, timeout=timeout
    )
    try:
        parsed = json.loads(raw_text) if raw_text else {}
    except Exception:
        parsed = {}
    return status_code, parsed, raw_text, network_error


def build_multipart_form_data(fields):
    boundary = "----EmmaNeighBoundary" + uuid.uuid4().hex
    chunks = []
    for key, value in fields.items():
        chunks.append(f"--{boundary}".encode("utf-8"))
        chunks.append(
            f'Content-Disposition: form-data; name="{key}"'.encode("utf-8")
        )
        chunks.append(b"")
        chunks.append(str(value).encode("utf-8"))
    chunks.append(f"--{boundary}--".encode("utf-8"))
    body = b"\r\n".join(chunks) + b"\r\n"
    return boundary, body


def call_anthropic_prompt(prompt, api_key, model_name):
    api_key = normalize_api_key(api_key)
    model_candidates = []
    for candidate in (
        model_name,
        os.environ.get("CLAUDE_MODEL"),
        DEFAULT_CLAUDE_MODEL,
        "claude-3-5-sonnet-20241022",
        "claude-3-5-haiku-20241022",
    ):
        candidate_value = str(candidate or "").strip()
        if candidate_value and candidate_value not in model_candidates:
            model_candidates.append(candidate_value)

    last_detail = "Unknown error"
    for candidate in model_candidates:
        model_unavailable = False
        for attempt_idx in range(MAX_HTTP_ATTEMPTS):
            payload = {
                "model": candidate,
                "max_tokens": 2048,
                "temperature": 0,
                "messages": [{"role": "user", "content": prompt}],
            }
            status_code, parsed, raw_text, network_error = perform_json_request(
                "https://api.anthropic.com/v1/messages",
                {
                    "x-api-key": api_key,
                    "anthropic-version": "2023-06-01",
                },
                payload,
            )

            if status_code != 200:
                detail = parse_error_detail(raw_text, network_error=network_error)
                last_detail = detail
                if should_retry_request(status_code, detail) and attempt_idx < (MAX_HTTP_ATTEMPTS - 1):
                    backoff_sleep(attempt_idx)
                    continue
                if is_likely_model_error(status_code, detail):
                    model_unavailable = True
                    break
                raise RuntimeError(format_provider_error("Anthropic", status_code, detail))

            content = parsed.get("content") if isinstance(parsed, dict) else None
            if isinstance(content, list) and content:
                first = content[0]
                if isinstance(first, dict) and first.get("text"):
                    return str(first.get("text")).strip(), candidate
            raise RuntimeError("Anthropic response was missing message content.")

        if model_unavailable:
            continue

    raise RuntimeError(
        f"No supported Claude model is available for this API key. Last error: {last_detail}"
    )


def call_openai_prompt(prompt, api_key, model_name):
    api_key = normalize_api_key(api_key)
    model_candidates = []
    for candidate in (
        model_name,
        os.environ.get("OPENAI_MODEL"),
        DEFAULT_OPENAI_MODEL,
        "gpt-4.1-mini",
        "gpt-4o-mini",
        "gpt-4.1",
        "gpt-4o",
    ):
        candidate_value = str(candidate or "").strip()
        if candidate_value and candidate_value not in model_candidates:
            model_candidates.append(candidate_value)

    last_detail = "Unknown error"
    for candidate in model_candidates:
        model_unavailable = False
        for attempt_idx in range(MAX_HTTP_ATTEMPTS):
            payload = {
                "model": candidate,
                "max_tokens": 2048,
                "temperature": 0,
                "messages": [{"role": "user", "content": prompt}],
            }
            status_code, parsed, raw_text, network_error = perform_json_request(
                "https://api.openai.com/v1/chat/completions",
                {"Authorization": f"Bearer {api_key}"},
                payload,
            )

            if status_code != 200:
                detail = parse_error_detail(raw_text, network_error=network_error)
                last_detail = detail
                if should_retry_request(status_code, detail) and attempt_idx < (MAX_HTTP_ATTEMPTS - 1):
                    backoff_sleep(attempt_idx)
                    continue
                if is_likely_model_error(status_code, detail):
                    model_unavailable = True
                    break
                raise RuntimeError(format_provider_error("OpenAI", status_code, detail))

            choices = parsed.get("choices") if isinstance(parsed, dict) else None
            message = choices[0].get("message") if isinstance(choices, list) and choices else {}
            content = extract_openai_text(message.get("content") if isinstance(message, dict) else "")
            if content:
                return content, candidate
            raise RuntimeError("OpenAI response was missing message content.")

        if model_unavailable:
            continue

    raise RuntimeError(
        f"No supported OpenAI model is available for this API key. Last error: {last_detail}"
    )


def call_openai_compatible_prompt(prompt, api_key, model_candidates, base_url, provider_label):
    api_key = normalize_api_key(api_key)
    candidates = []
    for candidate in model_candidates:
        candidate_value = str(candidate or "").strip()
        if candidate_value and candidate_value not in candidates:
            candidates.append(candidate_value)

    last_detail = "Unknown error"
    for candidate in candidates:
        model_unavailable = False
        for attempt_idx in range(MAX_HTTP_ATTEMPTS):
            payload = {
                "model": candidate,
                "max_tokens": 2048,
                "temperature": 0,
                "messages": [{"role": "user", "content": prompt}],
            }
            headers = {}
            if api_key:
                headers["Authorization"] = f"Bearer {api_key}"
            status_code, parsed, raw_text, network_error = perform_json_request(
                f"{base_url}/v1/chat/completions",
                headers,
                payload,
            )

            if status_code != 200:
                detail = parse_error_detail(raw_text, network_error=network_error)
                last_detail = detail
                if should_retry_request(status_code, detail) and attempt_idx < (MAX_HTTP_ATTEMPTS - 1):
                    backoff_sleep(attempt_idx)
                    continue
                if is_likely_model_error(status_code, detail):
                    model_unavailable = True
                    break
                raise RuntimeError(format_provider_error(provider_label, status_code, detail))

            choices = parsed.get("choices") if isinstance(parsed, dict) else None
            message = choices[0].get("message") if isinstance(choices, list) and choices else {}
            content = extract_openai_text(message.get("content") if isinstance(message, dict) else "")
            if content:
                return content, candidate
            raise RuntimeError(f"{provider_label} response was missing message content.")

        if model_unavailable:
            continue

    raise RuntimeError(
        f"No supported {provider_label} model is available. Last error: {last_detail}"
    )


def call_localai_prompt(prompt, api_key, model_name, provider_base_url):
    base_url = str(provider_base_url or os.environ.get("LOCALAI_BASE_URL") or DEFAULT_LOCALAI_BASE_URL).strip()
    if not base_url:
        base_url = DEFAULT_LOCALAI_BASE_URL
    if base_url.endswith("/"):
        base_url = base_url[:-1]
    return call_openai_compatible_prompt(
        prompt,
        api_key,
        [
            model_name,
            os.environ.get("LOCALAI_MODEL"),
            DEFAULT_LOCALAI_MODEL,
            "qwen2.5:1.5b",
            "qwen2.5:7b",
            "gemma3:1b",
            "gemma3",
        ],
        base_url,
        "Local AI",
    )


def call_ollama_prompt(prompt, api_key, model_name, provider_base_url):
    base_url = str(provider_base_url or os.environ.get("OLLAMA_BASE_URL") or DEFAULT_OLLAMA_BASE_URL).strip()
    if not base_url:
        base_url = DEFAULT_OLLAMA_BASE_URL
    if base_url.endswith("/"):
        base_url = base_url[:-1]
    return call_openai_compatible_prompt(
        prompt,
        api_key,
        [
            model_name,
            os.environ.get("OLLAMA_MODEL"),
            DEFAULT_OLLAMA_MODEL,
            "llama3.1:8b",
            "qwen2.5:7b",
        ],
        base_url,
        "Ollama",
    )


def call_lmstudio_prompt(prompt, api_key, model_name, provider_base_url):
    base_url = str(provider_base_url or os.environ.get("LMSTUDIO_BASE_URL") or DEFAULT_LMSTUDIO_BASE_URL).strip()
    if not base_url:
        base_url = DEFAULT_LMSTUDIO_BASE_URL
    if base_url.endswith("/"):
        base_url = base_url[:-1]
    return call_openai_compatible_prompt(
        prompt,
        api_key,
        [
            model_name,
            os.environ.get("LMSTUDIO_MODEL"),
            DEFAULT_LMSTUDIO_MODEL,
        ],
        base_url,
        "LM Studio",
    )


def call_managed_prompt(prompt, api_key, model_name, provider_base_url):
    endpoint = str(provider_base_url or os.environ.get("MANAGED_AI_PROMPT_URL") or "").strip()
    if not endpoint:
        raise RuntimeError("Managed AI prompt endpoint is not configured.")

    headers = {}
    token = normalize_api_key(api_key)
    if token:
        headers["Authorization"] = f"Bearer {token}"

    payload = {
        "prompt": prompt,
        "max_tokens": 2048,
    }
    if model_name:
        payload["model"] = str(model_name).strip()

    last_detail = "Unknown error"
    for attempt_idx in range(MAX_HTTP_ATTEMPTS):
        status_code, parsed, raw_text, network_error = perform_json_request(
            endpoint,
            headers,
            payload,
        )

        if status_code != 200:
            detail = parse_error_detail(raw_text, network_error=network_error)
            last_detail = detail
            if should_retry_request(status_code, detail) and attempt_idx < (MAX_HTTP_ATTEMPTS - 1):
                backoff_sleep(attempt_idx)
                continue
            raise RuntimeError(format_provider_error("Managed AI", status_code, detail))

        text = ""
        if isinstance(parsed, dict):
            text = str(parsed.get("text") or parsed.get("message") or "").strip()
        if text:
            return text, str(parsed.get("model") or model_name or "managed-ai").strip()
        raise RuntimeError("Managed AI response was missing message content.")

    raise RuntimeError(f"Managed AI error: {last_detail}")


def call_harvey_prompt(prompt, api_key, max_tokens, provider_base_url):
    api_key = normalize_api_key(api_key)
    base_url = str(
        provider_base_url or os.environ.get("HARVEY_BASE_URL") or DEFAULT_HARVEY_BASE_URL
    ).strip()
    if not base_url:
        base_url = DEFAULT_HARVEY_BASE_URL
    if base_url.endswith("/"):
        base_url = base_url[:-1]

    boundary, body = build_multipart_form_data(
        {
            "prompt": prompt,
            "mode": "assist",
            "stream": "false",
            "max_tokens": str(max_tokens),
        }
    )

    parsed = {}
    for attempt_idx in range(MAX_HTTP_ATTEMPTS):
        status_code, raw_text, network_error = perform_http_request(
            f"{base_url}/api/v2/completion?include_citations=false",
            {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": f"multipart/form-data; boundary={boundary}",
            },
            body,
        )
        detail = parse_error_detail(raw_text, network_error=network_error)
        if status_code != 200:
            if should_retry_request(status_code, detail) and attempt_idx < (MAX_HTTP_ATTEMPTS - 1):
                backoff_sleep(attempt_idx)
                continue
            raise RuntimeError(format_provider_error("Harvey", status_code, detail))

        try:
            parsed = json.loads(raw_text) if raw_text else {}
        except Exception:
            parsed = {}
        break

    text = ""
    if isinstance(parsed, dict):
        text = parsed.get("response") or parsed.get("text") or ""
    if not text:
        raise RuntimeError("Harvey response was missing message content.")

    return str(text).strip(), "harvey-assist"


def call_provider_prompt_json(prompt, api_key, provider, model_name=None, provider_base_url=None):
    provider_name = resolve_provider(provider, api_key)
    if provider_name == "managed":
        response_text, _ = call_managed_prompt(prompt, api_key, model_name, provider_base_url)
    elif provider_name == "localai":
        response_text, _ = call_localai_prompt(prompt, api_key, model_name, provider_base_url)
    elif provider_name == "openai":
        response_text, _ = call_openai_prompt(prompt, api_key, model_name)
    elif provider_name == "harvey":
        response_text, _ = call_harvey_prompt(prompt, api_key, 2048, provider_base_url)
    elif provider_name == "ollama":
        response_text, _ = call_ollama_prompt(prompt, api_key, model_name, provider_base_url)
    elif provider_name == "lmstudio":
        response_text, _ = call_lmstudio_prompt(prompt, api_key, model_name, provider_base_url)
    else:
        response_text, _ = call_anthropic_prompt(prompt, api_key, model_name)

    parsed = parse_json_response_text(response_text)
    if not isinstance(parsed, dict):
        raise ValueError("LLM response was not a JSON object.")
    return parsed


def normalize_cell_text(value):
    return re.sub(r'\s+', ' ', str(value or '')).strip()


def normalize_searchable_lookup_text(value):
    return re.sub(r'[^a-z0-9]+', ' ', normalize_cell_text(value).lower()).strip()


def normalize_document_lookup_text(value):
    text = os.path.splitext(os.path.basename(str(value or '')))[0]
    text = text.replace('_', ' ').replace('-', ' ')
    text = normalize_searchable_lookup_text(text)
    text = DOCUMENT_FILENAME_NOISE_RE.sub(' ', text)
    return re.sub(r'\s+', ' ', text).strip()


def normalize_bool(value, default=False):
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return value != 0
    normalized = str(value).strip().lower()
    if not normalized:
        return default
    if normalized in ('true', 'yes', 'y', '1', 'active', 'found'):
        return True
    if normalized in ('false', 'no', 'n', '0', 'none', 'not found'):
        return False
    return default


def extract_search_tokens(text):
    tokens = []
    seen = set()
    for token in re.findall(r'[a-z0-9]+', normalize_searchable_lookup_text(text)):
        if len(token) < 3 or token in SEARCH_TOKEN_STOPWORDS:
            continue
        if token not in seen:
            seen.add(token)
            tokens.append(token)
    return tokens


def extract_deal_tokens(text):
    tokens = []
    seen = set()
    for token in re.findall(r'[a-z0-9]+', normalize_searchable_lookup_text(text)):
        if len(token) < 4 or token in DEAL_TOKEN_STOPWORDS:
            continue
        if token not in seen:
            seen.add(token)
            tokens.append(token)
    return tokens


def build_document_match_profile(text):
    document_name = normalize_cell_text(text)
    lookup = normalize_document_lookup_text(document_name)
    tokens = extract_search_tokens(lookup or document_name)
    return {
        'document_name': document_name,
        'lookup': lookup,
        'tokens': tokens,
    }


def score_attachment_title_match(document_name, attachment_names):
    profile = build_document_match_profile(document_name)
    doc_lookup = profile['lookup']
    doc_tokens = profile['tokens']
    best_score = 0.0
    matched_titles = []

    for attachment_name in attachment_names or []:
        title = normalize_cell_text(attachment_name)
        if not title:
            continue
        title_lookup = normalize_document_lookup_text(title)
        token_hits = [token for token in doc_tokens if token in title_lookup]
        local_score = 0.0

        if doc_lookup and (doc_lookup in title_lookup or title_lookup in doc_lookup):
            local_score = max(local_score, 10.5)
        if token_hits:
            local_score = max(local_score, min(8.8, len(token_hits) * 2.2))
            if len(token_hits) >= 2:
                local_score += 0.9

        if local_score > 0:
            best_score = max(best_score, local_score)
            matched_titles.append(title)

    return round(best_score, 4), dedupe_preserve_order(matched_titles)[:4]


def classify_checklist_row(doc_name, row_data, headers, doc_col_idx, status_col_idx, notes_col_idx, previous_document_name=''):
    cleaned_doc_name = normalize_cell_text(doc_name)
    if not cleaned_doc_name:
        return {
            'row_type': 'blank',
            'should_update': False,
            'match_document_name': '',
        }

    lookup = normalize_document_lookup_text(cleaned_doc_name)
    doc_tokens = extract_search_tokens(lookup or cleaned_doc_name)
    has_document_keyword = any(keyword in lookup for keyword in DOCUMENT_ROW_KEYWORDS)
    has_section_keyword = any(keyword in lookup for keyword in SECTION_ROW_KEYWORDS)
    starts_with_subrow_marker = bool(ROW_NUMBER_PREFIX_RE.match(cleaned_doc_name))

    non_empty_cells = [normalize_cell_text(cell) for cell in row_data if normalize_cell_text(cell)]
    non_doc_cells = [
        normalize_cell_text(cell)
        for idx, cell in enumerate(row_data)
        if idx not in (doc_col_idx, status_col_idx, notes_col_idx) and normalize_cell_text(cell)
    ]
    short_generic_label = len(doc_tokens) <= 2 and not has_document_keyword
    likely_section_header = (
        len(non_empty_cells) <= 2 and len(non_doc_cells) <= 1 and (
            cleaned_doc_name.endswith(':')
            or has_section_keyword
            or (cleaned_doc_name.isupper() and len(cleaned_doc_name) <= 80)
        )
    )

    if likely_section_header:
        row_type = 'section_header'
    elif previous_document_name and (starts_with_subrow_marker or (short_generic_label and has_section_keyword)):
        row_type = 'subrow'
    elif previous_document_name and short_generic_label and len(non_doc_cells) <= 2:
        row_type = 'subrow'
    else:
        row_type = 'document'

    return {
        'row_type': row_type,
        'should_update': row_type == 'document',
        'match_document_name': cleaned_doc_name if row_type == 'document' else normalize_cell_text(previous_document_name),
    }


def dedupe_preserve_order(values):
    items = []
    seen = set()
    for value in values:
        cleaned = normalize_cell_text(value)
        if not cleaned:
            continue
        key = cleaned.lower()
        if key in seen:
            continue
        seen.add(key)
        items.append(cleaned)
    return items


def split_attachment_names(raw_value):
    return dedupe_preserve_order(re.split(r'[;,]', str(raw_value or '')))


def split_recipients(raw_value):
    return dedupe_preserve_order(re.split(r';', str(raw_value or '')))


def parse_datetime_value(value):
    text = normalize_cell_text(value)
    if not text:
        return None
    for candidate in (text, text.replace('Z', '+00:00')):
        try:
            return datetime.fromisoformat(candidate)
        except Exception:
            continue
    try:
        parsed = pd.to_datetime(text, errors='coerce')
        if pd.isna(parsed):
            return None
        return parsed.to_pydatetime() if hasattr(parsed, 'to_pydatetime') else None
    except Exception:
        return None


def format_date_value(value):
    parsed = parse_datetime_value(value)
    if parsed is not None:
        try:
            return parsed.strftime('%m/%d/%Y')
        except Exception:
            pass
    text = normalize_cell_text(value)
    return text[:10] if text else ''


def strip_thread_subject_markers(subject):
    text = normalize_cell_text(subject)
    previous = None
    while text and text != previous:
        previous = text
        text = THREAD_SUBJECT_PREFIX_RE.sub('', text)
        text = THREAD_SUBJECT_TAG_RE.sub('', text)
        text = normalize_cell_text(text)
    return text


def normalize_thread_subject(subject):
    return strip_thread_subject_markers(subject).lower()


def clean_email_body_for_analysis(body):
    lines = []
    header_lines_seen = 0
    for raw_line in str(body or '').splitlines():
        line = raw_line.strip()
        if not line:
            continue
        lower_line = line.lower()
        if line.startswith('>'):
            continue
        if lower_line.startswith('-----original message-----'):
            break
        if re.match(r'^on .+ wrote:$', lower_line):
            break
        if re.match(r'^(from|sent|to|cc|subject):', lower_line):
            header_lines_seen += 1
            if header_lines_seen >= 2:
                break
            continue
        if 'privileged and confidential' in lower_line or 'confidentiality notice' in lower_line:
            break
        lines.append(line)
        if len(lines) >= 80:
            break
    return " ".join(lines)


def extract_issue_snippets(text, limit=MAX_THREAD_ISSUES_PER_SUMMARY):
    cleaned = clean_email_body_for_analysis(text)
    if not cleaned:
        return []

    snippets = []
    for piece in re.split(r'(?<=[.!?])\s+|\n+', cleaned):
        snippet = normalize_cell_text(piece)
        if len(snippet) < 25 or len(snippet) > 240:
            continue
        lower_snippet = snippet.lower()
        if re.match(r'^(from|sent|to|cc|subject):', lower_snippet):
            continue
        if any(keyword in lower_snippet for keyword in ISSUE_KEYWORDS):
            snippets.append(snippet)
    return dedupe_preserve_order(snippets)[:limit]


def infer_email_direction(email):
    if email.get('is_sent_folder'):
        return 'sent'
    root_folder = normalize_cell_text(email.get('root_folder') or email.get('folder')).lower()
    sender = normalize_cell_text(email.get('from')).lower()
    if 'sent items' in root_folder or root_folder.endswith('sent items'):
        return 'sent'
    if sender in ('me', 'myself'):
        return 'sent'
    if email.get('date_sent') and not email.get('date_received'):
        return 'sent'
    return 'received'


def short_party_label(text, fallback='someone'):
    cleaned = normalize_cell_text(text)
    if not cleaned:
        return fallback
    primary = re.split(r';', cleaned)[0].strip()
    primary = re.sub(r'\s+', ' ', primary)
    return primary[:80] if primary else fallback


def highest_status_from_text(text):
    searchable = str(text or '').lower()
    best_status = None
    best_priority = 0
    for status_config in STATUS_PATTERNS:
        for pattern in status_config['patterns']:
            if re.search(pattern, searchable):
                if status_config['priority'] > best_priority:
                    best_status = status_config['status']
                    best_priority = status_config['priority']
                break
    return best_status, best_priority


def classify_email_activity(email):
    searchable = str(email.get('searchable', '') or '').lower()
    status, priority = highest_status_from_text(searchable)
    has_document_activity = bool(email.get('has_attachments')) or any(
        keyword in searchable for keyword in DOCUMENT_ACTIVITY_KEYWORDS
    )
    has_draft_language = 'draft' in searchable or ('revised' in searchable and bool(email.get('has_attachments')))
    has_markup_language = any(
        marker in searchable for marker in ('redline', 'blackline', 'markup', 'comments')
    )

    if status == 'Executed':
        label = 'executed copy'
    elif status == 'Execution Version':
        label = 'execution version'
    elif has_draft_language:
        label = 'draft'
    elif has_markup_language:
        label = 'markup/comments'
    elif 'draft' in searchable or has_document_activity:
        label = 'draft'
    else:
        label = ''

    return {
        'status': status,
        'priority': priority,
        'has_document_activity': has_document_activity or bool(status),
        'label': label,
        'issue_snippets': extract_issue_snippets(email.get('body', '')),
    }


def build_email_activity_event(email):
    activity = classify_email_activity(email)
    direction = infer_email_direction(email)
    counterparty = short_party_label(
        email.get('to') if direction == 'sent' else email.get('from'),
        fallback='counterparty',
    )
    date_text = format_date_value(email.get('date_received') or email.get('date_sent') or email.get('date'))
    actor = 'Sent' if direction == 'sent' else 'Received'
    preposition = 'to' if direction == 'sent' else 'from'
    object_label = activity.get('label') or 'document update'

    parts = [actor, object_label, preposition, counterparty]
    sentence = " ".join(part for part in parts if part)
    if date_text:
        sentence = f"{sentence} on {date_text}"
    sentence = sentence.strip()
    if sentence and not sentence.endswith('.'):
        sentence += '.'

    return {
        'text': sentence,
        'direction': direction,
        'party': counterparty,
        'date': date_text,
        'status': activity.get('status'),
        'priority': activity.get('priority', 0),
        'has_document_activity': activity.get('has_document_activity', False),
        'issue_snippets': activity.get('issue_snippets', []),
    }


def build_deal_profile(checklist_path, checklist_items):
    base_name = os.path.splitext(os.path.basename(str(checklist_path or '')))[0]
    anchor_tokens = extract_deal_tokens(base_name)

    token_counts = {}
    for item in checklist_items[:250]:
        combined = f"{item.get('document_name', '')} {item.get('row_context', '')}"
        for token in extract_deal_tokens(combined):
            token_counts[token] = token_counts.get(token, 0) + 1

    shared_tokens = [
        token for token, count in sorted(token_counts.items(), key=lambda item: (-item[1], item[0]))
        if count >= 2 and token not in anchor_tokens
    ]
    anchor_tokens.extend(shared_tokens[:6])

    return {
        'base_name': base_name,
        'anchor_tokens': anchor_tokens[:12],
    }


def build_thread_key(email, fallback_index):
    conversation_id = normalize_cell_text(email.get('conversation_id'))
    if conversation_id:
        return f"conversation:{conversation_id.lower()}"

    conversation_topic = normalize_thread_subject(email.get('conversation_topic'))
    if conversation_topic:
        return f"topic:{conversation_topic}"

    subject_key = normalize_thread_subject(email.get('subject'))
    if subject_key:
        return f"subject:{subject_key}"

    attachment_names = split_attachment_names(email.get('attachments'))
    if attachment_names:
        return f"attachment:{canonical_doc_key(os.path.splitext(attachment_names[0])[0])}"

    entry_id = normalize_cell_text(email.get('entry_id'))
    if entry_id:
        return f"entry:{entry_id.lower()}"

    return f"email:{fallback_index}"


def build_email_threads(emails):
    grouped = {}

    for idx, raw_email in enumerate(emails):
        email = dict(raw_email or {})
        email['index'] = idx
        email['sort_date'] = parse_datetime_value(
            email.get('date_received') or email.get('date_sent') or email.get('date')
        )
        email['thread_subject'] = strip_thread_subject_markers(email.get('subject'))
        email['activity_event'] = build_email_activity_event(email)
        thread_key = build_thread_key(email, idx)
        grouped.setdefault(thread_key, []).append(email)

    threads = []
    for thread_idx, (thread_key, thread_emails) in enumerate(grouped.items()):
        thread_emails.sort(key=lambda item: (
            item.get('sort_date') or datetime.min,
            item.get('index', 0),
        ))
        latest_subject = next(
            (normalize_cell_text(item.get('subject')) for item in reversed(thread_emails) if normalize_cell_text(item.get('subject'))),
            thread_emails[-1].get('thread_subject', '') if thread_emails else ''
        )
        aggregate_searchable = " ".join(str(item.get('searchable', '') or '') for item in thread_emails)
        attachment_names = []
        participants = []
        folders = []
        root_folders = []
        all_issue_snippets = []
        status_signal = None
        status_priority = 0

        for email in thread_emails:
            attachment_names.extend(split_attachment_names(email.get('attachments')))
            participants.append(email.get('from'))
            participants.extend(split_recipients(email.get('to')))
            participants.extend(split_recipients(email.get('cc')))
            folders.append(email.get('folder'))
            root_folders.append(email.get('root_folder'))
            all_issue_snippets.extend(email.get('activity_event', {}).get('issue_snippets', []))
            email_priority = email.get('activity_event', {}).get('priority', 0)
            if email_priority > status_priority:
                status_priority = email_priority
                status_signal = email.get('activity_event', {}).get('status')

        first_email = thread_emails[0] if thread_emails else {}
        last_email = thread_emails[-1] if thread_emails else {}
        latest_body_excerpt = clean_email_body_for_analysis(last_email.get('body_current') or last_email.get('body'))
        threads.append({
            'thread_id': f"thread-{thread_idx + 1}",
            'thread_key': thread_key,
            'subject': latest_subject or first_email.get('thread_subject') or '',
            'normalized_subject': normalize_thread_subject(latest_subject or first_email.get('thread_subject') or ''),
            'searchable': aggregate_searchable.lower(),
            'emails': thread_emails,
            'message_count': len(thread_emails),
            'folders': dedupe_preserve_order(folders),
            'root_folders': dedupe_preserve_order(root_folders),
            'participants': dedupe_preserve_order(participants)[:8],
            'attachment_names': dedupe_preserve_order(attachment_names),
            'attachments_text': " ".join(attachment_names).lower(),
            'first_date': format_date_value(first_email.get('date_received') or first_email.get('date_sent') or first_email.get('date')),
            'last_date': format_date_value(last_email.get('date_received') or last_email.get('date_sent') or last_email.get('date')),
            'latest_attachment_names': split_attachment_names(last_email.get('attachments')),
            'latest_body_excerpt': latest_body_excerpt[:260],
            'status_signal': status_signal,
            'status_priority': status_priority,
            'issue_snippets': dedupe_preserve_order(all_issue_snippets)[:MAX_THREAD_ISSUES_PER_SUMMARY],
        })

    threads.sort(
        key=lambda thread: (
            parse_datetime_value(thread.get('last_date')) or datetime.min,
            thread.get('thread_id', ''),
        ),
        reverse=True,
    )
    return threads


def score_email_candidate(checklist_item, email):
    searchable = normalize_searchable_lookup_text(email.get('searchable', ''))
    subject = normalize_searchable_lookup_text(email.get('subject', ''))
    body_text = normalize_searchable_lookup_text(email.get('body_current') or email.get('body'))
    attachments = split_attachment_names(email.get('attachments'))
    attachment_text = " ".join(normalize_document_lookup_text(title) for title in attachments)
    recipients = " ".join([
        str(email.get('from', '') or ''),
        str(email.get('to', '') or ''),
        str(email.get('cc', '') or ''),
    ])
    recipient_text = normalize_searchable_lookup_text(recipients)

    doc_name = normalize_cell_text(
        checklist_item.get('match_document_name') or checklist_item.get('document_name', '')
    )
    row_context = normalize_cell_text(checklist_item.get('row_context', ''))
    doc_profile = build_document_match_profile(doc_name)
    doc_lookup = doc_profile['lookup']
    doc_tokens = doc_profile['tokens']
    attachment_score, matched_attachments = score_attachment_title_match(doc_name, attachments)
    score = 0.0

    if attachment_score > 0:
        score += attachment_score
    elif doc_lookup and doc_lookup in subject:
        score += 6.5
    elif doc_lookup and (doc_lookup in body_text or doc_lookup in searchable):
        score += 3.2

    context_tokens = extract_search_tokens(row_context)

    doc_token_hits = 0
    for token in doc_tokens:
        if token in attachment_text:
            score += 2.4
            doc_token_hits += 1
        elif token in subject:
            score += 1.8
            doc_token_hits += 1
        elif token in body_text:
            score += 1.0
            doc_token_hits += 1
        elif token in searchable:
            score += 0.9
            doc_token_hits += 1

    if doc_token_hits >= 2:
        score += 1.6

    context_hits = 0
    for token in context_tokens[:8]:
        if token in subject or token in attachment_text:
            score += 0.9
            context_hits += 1
        elif token in recipient_text:
            score += 0.6
            context_hits += 1
        elif token in body_text or token in searchable:
            score += 0.35
            context_hits += 1

    if context_hits >= 2:
        score += 0.7
    if email.get('has_attachments'):
        score += 0.2
    if matched_attachments:
        score += min(0.8, 0.25 * len(matched_attachments))

    return score


def select_candidate_email_indices(checklist_item, emails, limit=MAX_CANDIDATE_EMAILS_PER_ROW):
    scored = []
    for idx, email in enumerate(emails):
        score = score_email_candidate(checklist_item, email)
        if score >= 1.2:
            scored.append((idx, round(score, 4)))

    scored.sort(key=lambda item: item[1], reverse=True)
    return scored[:limit]


def score_thread_candidate(checklist_item, thread, deal_profile=None):
    doc_name = normalize_cell_text(
        checklist_item.get('match_document_name') or checklist_item.get('document_name', '')
    )
    doc_profile = build_document_match_profile(doc_name)
    doc_lookup = doc_profile['lookup']
    doc_tokens = doc_profile['tokens']
    thread_subject = normalize_searchable_lookup_text(thread.get('subject', ''))
    thread_searchable = normalize_searchable_lookup_text(thread.get('searchable', ''))
    thread_attachments = " ".join(
        normalize_document_lookup_text(name) for name in thread.get('attachment_names', [])
    )
    relevant_email_scores = []
    attachment_score, matched_attachments = score_attachment_title_match(doc_name, thread.get('attachment_names', []))

    for email in thread.get('emails', []):
        base_score = score_email_candidate(checklist_item, email)
        if base_score >= 0.9:
            relevant_email_scores.append((email.get('index', -1), round(base_score, 4)))

    exact_doc_hit = bool(doc_lookup and (
        doc_lookup in thread_attachments or
        doc_lookup in thread_subject or
        doc_lookup in thread_searchable
    ))

    score = 0.0
    if relevant_email_scores:
        relevant_email_scores.sort(key=lambda item: item[1], reverse=True)
        score = relevant_email_scores[0][1]
        score += sum(max(item[1] - 0.75, 0.0) * 0.35 for item in relevant_email_scores[1:3])
        if len(relevant_email_scores) >= 2:
            score += 1.2
    elif exact_doc_hit:
        score += 4.5

    score += attachment_score
    attachment_hits = sum(1 for token in doc_tokens[:8] if token in thread_attachments)
    subject_hits = sum(1 for token in doc_tokens[:8] if token in thread_subject)
    body_hits = sum(1 for token in doc_tokens[:8] if token in thread_searchable)
    score += (attachment_hits * 0.9) + (subject_hits * 0.6) + (body_hits * 0.25)

    deal_anchor_hits = []
    if deal_profile and deal_profile.get('anchor_tokens'):
        for token in deal_profile.get('anchor_tokens', []):
            if token in thread_searchable:
                deal_anchor_hits.append(token)
        if deal_anchor_hits:
            score += min(2.5, 0.8 * len(deal_anchor_hits))
        elif score < 6.0:
            score -= 1.8

    score += min(1.0, 0.2 * max(0, int(thread.get('message_count', 0)) - 1))
    score += min(0.6, 0.2 * len(thread.get('issue_snippets', [])))
    if thread.get('status_priority', 0) > 0:
        score += min(0.8, 0.2 * thread.get('status_priority', 0))

    return {
        'score': round(score, 4),
        'relevant_email_indices': [idx for idx, _ in relevant_email_scores[:MAX_CANDIDATE_EMAILS_PER_ROW] if idx >= 0],
        'email_scores': relevant_email_scores[:MAX_CANDIDATE_EMAILS_PER_ROW],
        'deal_anchor_hits': deal_anchor_hits,
        'exact_doc_hit': exact_doc_hit,
        'matched_attachment_titles': matched_attachments,
    }


def derive_candidate_status(relevant_emails, thread):
    best_status = None
    best_priority = 0
    sent_count = 0
    received_count = 0
    issue_count = 0

    for email in relevant_emails:
        event = email.get('activity_event', {})
        priority = int(event.get('priority', 0) or 0)
        if priority > best_priority:
            best_priority = priority
            best_status = event.get('status')
        if event.get('direction') == 'sent':
            sent_count += 1
        else:
            received_count += 1
        issue_count += len(event.get('issue_snippets', []))

    if best_priority <= 1 and sent_count and (received_count or issue_count):
        return 'With Opposing Counsel'
    if best_status:
        return best_status
    if sent_count and (received_count or issue_count):
        return 'With Opposing Counsel'
    if sent_count and thread.get('message_count', 0) >= 1:
        return 'With Opposing Counsel'
    if received_count:
        return 'Draft Circulated'
    return thread.get('status_signal')


def compose_candidate_comment(events, issues):
    event_sentences = [normalize_cell_text(event) for event in events if normalize_cell_text(event)]
    issue_sentences = [normalize_cell_text(issue) for issue in issues if normalize_cell_text(issue)]

    comment = " ".join(event_sentences[:2]).strip()
    if issue_sentences:
        issues_text = "; ".join(issue_sentences[:2])
        comment = f"{comment} Issues flagged: {issues_text}.".strip() if comment else f"Issues flagged: {issues_text}."
    return comment.strip()


def summarize_thread_for_row(checklist_item, thread, score_info):
    relevant_index_set = set(score_info.get('relevant_email_indices', []))
    relevant_emails = [
        email for email in thread.get('emails', [])
        if email.get('index') in relevant_index_set
    ]
    if not relevant_emails and score_info.get('exact_doc_hit'):
        relevant_emails = thread.get('emails', [])[:MAX_THREAD_EVENTS_PER_SUMMARY]

    relevant_emails = relevant_emails[:MAX_THREAD_EVENTS_PER_SUMMARY]
    events = []
    issues = []
    matching_subjects = []
    relevant_attachment_titles = []
    for email in relevant_emails:
        event = email.get('activity_event', {})
        event_text = normalize_cell_text(event.get('text'))
        if event_text and event_text not in events:
            events.append(event_text)
        matching_subjects.append(normalize_cell_text(email.get('subject')) or 'No subject')
        relevant_attachment_titles.extend(
            score_attachment_title_match(
                checklist_item.get('match_document_name') or checklist_item.get('document_name', ''),
                split_attachment_names(email.get('attachments')),
            )[1]
        )
        issues.extend(event.get('issue_snippets', []))

    issues = dedupe_preserve_order(issues)[:MAX_THREAD_ISSUES_PER_SUMMARY]
    status_signal = derive_candidate_status(relevant_emails, thread)
    suggested_comment = compose_candidate_comment(events, issues)
    latest_relevant_email = relevant_emails[-1] if relevant_emails else {}

    return {
        'thread_id': thread.get('thread_id'),
        'score': score_info.get('score', 0.0),
        'subject': thread.get('subject', ''),
        'message_count': thread.get('message_count', 0),
        'folders': thread.get('folders', []),
        'participants': thread.get('participants', []),
        'first_date': thread.get('first_date', ''),
        'last_date': thread.get('last_date', ''),
        'deal_anchor_hits': score_info.get('deal_anchor_hits', []),
        'matching_attachment_titles': dedupe_preserve_order(
            score_info.get('matched_attachment_titles', []) + relevant_attachment_titles
        )[:4],
        'latest_attachment_titles': dedupe_preserve_order(thread.get('latest_attachment_names', []))[:4],
        'latest_body_excerpt': normalize_cell_text(
            clean_email_body_for_analysis(latest_relevant_email.get('body_current') or latest_relevant_email.get('body'))
            or thread.get('latest_body_excerpt', '')
        )[:260],
        'relevant_email_indices': score_info.get('relevant_email_indices', []),
        'matching_subjects': dedupe_preserve_order(matching_subjects)[:3],
        'events': events[:MAX_THREAD_EVENTS_PER_SUMMARY],
        'issues': issues,
        'status_signal': status_signal or '',
        'exact_doc_hit': bool(score_info.get('exact_doc_hit')),
        'suggested_comment': suggested_comment,
    }


def build_row_thread_candidates(checklist_path, checklist_items, emails):
    deal_profile = build_deal_profile(checklist_path, checklist_items)
    threads = build_email_threads(emails)
    candidate_map = {}
    thread_lookup = {thread.get('thread_id'): thread for thread in threads}

    for item in checklist_items:
        row_id = item.get('row_id')
        thread_candidates = []
        for thread in threads:
            score_info = score_thread_candidate(item, thread, deal_profile=deal_profile)
            if score_info.get('score', 0.0) < 2.0:
                continue
            thread_candidates.append(summarize_thread_for_row(item, thread, score_info))

        thread_candidates.sort(
            key=lambda candidate: (
                float(candidate.get('score', 0.0)),
                parse_datetime_value(candidate.get('last_date')) or datetime.min,
            ),
            reverse=True,
        )
        candidate_map[row_id] = thread_candidates[:MAX_CANDIDATE_THREADS_PER_ROW]

    return candidate_map, thread_lookup, deal_profile


def chunk_items(items, chunk_size):
    for start in range(0, len(items), chunk_size):
        yield items[start:start + chunk_size]


def format_email_evidence(email):
    if not isinstance(email, dict):
        return ''

    date_value = normalize_cell_text(
        email.get('date_received') or email.get('date_sent') or email.get('date')
    )
    if date_value:
        try:
            parsed = datetime.fromisoformat(date_value.replace('Z', '+00:00'))
            date_value = parsed.strftime('%m/%d/%Y')
        except Exception:
            date_value = date_value[:10]

    direction = ''
    sender = normalize_cell_text(email.get('from'))
    recipients = normalize_cell_text(email.get('to'))
    subject = normalize_cell_text(email.get('subject')) or 'No subject'
    attachments = normalize_cell_text(email.get('attachments'))

    if sender:
        direction = f'from {sender}'
    elif recipients:
        direction = f'to {recipients}'

    parts = [part for part in [date_value, direction, f'"{subject}"'] if part]
    summary = " ".join(parts).strip()
    if attachments:
        summary = f"{summary} (attachments: {attachments[:140]})".strip()
    return summary[:320]


def build_checklist_comment(match, row_candidates=None, emails=None):
    if not isinstance(match, dict):
        return ''

    llm_comment = normalize_cell_text(match.get('checklist_comment') or match.get('comment'))
    candidate_lookup = {}
    for candidate in row_candidates or []:
        thread_id = normalize_cell_text(candidate.get('thread_id'))
        if thread_id:
            candidate_lookup[thread_id] = candidate

    candidate_comments = []
    candidate_subjects = []
    for thread_id in match.get('matching_thread_ids', [])[:2]:
        candidate = candidate_lookup.get(normalize_cell_text(thread_id))
        if not isinstance(candidate, dict):
            continue
        comment = normalize_cell_text(candidate.get('suggested_comment'))
        if comment:
            candidate_comments.append(comment)
        subject = normalize_cell_text(candidate.get('subject'))
        if subject:
            candidate_subjects.append(subject)

    if candidate_comments:
        deterministic_comment = " ".join(candidate_comments[:2]).strip()
        if llm_comment and llm_comment.lower() not in deterministic_comment.lower():
            return f"{deterministic_comment} Notes: {llm_comment}"
        return deterministic_comment

    evidence_parts = []
    email_items = emails or []
    for idx in match.get('matching_email_indices', [])[:2]:
        if 0 <= idx < len(email_items):
            evidence = format_email_evidence(email_items[idx])
            if evidence:
                evidence_parts.append(evidence)

    if llm_comment and evidence_parts:
        return f"{llm_comment} Evidence: {'; '.join(evidence_parts)}"
    if llm_comment:
        return llm_comment
    if candidate_subjects:
        status = normalize_cell_text(match.get('status'))
        prefix = f"{status}: " if status else ''
        return f"{prefix}{'; '.join(candidate_subjects[:2])}".strip()
    if evidence_parts:
        status = normalize_cell_text(match.get('status'))
        prefix = f"{status}: " if status else ''
        return f"{prefix}{'; '.join(evidence_parts)}".strip()

    reasoning = normalize_cell_text(match.get('reasoning'))
    return reasoning[:400]


def merge_checklist_comment(existing_text, new_comment):
    existing = normalize_cell_text(existing_text)
    comment = normalize_cell_text(new_comment)
    if not comment:
        return existing
    generated = f"EmmaNeigh: {comment}"
    if not existing:
        return generated
    if existing == generated or generated in existing:
        return existing
    if existing.startswith('EmmaNeigh:'):
        return generated
    return f"{existing} | {generated}"


def infer_document_column(headers, rows):
    """
    Infer the document/title column when header names differ across templates/firms.
    """
    if not headers:
        return -1

    max_cols = max(len(headers), max((len(r) for r in rows), default=0))
    best_col = -1
    best_score = -1
    blocked_terms = ('status', 'state', 'progress', 'date', 'sent', 'received')

    for col_idx in range(max_cols):
        header_value = normalize_cell_text(headers[col_idx] if col_idx < len(headers) else '')
        header_lower = header_value.lower()
        if any(term in header_lower for term in blocked_terms):
            continue

        non_empty = 0
        rich_text = 0
        penalties = 0
        for row in rows[:150]:
            cell = normalize_cell_text(row[col_idx] if col_idx < len(row) else '')
            if not cell:
                continue
            non_empty += 1
            if re.search(r'[a-zA-Z]', cell):
                rich_text += 1
            if len(cell) > 220 or re.fullmatch(r'[\d\W_]+', cell):
                penalties += 1

        score = (non_empty * 2) + (rich_text * 3) - (penalties * 2)
        if score > best_score:
            best_score = score
            best_col = col_idx

    return best_col


def score_header_candidate(all_rows, header_idx):
    headers = all_rows[header_idx]
    non_empty_headers = [normalize_cell_text(h) for h in headers if normalize_cell_text(h)]
    if len(non_empty_headers) < 2:
        return None

    doc_col = find_column_index(headers, DOCUMENT_COLUMN_PATTERNS)
    status_col = find_column_index(headers, STATUS_COLUMN_PATTERNS)
    party_col = find_column_index(headers, PARTY_COLUMN_PATTERNS)
    notes_col = find_column_index(headers, NOTES_COLUMN_PATTERNS)

    avg_header_len = (
        sum(len(x) for x in non_empty_headers) / len(non_empty_headers)
        if non_empty_headers else 0
    )
    looks_like_header_text = 1 if avg_header_len <= 40 else -5

    # Score based on quality of the first rows after this candidate header.
    data_rows = all_rows[header_idx + 1: header_idx + 11]
    non_empty_data_rows = sum(1 for row in data_rows if any(normalize_cell_text(c) for c in row))

    score = len(non_empty_headers) + (non_empty_data_rows * 3) + looks_like_header_text
    if doc_col != -1:
        score += 40
    if status_col != -1:
        score += 20
    if party_col != -1:
        score += 6
    if notes_col != -1:
        score += 4

    return {
        'score': score,
        'header_idx': header_idx,
        'doc_col': doc_col,
        'status_col': status_col,
        'notes_col': notes_col,
    }


def iter_container_paragraphs(container):
    for paragraph in getattr(container, 'paragraphs', []):
        yield paragraph
    for table in getattr(container, 'tables', []):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in iter_container_paragraphs(cell):
                    yield paragraph


def format_draft_stamp_date(existing_text, stamp_date):
    now = stamp_date or datetime.now()
    text = normalize_cell_text(existing_text)
    if re.search(r'\b\d{1,2}-\d{1,2}-\d{2,4}\b', text):
        return now.strftime('%m-%d-%Y')
    if re.search(r'\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\s+\d{1,2},\s+\d{4}\b', text, re.IGNORECASE):
        return now.strftime('%B %d, %Y').replace(' 0', ' ')
    return now.strftime('%m/%d/%Y')


def replace_inline_date(text, replacement):
    updated = str(text or '')
    for pattern in INLINE_DATE_PATTERNS:
        if pattern.search(updated):
            return pattern.sub(replacement, updated, count=1)
    return updated


def refresh_document_draft_stamp(doc, stamp_date=None):
    if WD_ALIGN_PARAGRAPH is None:
        return False

    stamp_updated = False
    for section in getattr(doc, 'sections', []):
        header = getattr(section, 'header', None)
        if header is None:
            continue

        matched_paragraph = None
        for paragraph in iter_container_paragraphs(header):
            text = normalize_cell_text(paragraph.text)
            if 'draft' not in text.lower():
                continue
            replacement_date = format_draft_stamp_date(text, stamp_date)
            new_text = replace_inline_date(text, replacement_date)
            if new_text == text and replacement_date not in text:
                new_text = f"DRAFT {replacement_date}"
            paragraph.text = new_text
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            matched_paragraph = paragraph
            stamp_updated = True
            break

        if matched_paragraph is not None:
            continue

        if getattr(header, 'paragraphs', None):
            paragraph = header.paragraphs[0]
        else:
            paragraph = header.add_paragraph()
        paragraph.text = f"DRAFT {format_draft_stamp_date('', stamp_date)}"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        stamp_updated = True

    return stamp_updated


def parse_checklist_table(doc):
    """
    Parse a checklist-like table from the document, allowing flexible header position
    and non-standard column names.

    Returns:
        headers, rows, row_indices, table, doc_col_idx, status_col_idx, notes_col_idx, data_row_start_idx
    """
    if not doc.tables:
        return None, None, None, None, -1, -1, -1, -1

    best_candidate = None

    for table in doc.tables:
        all_rows = []
        for row in table.rows:
            all_rows.append([normalize_cell_text(cell.text) for cell in row.cells])

        if not all_rows:
            continue

        max_probe = min(6, len(all_rows) - 1)
        for header_idx in range(max_probe + 1):
            candidate = score_header_candidate(all_rows, header_idx)
            if not candidate:
                continue
            candidate['all_rows'] = all_rows
            candidate['table'] = table
            if best_candidate is None or candidate['score'] > best_candidate['score']:
                best_candidate = candidate

    if not best_candidate:
        return None, None, None, None, -1, -1, -1, -1

    table = best_candidate['table']
    all_rows = best_candidate['all_rows']
    header_idx = best_candidate['header_idx']
    headers = all_rows[header_idx]
    row_entries = [
        (absolute_idx, row)
        for absolute_idx, row in enumerate(all_rows[header_idx + 1:], start=header_idx + 1)
        if any(normalize_cell_text(cell) for cell in row)
    ]
    rows = [row for _, row in row_entries]
    row_indices = [absolute_idx for absolute_idx, _ in row_entries]

    doc_col_idx = best_candidate['doc_col']
    if doc_col_idx == -1:
        doc_col_idx = infer_document_column(headers, rows)

    status_col_idx = best_candidate['status_col']
    if status_col_idx == -1:
        status_col_idx = len(headers)  # Signals "no existing status column".

    notes_col_idx = best_candidate.get('notes_col', -1)

    return headers, rows, row_indices, table, doc_col_idx, status_col_idx, notes_col_idx, header_idx + 1


def detect_document_status(doc_name, emails):
    """
    Search emails for mentions of a document and detect its status.

    Returns:
        - status: detected status string or None
        - priority: status priority (higher = more advanced)
        - matching_emails: list of email subjects that matched
    """
    if not doc_name:
        return None, 0, []

    doc_name_lower = doc_name.lower()

    # Create search patterns from document name
    # Handle common variations
    doc_words = re.findall(r'\b\w+\b', doc_name_lower)

    best_status = None
    best_priority = 0
    matching_emails = []

    for email in emails:
        searchable = email.get('searchable', '')

        # Check if email mentions this document
        # Use fuzzy matching - at least 2 key words must match
        word_matches = sum(1 for word in doc_words if len(word) > 3 and word in searchable)

        if word_matches < 2 and doc_name_lower not in searchable:
            continue

        # Email mentions this document - check for status patterns
        for status_config in STATUS_PATTERNS:
            for pattern in status_config['patterns']:
                if re.search(pattern, searchable):
                    if status_config['priority'] > best_priority:
                        best_status = status_config['status']
                        best_priority = status_config['priority']
                        matching_emails.append(email.get('subject', 'No subject'))
                    break

    return best_status, best_priority, matching_emails


def _parse_row_id(value):
    try:
        return int(str(value).strip())
    except Exception:
        return None


def normalize_checklist_llm_matches(raw_payload):
    """
    Normalize LLM outputs into a stable structure:
      {
        "by_row": { row_id: {...} },
        "by_doc": { canonical_doc_key: {...} }
      }
    Supports both legacy doc-name keyed responses and new row-id keyed payloads.
    """
    normalized = {"by_row": {}, "by_doc": {}}
    payload = raw_payload if isinstance(raw_payload, dict) else {}

    def coerce_match(entry, default_doc_name=None):
        if not isinstance(entry, dict):
            return None
        status = normalize_llm_status(entry.get("status"))
        raw_indices = entry.get("matching_email_indices", [])
        email_indices = []
        if isinstance(raw_indices, list):
            for idx in raw_indices[:15]:
                try:
                    parsed_idx = int(idx)
                except Exception:
                    continue
                if parsed_idx >= 0:
                    email_indices.append(parsed_idx)
        raw_thread_ids = entry.get("matching_thread_ids", [])
        thread_ids = []
        if isinstance(raw_thread_ids, list):
            for raw_thread_id in raw_thread_ids[:8]:
                parsed_thread_id = str(raw_thread_id or "").strip()
                if parsed_thread_id and parsed_thread_id not in thread_ids:
                    thread_ids.append(parsed_thread_id)

        confidence = entry.get("confidence", 0.5)
        try:
            confidence = float(confidence)
        except Exception:
            confidence = 0.5
        confidence = max(0.0, min(confidence, 1.0))

        doc_name = str(entry.get("document_name") or default_doc_name or "").strip()
        row_id = _parse_row_id(entry.get("row_id"))
        has_activity = normalize_bool(entry.get("has_activity"), default=bool(status or email_indices))
        checklist_comment = str(
            entry.get("checklist_comment") or entry.get("comment") or ""
        ).strip()

        if not doc_name and row_id is None:
            return None

        return {
            "row_id": row_id,
            "document_name": doc_name,
            "status": status,
            "has_activity": has_activity,
            "matching_email_indices": email_indices,
            "matching_thread_ids": thread_ids,
            "confidence": confidence,
            "reasoning": str(entry.get("reasoning") or "").strip(),
            "checklist_comment": checklist_comment,
        }

    # New format: { "matches": [ ... ] }
    if isinstance(payload.get("matches"), list):
        for item in payload.get("matches", []):
            parsed = coerce_match(item)
            if not parsed:
                continue
            row_id = parsed.get("row_id")
            doc_name = parsed.get("document_name", "")
            if row_id is not None:
                normalized["by_row"][row_id] = parsed
            if doc_name:
                normalized["by_doc"][canonical_doc_key(doc_name)] = parsed

    # Legacy format: { "Document Name": { ... } }
    for key, value in payload.items():
        if key == "matches":
            continue
        parsed = coerce_match(value, default_doc_name=key)
        if not parsed:
            continue
        row_id = parsed.get("row_id")
        doc_name = parsed.get("document_name", "") or str(key)
        if row_id is not None:
            normalized["by_row"][row_id] = parsed
        normalized["by_doc"][canonical_doc_key(doc_name)] = parsed

    return normalized


def match_documents_with_llm(
    checklist_items,
    row_thread_candidates,
    api_key,
    provider="anthropic",
    model_name=None,
    provider_base_url=None,
):
    """
    Use the selected LLM provider to match checklist rows to email activity and infer status.
    """
    if provider_requires_api_key(provider) and not api_key:
        raise ValueError("API key is required for checklist LLM analysis.")
    normalized = {"by_row": {}, "by_doc": {}, "warnings": []}

    for batch in chunk_items(checklist_items[:220], CHECKLIST_LLM_BATCH_SIZE):
        checklist_context = []
        for item in batch:
            row_id = item.get("row_id")
            candidates = []
            for candidate in (row_thread_candidates or {}).get(row_id, [])[:MAX_CANDIDATE_THREADS_PER_ROW]:
                candidates.append({
                    "thread_id": candidate.get("thread_id"),
                    "score": candidate.get("score"),
                    "subject": candidate.get("subject", ""),
                    "message_count": candidate.get("message_count", 0),
                    "folders": candidate.get("folders", []),
                    "participants": candidate.get("participants", []),
                    "date_range": {
                        "first": candidate.get("first_date", ""),
                        "last": candidate.get("last_date", ""),
                    },
                    "deal_anchor_hits": candidate.get("deal_anchor_hits", []),
                    "status_signal": candidate.get("status_signal", ""),
                    "matching_subjects": candidate.get("matching_subjects", []),
                    "matching_attachment_titles": candidate.get("matching_attachment_titles", []),
                    "latest_attachment_titles": candidate.get("latest_attachment_titles", []),
                    "latest_body_excerpt": candidate.get("latest_body_excerpt", ""),
                    "relevant_events": candidate.get("events", []),
                    "issues": candidate.get("issues", []),
                    "suggested_comment": candidate.get("suggested_comment", ""),
                })
            checklist_context.append({
                "row_id": row_id,
                "document_name": item.get("document_name", ""),
                "match_document_name": item.get("match_document_name", item.get("document_name", "")),
                "row_type": item.get("row_type", "document"),
                "current_status": item.get("current_status", ""),
                "row_context": item.get("row_context", ""),
                "candidate_threads": candidates,
            })

        if not any(row.get("candidate_threads") for row in checklist_context):
            for item in batch:
                row_id = item.get("row_id")
                doc_name = item.get("document_name", "")
                normalized["by_row"][row_id] = {
                    "row_id": row_id,
                    "document_name": doc_name,
                    "status": "",
                    "has_activity": False,
                    "matching_email_indices": [],
                    "matching_thread_ids": [],
                    "confidence": 0.0,
                    "reasoning": "No candidate threads were found for this checklist row.",
                    "checklist_comment": "",
                }
                normalized["by_doc"][canonical_doc_key(doc_name)] = normalized["by_row"][row_id]
            continue

        prompt = f"""You are a legal transaction assistant updating a checklist from email evidence.

For EACH checklist row, decide whether any candidate thread directly concerns the exact checklist document for the correct deal.

CHECKLIST ROWS:
{json.dumps(checklist_context, indent=2)}

Status options (use exactly one when there is activity):
- "Pending Draft"
- "Draft Circulated"
- "With Opposing Counsel"
- "Agreed Form"
- "Execution Version"
- "Executed"

Interpretation guidance:
- "Draft Circulated": a draft or markup was sent internally or generally circulated.
- "With Opposing Counsel": sent to counterparty/opposing counsel/client for review or comments.
- "Agreed Form": settled/final form with no material open comments.
- "Execution Version": ready for signature, signature pages circulated, or sent for execution.
- "Executed": fully signed, executed copies received, or executed versions circulated.

Instructions:
- Use ONLY the listed candidate_threads for each row.
- A thread is relevant only if it is clearly about the same deal and the same checklist document, not just the same workstream.
- Start with the thread subject, then matching_attachment_titles/latest_attachment_titles, then the latest_body_excerpt, then deal_anchor_hits and the concrete draft/execution events.
- Prefer threads with direct attachment filename matches over body-only references.
- Ignore unrelated admin traffic, scheduling, or threads about different deal documents.
- When there is activity, the checklist_comment should say who sent or received the draft/version, on what date, and note material issues or comments raised in the thread.
- Return EVERY row_id exactly once.

Return JSON only in this exact shape:
{{
  "matches": [
    {{
      "row_id": 12,
      "document_name": "Credit Agreement",
      "has_activity": true,
      "status": "With Opposing Counsel",
      "matching_thread_ids": ["thread-2"],
      "confidence": 0.82,
      "reasoning": "Thread 2 is for the same agreement and deal, and it shows a revised draft sent to buyer counsel for review.",
      "checklist_comment": "Received revised draft from seller counsel on 03/11/2026 and sent markup to buyer counsel on 03/13/2026. Issues flagged: open point on working capital adjustment."
    }},
    {{
      "row_id": 13,
      "document_name": "Security Agreement",
      "has_activity": false,
      "status": "",
      "matching_thread_ids": [],
      "confidence": 0.12,
      "reasoning": "No candidate thread directly concerns this checklist row.",
      "checklist_comment": ""
    }}
  ]
}}
"""

        try:
            parsed = call_provider_prompt_json(
                prompt,
                api_key,
                provider,
                model_name=model_name,
                provider_base_url=provider_base_url,
            )
        except Exception as exc:
            normalized["warnings"].append(str(exc))
            continue

        batch_matches = normalize_checklist_llm_matches(parsed)
        normalized["by_row"].update(batch_matches.get("by_row", {}))
        normalized["by_doc"].update(batch_matches.get("by_doc", {}))

        for item in batch:
            row_id = item.get("row_id")
            doc_name = item.get("document_name", "")
            if row_id not in normalized["by_row"]:
                normalized["by_row"][row_id] = {
                    "row_id": row_id,
                    "document_name": doc_name,
                    "status": "",
                    "has_activity": False,
                    "matching_email_indices": [],
                    "matching_thread_ids": [],
                    "confidence": 0.0,
                    "reasoning": "Model did not return a decision for this row.",
                    "checklist_comment": "",
                }
                normalized["by_doc"][canonical_doc_key(doc_name)] = normalized["by_row"][row_id]

    return normalized


def update_checklist(
    checklist_path,
    email_input_path,
    output_folder,
    api_key=None,
    provider="anthropic",
    model_name=None,
    provider_base_url=None,
):
    """
    Main function to update checklist based on email activity.

    Args:
        checklist_path: Path to Word document with checklist table
        email_input_path: Path to CSV or JSON email activity input
        output_folder: Folder to save updated checklist
        api_key: Provider API key
        provider: LLM provider (anthropic/openai/harvey)
        model_name: Optional provider model override
        provider_base_url: Optional provider API base URL

    Returns:
        dict with: success, output_path, items_updated, details
    """
    result = {
        'success': False,
        'output_path': None,
        'items_updated': 0,
        'details': [],
        'error': None
    }

    try:
        api_key = normalize_api_key(api_key)
        provider = resolve_provider(provider, api_key)

        if provider_requires_api_key(provider) and not api_key:
            result['error'] = 'LLM API key is required to update checklist from email activity.'
            return result

        # Parse email activity
        emails = parse_email_input(email_input_path)
        if not emails:
            result['error'] = 'No emails found in the supplied email activity.'
            return result

        if Document is None:
            result['error'] = 'python-docx is required to update checklist documents in this environment.'
            return result

        # Open checklist document
        doc = Document(checklist_path)

        # Parse the checklist table
        headers, rows, row_indices, table, doc_col_idx, status_col_idx, notes_col_idx, data_row_start_idx = parse_checklist_table(doc)

        if table is None:
            result['error'] = 'No table found in checklist document'
            return result

        if doc_col_idx == -1:
            result['error'] = 'Could not identify document name column in checklist'
            return result

        # Check if we need to add a status column
        needs_new_status_col = status_col_idx >= len(headers)

        if needs_new_status_col:
            # Add "Status" header to first row
            header_row_idx = data_row_start_idx - 1 if data_row_start_idx > 0 else 0
            header_row = table.rows[header_row_idx]
            # Add new cell (this is complex in python-docx, so we'll update existing empty column if available)
            # For now, we'll look for an empty column or skip
            for i, header in enumerate(headers):
                if not header.strip():
                    status_col_idx = i
                    header_row.cells[i].text = "Status (Updated)"
                    needs_new_status_col = False
                    break

            if needs_new_status_col:
                result['error'] = 'No status column found and cannot add new column. Please add a Status column to your checklist.'
                return result

        # Build checklist row payloads for semantic LLM matching.
        checklist_items = []
        previous_document_name = ''
        for row_idx, row_data in enumerate(rows):
            if doc_col_idx < 0 or doc_col_idx >= len(row_data):
                continue
            doc_name = normalize_cell_text(row_data[doc_col_idx])
            if not doc_name:
                continue
            row_index = row_indices[row_idx] if row_indices and row_idx < len(row_indices) else data_row_start_idx + row_idx

            row_classification = classify_checklist_row(
                doc_name,
                row_data,
                headers,
                doc_col_idx,
                status_col_idx,
                notes_col_idx,
                previous_document_name=previous_document_name,
            )
            if row_classification.get('row_type') == 'document':
                previous_document_name = doc_name
            if not row_classification.get('should_update'):
                continue

            current_status = ''
            if status_col_idx < len(row_data):
                current_status = normalize_cell_text(row_data[status_col_idx])

            existing_notes = ''
            if notes_col_idx >= 0 and notes_col_idx < len(row_data):
                existing_notes = normalize_cell_text(row_data[notes_col_idx])

            context_parts = []
            for col_idx, cell_value in enumerate(row_data):
                if col_idx in (doc_col_idx, status_col_idx):
                    continue
                cleaned = normalize_cell_text(cell_value)
                if not cleaned:
                    continue
                header_label = normalize_cell_text(headers[col_idx] if col_idx < len(headers) else f"Column {col_idx + 1}")
                if header_label:
                    context_parts.append(f"{header_label}: {cleaned}")
                else:
                    context_parts.append(cleaned)

            checklist_items.append({
                'row_id': row_index,
                'document_name': doc_name,
                'match_document_name': row_classification.get('match_document_name') or doc_name,
                'row_type': row_classification.get('row_type', 'document'),
                'current_status': current_status,
                'existing_notes': existing_notes,
                'row_context': ' | '.join(context_parts[:8]),
                'row_data': row_data,
            })

        if not checklist_items:
            result['error'] = 'No checklist document names found to analyze.'
            return result

        row_thread_candidates, _, deal_profile = build_row_thread_candidates(
            checklist_path,
            checklist_items,
            emails,
        )

        # LLM matching is mandatory for checklist updates in this workflow.
        llm_warning = None
        try:
            llm_matches = match_documents_with_llm(
                checklist_items,
                row_thread_candidates,
                api_key,
                provider=provider,
                model_name=model_name,
                provider_base_url=provider_base_url,
            )
        except Exception as e:
            llm_matches = {"by_row": {}, "by_doc": {}, "warnings": [str(e)]}
            llm_warning = str(e)

        # Process each row
        items_updated = 0
        details = []

        for item in checklist_items:
            doc_name = item['document_name']
            row_data = item.get('row_data', [])
            row_id = item.get('row_id')

            current_status = item.get('current_status', '')
            existing_notes = item.get('existing_notes', '')
            new_status = None
            matching_emails = []
            generated_comment = ''
            row_candidates = row_thread_candidates.get(row_id, [])
            row_candidate_lookup = {
                normalize_cell_text(candidate.get('thread_id')): candidate
                for candidate in row_candidates
                if normalize_cell_text(candidate.get('thread_id'))
            }

            llm_result = None
            if isinstance(llm_matches, dict):
                by_row = llm_matches.get('by_row', {})
                by_doc = llm_matches.get('by_doc', {})
                llm_result = by_row.get(row_id) if isinstance(by_row, dict) else None
                if not llm_result and isinstance(by_doc, dict):
                    llm_result = by_doc.get(canonical_doc_key(doc_name))

            if llm_result and llm_result.get('has_activity'):
                new_status = normalize_llm_status(llm_result.get('status'))
                matching_thread_ids = llm_result.get('matching_thread_ids', [])
                if not new_status and matching_thread_ids:
                    primary_candidate = row_candidate_lookup.get(normalize_cell_text(matching_thread_ids[0]))
                    if isinstance(primary_candidate, dict):
                        new_status = normalize_llm_status(primary_candidate.get('status_signal'))

                generated_comment = build_checklist_comment(
                    llm_result,
                    row_candidates=row_candidates,
                    emails=emails,
                )

                for thread_id in matching_thread_ids[:2]:
                    candidate = row_candidate_lookup.get(normalize_cell_text(thread_id))
                    if not isinstance(candidate, dict):
                        continue
                    matching_emails.extend(candidate.get('matching_subjects', [])[:2])

                if not matching_emails:
                    email_indices = llm_result.get('matching_email_indices', [])
                    for idx in email_indices[:3]:
                        if 0 <= idx < len(emails):
                            matching_emails.append(emails[idx].get('subject', 'No subject'))
            else:
                strongest_candidate = row_candidates[0] if row_candidates else None
                if strongest_candidate and float(strongest_candidate.get('score', 0.0) or 0.0) >= HIGH_CONFIDENCE_THREAD_SCORE:
                    new_status = normalize_llm_status(strongest_candidate.get('status_signal')) or 'Draft Circulated'
                    generated_comment = normalize_cell_text(strongest_candidate.get('suggested_comment'))
                    matching_emails = strongest_candidate.get('matching_subjects', [])[:3]
                else:
                    fallback_status, _, fallback_subjects = detect_document_status(doc_name, emails)
                    if fallback_status:
                        new_status = fallback_status
                        matching_emails = fallback_subjects[:3]
                        generated_comment = f"Email activity indicates {fallback_status}. Example emails: {'; '.join(matching_emails[:2])}".strip()

            merged_comment = merge_checklist_comment(existing_notes, generated_comment)
            status_changed = bool(new_status and new_status != current_status)
            notes_changed = bool(merged_comment and merged_comment != existing_notes and notes_col_idx >= 0)

            if not status_changed and not notes_changed:
                continue

            if row_id is None or row_id < 0 or row_id >= len(table.rows):
                continue

            table_row = table.rows[row_id]
            row_updated = False

            if status_changed and status_col_idx < len(table_row.cells):
                status_cell = table_row.cells[status_col_idx]
                if status_cell.paragraphs:
                    status_cell.paragraphs[0].text = new_status
                else:
                    status_cell.text = new_status
                row_updated = True

            if notes_changed and notes_col_idx < len(table_row.cells):
                notes_cell = table_row.cells[notes_col_idx]
                if notes_cell.paragraphs:
                    notes_cell.paragraphs[0].text = merged_comment
                else:
                    notes_cell.text = merged_comment
                row_updated = True

            if row_updated:
                items_updated += 1
                details.append({
                    'document': doc_name,
                    'old_status': current_status,
                    'new_status': new_status or current_status,
                    'comment': merged_comment,
                    'confidence': llm_result.get('confidence') if isinstance(llm_result, dict) else None,
                    'emails': matching_emails[:3]
                })

        refresh_document_draft_stamp(doc)

        # Save updated document
        os.makedirs(output_folder, exist_ok=True)

        base_name = os.path.splitext(os.path.basename(checklist_path))[0]
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"{base_name}_updated_{timestamp}.docx"
        output_path = os.path.join(output_folder, output_filename)

        doc.save(output_path)

        result['success'] = True
        result['output_path'] = output_path
        result['items_updated'] = items_updated
        result['details'] = details
        result['emails_processed'] = len(emails)
        result['llm_documents_with_activity'] = sum(
            1 for item in (llm_matches or {}).get('by_row', {}).values()
            if isinstance(item, dict) and item.get('has_activity')
        )
        result['deal_anchor_tokens'] = (deal_profile or {}).get('anchor_tokens', [])
        result['warning'] = llm_warning or '; '.join((llm_matches or {}).get('warnings', []))

        return result

    except Exception as e:
        result['error'] = str(e)
        return result


def main():
    if len(sys.argv) < 4:
        print(json.dumps({
            'success': False,
            'error': 'Usage: checklist_updater.py <checklist_path> <email_input_path> <output_folder> <api_key> [provider] [model] [provider_base_url]'
        }))
        sys.exit(1)

    checklist_path = sys.argv[1]
    email_input_path = sys.argv[2]
    output_folder = sys.argv[3]

    # API key is required for LLM-driven checklist updates.
    api_key = sys.argv[4] if len(sys.argv) > 4 else (
        os.environ.get('ANTHROPIC_API_KEY')
        or os.environ.get('OPENAI_API_KEY')
        or os.environ.get('HARVEY_API_KEY')
    )
    provider = resolve_provider(
        sys.argv[5] if len(sys.argv) > 5 else os.environ.get('AI_PROVIDER', 'anthropic'),
        api_key,
    )
    model_name = sys.argv[6] if len(sys.argv) > 6 else (
        os.environ.get('OPENAI_MODEL') if provider == 'openai'
        else os.environ.get('LOCALAI_MODEL') if provider == 'localai'
        else os.environ.get('OLLAMA_MODEL') if provider == 'ollama'
        else os.environ.get('LMSTUDIO_MODEL') if provider == 'lmstudio'
        else os.environ.get('CLAUDE_MODEL')
    )
    provider_base_url = sys.argv[7] if len(sys.argv) > 7 else (
        os.environ.get('HARVEY_BASE_URL', DEFAULT_HARVEY_BASE_URL) if provider == 'harvey'
        else os.environ.get('LOCALAI_BASE_URL', DEFAULT_LOCALAI_BASE_URL) if provider == 'localai'
        else os.environ.get('OLLAMA_BASE_URL', DEFAULT_OLLAMA_BASE_URL) if provider == 'ollama'
        else os.environ.get('LMSTUDIO_BASE_URL', DEFAULT_LMSTUDIO_BASE_URL) if provider == 'lmstudio'
        else None
    )

    result = update_checklist(
        checklist_path,
        email_input_path,
        output_folder,
        api_key,
        provider=provider,
        model_name=model_name,
        provider_base_url=provider_base_url,
    )
    print(json.dumps(result))


if __name__ == '__main__':
    main()
