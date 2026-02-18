#!/usr/bin/env python3
"""
EmmaNeigh - Document Collator
v5.1.3: Format-preserving merge of track changes from multiple Word documents.

Takes a precedent (base) document and multiple modified versions,
then merges all edits into one combined document showing all changes
while PRESERVING the original document's formatting (fonts, styles, margins).

Key change from v5.1.2: Instead of creating a new document, we copy the base
document and modify it in-place to preserve all formatting.
"""

import os
import sys
import json
import shutil
import zipfile
import difflib
import copy
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def emit(msg_type, **kwargs):
    """Output JSON message to stdout for the Electron app."""
    print(json.dumps({"type": msg_type, **kwargs}), flush=True)


def get_author_from_docx(docx_path):
    """Try to extract author from document properties or use filename."""
    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            if 'docProps/core.xml' in zf.namelist():
                import re
                core_xml = zf.read('docProps/core.xml').decode('utf-8')
                match = re.search(r'<dc:creator[^>]*>([^<]+)</dc:creator>', core_xml)
                if match:
                    return match.group(1).strip()
    except:
        pass
    # Fall back to filename without extension
    return os.path.splitext(os.path.basename(docx_path))[0]


def copy_run_formatting(source_run, target_run):
    """Copy formatting from source run to target run."""
    try:
        # Copy font properties
        if source_run.font.name:
            target_run.font.name = source_run.font.name
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        target_run.font.bold = source_run.font.bold
        target_run.font.italic = source_run.font.italic
        target_run.font.underline = source_run.font.underline
        # Don't copy color - we need to set our own for track changes
    except:
        pass


def get_paragraph_text_with_runs(paragraph):
    """Get paragraph text and run boundaries for precise editing."""
    runs_info = []
    for run in paragraph.runs:
        runs_info.append({
            'text': run.text,
            'run': run
        })
    return runs_info


def clear_paragraph_content(paragraph):
    """Remove all runs from a paragraph while keeping the paragraph itself."""
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def add_tracked_change_run(paragraph, text, change_type, reference_run=None):
    """
    Add a run with track change formatting.
    change_type: 'deleted', 'inserted', 'unchanged'
    """
    run = paragraph.add_run(text)

    # Copy formatting from reference run if available
    if reference_run:
        copy_run_formatting(reference_run, run)

    if change_type == 'deleted':
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.font.strike = True
    elif change_type == 'inserted':
        run.font.color.rgb = RGBColor(0, 0, 255)
    # 'unchanged' keeps original formatting

    return run


def create_merged_document_preserving_format(base_path, modified_paths, output_path):
    """
    Create a merged document showing all changes while PRESERVING base document formatting.

    Strategy:
    1. Copy the base document file
    2. Open the copy and modify paragraphs in-place
    3. This preserves: styles, fonts, margins, headers, footers, page layout
    """
    import re

    # Step 1: Copy the base document as our starting point
    shutil.copy2(base_path, output_path)

    # Step 2: Open the copy for editing
    output_doc = Document(output_path)

    # Get base paragraphs text for comparison
    base_paragraphs = [p.text for p in output_doc.paragraphs]

    # Step 3: Collect all modifications from all sources
    all_modifications = []  # List of modification records

    for mod_path in modified_paths:
        author = get_author_from_docx(mod_path)
        mod_doc = Document(mod_path)
        mod_paragraphs = [p.text for p in mod_doc.paragraphs]

        # Use difflib to align paragraphs
        matcher = difflib.SequenceMatcher(None, base_paragraphs, mod_paragraphs)

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'replace':
                # Paragraphs were modified
                for bi, mi in zip(range(i1, i2), range(j1, j2)):
                    if bi < len(base_paragraphs) and mi < len(mod_paragraphs):
                        if base_paragraphs[bi] != mod_paragraphs[mi]:
                            all_modifications.append({
                                'author': author,
                                'base_text': base_paragraphs[bi],
                                'modified_text': mod_paragraphs[mi],
                                'para_index': bi,
                                'type': 'modified'
                            })
            elif tag == 'delete':
                # Paragraphs deleted
                for bi in range(i1, i2):
                    all_modifications.append({
                        'author': author,
                        'base_text': base_paragraphs[bi],
                        'modified_text': '',
                        'para_index': bi,
                        'type': 'deleted'
                    })
            elif tag == 'insert':
                # New paragraphs added
                insert_after = i1 - 1 if i1 > 0 else 0
                for mi in range(j1, j2):
                    all_modifications.append({
                        'author': author,
                        'base_text': '',
                        'modified_text': mod_paragraphs[mi],
                        'para_index': insert_after,
                        'type': 'inserted'
                    })

    # Group modifications by paragraph index
    mods_by_para = {}
    for mod in all_modifications:
        idx = mod['para_index']
        if idx not in mods_by_para:
            mods_by_para[idx] = []
        mods_by_para[idx].append(mod)

    changes_count = 0

    # Collect insertions to add after processing
    insertions_to_add = {}  # para_index -> list of insertions
    for idx, mods in mods_by_para.items():
        for mod in mods:
            if mod['type'] == 'inserted':
                if idx not in insertions_to_add:
                    insertions_to_add[idx] = []
                insertions_to_add[idx].append(mod)

    # Step 4: Modify paragraphs in-place (preserving formatting)
    for i, para in enumerate(output_doc.paragraphs):
        mods = mods_by_para.get(i, [])

        # Filter to non-insertions for this paragraph
        para_mods = [m for m in mods if m['type'] != 'inserted']

        if not para_mods:
            # No changes - keep paragraph as-is (formatting preserved)
            continue

        # Get a reference run for formatting (use first run if available)
        reference_run = para.runs[0] if para.runs else None
        original_text = para.text

        deletions = [m for m in para_mods if m['type'] == 'deleted']
        modifications = [m for m in para_mods if m['type'] == 'modified']

        if deletions and not modifications:
            # Paragraph was deleted - show as strikethrough red
            if original_text.strip():
                # Clear and rebuild with strikethrough
                clear_paragraph_content(para)
                run = add_tracked_change_run(para, original_text, 'deleted', reference_run)
                # Add author attribution
                attr = para.add_run(f" [{deletions[0]['author']}]")
                attr.font.size = Pt(8)
                attr.font.color.rgb = RGBColor(128, 128, 128)
                changes_count += 1

        elif modifications:
            # Apply word-level diff while preserving paragraph formatting
            mod = modifications[0]
            clear_paragraph_content(para)
            apply_word_diff_in_place(para, original_text, mod['modified_text'],
                                     mod['author'], reference_run)
            changes_count += 1

    # Step 5: Handle insertions (add new paragraphs after existing ones)
    # We need to insert in reverse order to maintain correct indices
    sorted_indices = sorted(insertions_to_add.keys(), reverse=True)

    for idx in sorted_indices:
        insertions = insertions_to_add[idx]

        # Get the paragraph to insert after
        if idx < len(output_doc.paragraphs):
            target_para = output_doc.paragraphs[idx]

            for ins in reversed(insertions):  # Reverse to maintain order
                if ins['modified_text'].strip():
                    # Create new paragraph after target
                    new_para = insert_paragraph_after(target_para)

                    run = new_para.add_run(ins['modified_text'])
                    run.font.color.rgb = RGBColor(0, 0, 255)

                    # Add author attribution
                    attr = new_para.add_run(f" [Added by {ins['author']}]")
                    attr.font.size = Pt(8)
                    attr.font.color.rgb = RGBColor(128, 128, 128)
                    changes_count += 1

    # Step 6: Add legend at the very top
    if changes_count > 0:
        # Insert legend as first paragraph
        first_para = output_doc.paragraphs[0] if output_doc.paragraphs else output_doc.add_paragraph()
        legend_para = insert_paragraph_before(first_para)

        bold_run = legend_para.add_run("COLLATED DOCUMENT - ")
        bold_run.bold = True
        legend_para.add_run("Legend: ")
        del_run = legend_para.add_run("Deleted")
        del_run.font.color.rgb = RGBColor(255, 0, 0)
        del_run.font.strike = True
        legend_para.add_run(" | ")
        ins_run = legend_para.add_run("Added")
        ins_run.font.color.rgb = RGBColor(0, 0, 255)
        legend_para.add_run(" | ")
        legend_para.add_run("[Author attribution in gray]")

        # Add blank line after legend
        blank_para = insert_paragraph_after(legend_para)

    # Save the modified document
    output_doc.save(output_path)
    return changes_count


def insert_paragraph_after(paragraph):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    new_para = paragraph._element.getnext()

    # Create proper Paragraph object
    from docx.text.paragraph import Paragraph
    return Paragraph(new_para, paragraph._parent)


def insert_paragraph_before(paragraph):
    """Insert a new paragraph before the given paragraph."""
    new_p = OxmlElement('w:p')
    paragraph._element.addprevious(new_p)
    new_para = paragraph._element.getprevious()

    # Create proper Paragraph object
    from docx.text.paragraph import Paragraph
    return Paragraph(new_para, paragraph._parent)


def apply_word_diff_in_place(para, base_text, modified_text, author, reference_run=None):
    """Apply word-level diff to a paragraph while preserving formatting."""
    import re

    # Split into words preserving whitespace
    base_words = re.findall(r'\S+|\s+', base_text)
    mod_words = re.findall(r'\S+|\s+', modified_text)

    matcher = difflib.SequenceMatcher(None, base_words, mod_words)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            text = ''.join(base_words[i1:i2])
            if text:
                run = para.add_run(text)
                if reference_run:
                    copy_run_formatting(reference_run, run)

        elif tag == 'delete':
            text = ''.join(base_words[i1:i2])
            if text.strip():
                run = para.add_run(text)
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.strike = True
                if reference_run:
                    # Copy font but override color
                    if reference_run.font.name:
                        run.font.name = reference_run.font.name
                    if reference_run.font.size:
                        run.font.size = reference_run.font.size

        elif tag == 'insert':
            text = ''.join(mod_words[j1:j2])
            if text.strip():
                run = para.add_run(text)
                run.font.color.rgb = RGBColor(0, 0, 255)
                if reference_run:
                    if reference_run.font.name:
                        run.font.name = reference_run.font.name
                    if reference_run.font.size:
                        run.font.size = reference_run.font.size

        elif tag == 'replace':
            # Show deletion then insertion
            del_text = ''.join(base_words[i1:i2])
            ins_text = ''.join(mod_words[j1:j2])

            if del_text.strip():
                run = para.add_run(del_text)
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.strike = True
                if reference_run:
                    if reference_run.font.name:
                        run.font.name = reference_run.font.name
                    if reference_run.font.size:
                        run.font.size = reference_run.font.size

            if ins_text.strip():
                run = para.add_run(ins_text)
                run.font.color.rgb = RGBColor(0, 0, 255)
                if reference_run:
                    if reference_run.font.name:
                        run.font.name = reference_run.font.name
                    if reference_run.font.size:
                        run.font.size = reference_run.font.size

    # Add author attribution at end
    attr = para.add_run(f" [{author}]")
    attr.font.size = Pt(8)
    attr.font.color.rgb = RGBColor(128, 128, 128)


def create_summary(base_path, modified_paths, changes_count, output_folder):
    """Create a simple summary document."""
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Document Collation Summary")
    title_run.bold = True
    title_run.font.size = Pt(16)

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")

    doc.add_paragraph(f"Base Document: {os.path.basename(base_path)}")
    doc.add_paragraph(f"Modified Versions: {len(modified_paths)}")
    doc.add_paragraph("")

    # List modified docs with authors
    sources_heading = doc.add_paragraph()
    sources_heading.add_run("Sources:").bold = True

    for path in modified_paths:
        author = get_author_from_docx(path)
        doc.add_paragraph(f"  • {os.path.basename(path)} (Author: {author})")

    doc.add_paragraph("")
    doc.add_paragraph(f"Total Changes Applied: {changes_count}")

    summary_path = os.path.join(output_folder, "Collation_Summary.docx")
    doc.save(summary_path)
    return summary_path


def collate_documents(base_path, modified_paths, output_folder):
    """
    Main collation function.

    Args:
        base_path: Path to the precedent/base document
        modified_paths: List of paths to modified versions
        output_folder: Where to save output

    Returns:
        dict with results
    """
    os.makedirs(output_folder, exist_ok=True)

    emit("progress", percent=10, message="Reading base document...")

    # Validate inputs
    if not os.path.isfile(base_path):
        raise ValueError(f"Base document not found: {base_path}")

    valid_modified = []
    for p in modified_paths:
        if os.path.isfile(p):
            valid_modified.append(p)
        else:
            emit("progress", percent=0, message=f"Warning: Skipping missing file {p}")

    if not valid_modified:
        raise ValueError("No valid modified documents found")

    emit("progress", percent=30, message=f"Processing {len(valid_modified)} modified documents...")

    # Create merged document (preserving formatting)
    base_name = os.path.splitext(os.path.basename(base_path))[0]
    output_path = os.path.join(output_folder, f"{base_name}_Collated.docx")

    emit("progress", percent=50, message="Merging changes (preserving formatting)...")

    changes_count = create_merged_document_preserving_format(base_path, valid_modified, output_path)

    emit("progress", percent=80, message="Creating summary...")

    # Create summary
    summary_path = create_summary(base_path, valid_modified, changes_count, output_folder)

    emit("progress", percent=100, message="Complete!")

    return {
        'output_folder': output_folder,
        'output_document': output_path,
        'summary_document': summary_path,
        'total_changes': changes_count,
        'documents_processed': len(valid_modified)
    }


def main():
    """CLI entry point."""
    if len(sys.argv) < 2:
        emit("error", message="Usage: document_collator.py <config_json_path>")
        sys.exit(1)

    config_path = sys.argv[1]

    if not os.path.isfile(config_path):
        emit("error", message=f"Config file not found: {config_path}")
        sys.exit(1)

    try:
        with open(config_path, 'r') as f:
            config = json.load(f)
    except json.JSONDecodeError as e:
        emit("error", message=f"Invalid JSON config: {str(e)}")
        sys.exit(1)

    base_path = config.get('base_document')
    modified_paths = config.get('commented_documents', [])
    output_folder = config.get('output_folder')

    if not base_path:
        emit("error", message="No base document specified")
        sys.exit(1)

    if not modified_paths:
        emit("error", message="No modified documents specified")
        sys.exit(1)

    if not output_folder:
        output_folder = os.path.join(os.path.dirname(base_path), "collated_output")

    try:
        results = collate_documents(base_path, modified_paths, output_folder)

        emit("result",
             success=True,
             output_folder=results['output_folder'],
             output_document=results['output_document'],
             summary_document=results['summary_document'],
             total_changes=results['total_changes'],
             total_comments=0,
             conflicts=0,
             unresolved_conflicts=0,
             changes_by_source={},
             comments_by_source={})

    except Exception as e:
        import traceback
        emit("error", message=f"{str(e)}\n{traceback.format_exc()}")
        sys.exit(1)


if __name__ == "__main__":
    main()
