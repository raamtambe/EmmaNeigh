#!/usr/bin/env python3
"""
EmmaNeigh document editor.

Supports two modes:
- extract: read visible text from a DOCX or PDF so the LLM can plan edits
- apply: apply exact text substitutions to a copy of a DOCX or PDF

Usage:
  document_editor.py <config_json_path>
"""

from __future__ import annotations

import json
import os
import sys
from typing import Dict, Iterable, List, Sequence, Tuple

try:
    import fitz  # type: ignore
except ImportError:
    fitz = None

try:
    from docx import Document  # type: ignore
except ImportError:
    Document = None


def emit(payload: Dict) -> None:
    print(json.dumps(payload), flush=True)


def emit_progress(percent: int, message: str) -> None:
    emit({"type": "progress", "percent": percent, "message": message})


def emit_result(payload: Dict) -> None:
    emit({"type": "result", **payload})


def emit_error(message: str) -> None:
    emit({"type": "error", "message": message})


def normalize_replacements(items: Sequence[Dict]) -> List[Dict[str, str]]:
    normalized: List[Dict[str, str]] = []
    seen = set()
    for item in items or []:
        if not isinstance(item, dict):
            continue
        find_text = str(item.get("find_text") or "").strip()
        replace_text = str(item.get("replace_text") or "")
        reason = str(item.get("reason") or "").strip()
        if not find_text:
            continue
        key = (find_text, replace_text)
        if key in seen:
            continue
        seen.add(key)
        normalized.append({
            "find_text": find_text,
            "replace_text": replace_text,
            "reason": reason,
        })
    return normalized


def iter_table_paragraphs(table) -> Iterable:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def iter_docx_paragraphs(doc) -> Iterable:
    seen = set()

    def yield_paragraph(paragraph) -> Iterable:
        key = id(paragraph)
        if key in seen:
            return
        seen.add(key)
        yield paragraph

    for paragraph in doc.paragraphs:
        yield from yield_paragraph(paragraph)
    for table in doc.tables:
        for paragraph in iter_table_paragraphs(table):
            yield from yield_paragraph(paragraph)

    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                yield from yield_paragraph(paragraph)
            for table in container.tables:
                for paragraph in iter_table_paragraphs(table):
                    yield from yield_paragraph(paragraph)


def extract_text_from_docx(file_path: str) -> str:
    if Document is None:
        raise ImportError("python-docx is not installed")
    doc = Document(file_path)
    blocks: List[str] = []
    for paragraph in iter_docx_paragraphs(doc):
        text = (paragraph.text or "").strip()
        if text:
            blocks.append(text)
    return "\n".join(blocks)


def _run_boundaries(paragraph) -> List[Tuple[int, int, int]]:
    boundaries: List[Tuple[int, int, int]] = []
    cursor = 0
    for idx, run in enumerate(paragraph.runs):
        run_text = run.text or ""
        next_cursor = cursor + len(run_text)
        boundaries.append((idx, cursor, next_cursor))
        cursor = next_cursor
    return boundaries


def _find_run_position(boundaries: Sequence[Tuple[int, int, int]], char_index: int) -> Tuple[int, int]:
    for run_idx, start, end in boundaries:
        if start <= char_index < end:
            return run_idx, char_index - start
    if boundaries:
        run_idx, start, end = boundaries[-1]
        return run_idx, max(0, min(char_index - start, max(0, end - start)))
    return 0, 0


def replace_text_in_paragraph(paragraph, find_text: str, replace_text: str) -> int:
    if not find_text:
        return 0

    if not paragraph.runs:
        original = paragraph.text or ""
        count = original.count(find_text)
        if count:
            paragraph.text = original.replace(find_text, replace_text)
        return count

    replacements = 0
    while True:
        full_text = "".join(run.text or "" for run in paragraph.runs)
        match_index = full_text.find(find_text)
        if match_index < 0:
            break

        boundaries = _run_boundaries(paragraph)
        start_run_idx, start_offset = _find_run_position(boundaries, match_index)
        end_run_idx, end_offset = _find_run_position(boundaries, match_index + len(find_text) - 1)

        start_run = paragraph.runs[start_run_idx]
        end_run = paragraph.runs[end_run_idx]
        prefix = (start_run.text or "")[:start_offset]
        suffix = (end_run.text or "")[end_offset + 1:]
        start_run.text = prefix + replace_text + suffix

        for idx in range(start_run_idx + 1, end_run_idx + 1):
            paragraph.runs[idx].text = ""

        replacements += 1

    return replacements


def apply_docx_replacements(input_path: str, output_path: str, replacements: Sequence[Dict[str, str]]) -> Dict:
    if Document is None:
        raise ImportError("python-docx is not installed")
    doc = Document(input_path)
    total = 0
    applied: List[Dict] = []

    for item in replacements:
        count = 0
        for paragraph in iter_docx_paragraphs(doc):
            count += replace_text_in_paragraph(paragraph, item["find_text"], item["replace_text"])
        total += count
        applied.append({
            "find_text": item["find_text"],
            "replace_text": item["replace_text"],
            "reason": item.get("reason", ""),
            "count": count,
        })

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return {
        "success": True,
        "output_path": output_path,
        "applied_edits": applied,
        "total_replacements": total,
        "warnings": [],
    }


def extract_text_from_pdf(file_path: str) -> str:
    if fitz is None:
        raise ImportError("PyMuPDF is not installed")
    doc = fitz.open(file_path)
    try:
        pages: List[str] = []
        for page_index, page in enumerate(doc):
            text = (page.get_text("text") or "").strip()
            if text:
                pages.append(f"[Page {page_index + 1}]\n{text}")
        return "\n\n".join(pages)
    finally:
        doc.close()


def apply_pdf_replacements(input_path: str, output_path: str, replacements: Sequence[Dict[str, str]]) -> Dict:
    if fitz is None:
        raise ImportError("PyMuPDF is not installed")

    doc = fitz.open(input_path)
    warnings: List[str] = []
    applied: List[Dict] = []
    total = 0

    try:
        for item in replacements:
            count = 0
            find_text = item["find_text"]
            replace_text = item["replace_text"]

            for page_index, page in enumerate(doc):
                rects = page.search_for(find_text)
                if not rects:
                    continue

                saved_rects = [fitz.Rect(rect) for rect in rects]
                for rect in saved_rects:
                    page.add_redact_annot(rect, fill=(1, 1, 1))
                page.apply_redactions()

                for rect in saved_rects:
                    fontsize = max(8, min(18, rect.height * 0.78))
                    overflow = page.insert_textbox(
                        rect,
                        replace_text,
                        fontsize=fontsize,
                        fontname="helv",
                        color=(0, 0, 0),
                        align=0,
                    )
                    if overflow < 0:
                        warnings.append(
                            f'Page {page_index + 1}: replacement "{find_text}" may not have fully fit in its original box.'
                        )
                    count += 1

            total += count
            applied.append({
                "find_text": find_text,
                "replace_text": replace_text,
                "reason": item.get("reason", ""),
                "count": count,
            })

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
    finally:
        doc.close()

    return {
        "success": True,
        "output_path": output_path,
        "applied_edits": applied,
        "total_replacements": total,
        "warnings": warnings,
    }


def extract_document_text(input_path: str, max_chars: int) -> Dict:
    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".pdf":
        text = extract_text_from_pdf(input_path)
    elif ext == ".docx":
        text = extract_text_from_docx(input_path)
    else:
        raise ValueError(f"Unsupported file type for document editing: {ext or '(none)'}")

    original_length = len(text or "")
    text = text or ""
    if max_chars > 0 and original_length > max_chars:
        head = text[: max_chars // 2]
        tail = text[-(max_chars - len(head)) :]
        text = head + "\n\n...[truncated]...\n\n" + tail
        truncated = True
    else:
        truncated = False

    return {
        "success": True,
        "text": text,
        "truncated": truncated,
        "full_text_length": original_length,
    }


def apply_document_replacements(input_path: str, output_path: str, replacements: Sequence[Dict[str, str]]) -> Dict:
    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".pdf":
        return apply_pdf_replacements(input_path, output_path, replacements)
    if ext == ".docx":
        return apply_docx_replacements(input_path, output_path, replacements)
    raise ValueError(f"Unsupported file type for document editing: {ext or '(none)'}")


def main() -> None:
    if len(sys.argv) < 2:
        emit_error("Usage: document_editor.py <config_json_path>")
        sys.exit(1)

    config_path = sys.argv[1]
    try:
        with open(config_path, "r", encoding="utf-8") as handle:
            config = json.load(handle)
    except Exception as exc:
        emit_error(f"Failed to read config: {exc}")
        sys.exit(1)

    mode = str(config.get("mode") or "apply").strip().lower()
    input_path = str(config.get("input_path") or "").strip()
    if not input_path or not os.path.isfile(input_path):
        emit_error("A valid input_path is required.")
        sys.exit(1)

    try:
        if mode == "extract":
            emit_progress(15, "Extracting document text...")
            max_chars = int(config.get("max_chars") or 60000)
            result = extract_document_text(input_path, max_chars=max_chars)
            emit_result(result)
            return

        if mode != "apply":
            raise ValueError(f"Unsupported mode: {mode}")

        replacements = normalize_replacements(config.get("replacements") or [])
        if not replacements:
            raise ValueError("At least one replacement is required.")

        output_path = str(config.get("output_path") or "").strip()
        if not output_path:
            raise ValueError("output_path is required in apply mode.")

        emit_progress(20, "Applying document edits...")
        result = apply_document_replacements(input_path, output_path, replacements)
        emit_result(result)
    except Exception as exc:
        emit_error(str(exc))
        sys.exit(1)


if __name__ == "__main__":
    main()
