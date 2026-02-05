#!/usr/bin/env python3
"""
EmmaNeigh - Incumbency Certificate Parser
Extracts authorized signers (name, title) from incumbency certificates.
"""

import fitz
import os
import re
import sys
import json
from docx import Document


# Keywords for name column in incumbency tables
NAME_HEADERS = ["NAME", "OFFICER", "DIRECTOR", "AUTHORIZED"]
TITLE_HEADERS = ["TITLE", "POSITION", "OFFICE", "CAPACITY"]


def emit(msg_type, **kwargs):
    """Output JSON message to stdout for the Electron app."""
    print(json.dumps({"type": msg_type, **kwargs}), flush=True)


def find_column_index(headers, keywords):
    """Find the index of a column matching any keyword."""
    if not headers:
        return None

    headers_upper = [(h.upper().strip() if h else "") for h in headers]

    for i, h in enumerate(headers_upper):
        for keyword in keywords:
            if keyword in h:
                return i
    return None


def is_incumbency_table(headers):
    """Check if table headers indicate an incumbency/officer table."""
    if not headers:
        return False

    headers_upper = [(h.upper().strip() if h else "") for h in headers]

    # Must have a name-like column
    has_name = any(
        any(nh in h for nh in NAME_HEADERS)
        for h in headers_upper
    )

    # Should have title column
    has_title = any(
        any(th in h for th in TITLE_HEADERS)
        for h in headers_upper
    )

    return has_name and has_title


def extract_entity_name_from_text(text):
    """
    Try to extract entity name from incumbency certificate text.
    Looks for patterns like "I, [name], Secretary of [ENTITY NAME]"
    or "[ENTITY NAME] (the 'Company')"
    """
    # Pattern 1: "of [ENTITY NAME], a [state] [type]"
    match = re.search(
        r'of\s+([A-Z][A-Za-z0-9\s,\.]+(?:LLC|Inc|Corp|Corporation|LP|LLP|L\.P\.|L\.L\.C\.))',
        text
    )
    if match:
        return match.group(1).strip()

    # Pattern 2: "[ENTITY NAME] (the "Company")" or "(the "Corporation")"
    match = re.search(
        r'([A-Z][A-Za-z0-9\s,\.]+(?:LLC|Inc|Corp|Corporation|LP|LLP))\s*\((?:the\s+)?["\']?(?:Company|Corporation|Borrower)',
        text,
        re.IGNORECASE
    )
    if match:
        return match.group(1).strip()

    # Pattern 3: "CERTIFICATE OF INCUMBENCY OF [ENTITY]"
    match = re.search(
        r'CERTIFICATE\s+OF\s+(?:INCUMBENCY|SECRETARY)\s+OF\s+([A-Z][A-Za-z0-9\s,\.]+)',
        text,
        re.IGNORECASE
    )
    if match:
        entity = match.group(1).strip()
        # Clean up - remove trailing words like "I hereby"
        entity = re.sub(r'\s+I\s+hereby.*$', '', entity, flags=re.IGNORECASE)
        if len(entity) > 3:
            return entity

    return None


def is_valid_name(name):
    """Check if a string looks like a person's name."""
    if not name or len(name) < 3:
        return False

    # Filter out entity indicators
    entity_terms = ["LLC", "INC", "CORP", "CORPORATION", "LP", "LLP", "TRUST", "N.A."]
    name_upper = name.upper()
    if any(term in name_upper for term in entity_terms):
        return False

    # Should have 2-4 words
    words = name.split()
    if not (2 <= len(words) <= 5):
        return False

    # First word should start with capital
    if not name[0].isupper():
        return False

    return True


def extract_signers_from_table(table_data):
    """Extract signer information from a table."""
    if not table_data or len(table_data) < 2:
        return []

    headers = table_data[0]
    name_idx = find_column_index(headers, NAME_HEADERS)
    title_idx = find_column_index(headers, TITLE_HEADERS)

    # If no explicit name column, try first column
    if name_idx is None:
        name_idx = 0

    signers = []
    for row in table_data[1:]:
        if name_idx >= len(row):
            continue

        name = row[name_idx].strip() if row[name_idx] else ""

        # Get title if available
        title = ""
        if title_idx is not None and title_idx < len(row):
            title = row[title_idx].strip() if row[title_idx] else ""

        if name and is_valid_name(name):
            signers.append({
                "name": name,
                "title": title
            })

    return signers


def parse_incumbency_pdf(filepath):
    """
    Parse an incumbency certificate PDF.

    Returns:
        dict: {
            "entity_name": str or None,
            "signers": [{"name": str, "title": str}],
            "source_file": str
        }
    """
    doc = fitz.open(filepath)
    entity_name = None
    all_signers = []

    # First pass: try to get entity name from text
    full_text = ""
    for page in doc:
        full_text += page.get_text()

    entity_name = extract_entity_name_from_text(full_text)

    # Second pass: find tables and extract signers
    for page in doc:
        try:
            tables = page.find_tables()
            for table in tables.tables:
                data = table.extract()
                if data and len(data) > 0:
                    if is_incumbency_table(data[0]):
                        signers = extract_signers_from_table(data)
                        all_signers.extend(signers)
        except Exception:
            pass

    # If no tables found, try text-based extraction
    if not all_signers:
        # Look for patterns like "Name: John Smith" "Title: CEO"
        lines = full_text.split('\n')
        current_name = None

        for line in lines:
            line = line.strip()

            # Look for Name: pattern
            match = re.match(r'Name:\s*(.+)', line, re.IGNORECASE)
            if match:
                current_name = match.group(1).strip()
                continue

            # Look for Title: pattern
            match = re.match(r'Title:\s*(.+)', line, re.IGNORECASE)
            if match and current_name:
                title = match.group(1).strip()
                if is_valid_name(current_name):
                    all_signers.append({
                        "name": current_name,
                        "title": title
                    })
                current_name = None

    doc.close()

    return {
        "entity_name": entity_name,
        "signers": all_signers,
        "source_file": os.path.basename(filepath)
    }


def parse_incumbency_docx(filepath):
    """
    Parse an incumbency certificate Word document.

    Returns same format as parse_incumbency_pdf.
    """
    doc = Document(filepath)
    entity_name = None
    all_signers = []

    # Get full text for entity name extraction
    full_text = "\n".join([p.text for p in doc.paragraphs])
    entity_name = extract_entity_name_from_text(full_text)

    # Extract from tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(row_data)

        if rows and len(rows) > 0:
            if is_incumbency_table(rows[0]):
                signers = extract_signers_from_table(rows)
                all_signers.extend(signers)

    return {
        "entity_name": entity_name,
        "signers": all_signers,
        "source_file": os.path.basename(filepath)
    }


def parse_incumbency(filepath):
    """
    Parse an incumbency certificate (PDF or DOCX).
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.pdf':
        return parse_incumbency_pdf(filepath)
    elif ext == '.docx':
        return parse_incumbency_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def main():
    """CLI entry point for testing."""
    if len(sys.argv) < 2:
        emit("error", message="Usage: incumbency_parser.py <incumbency_file>")
        sys.exit(1)

    filepath = sys.argv[1]

    if not os.path.isfile(filepath):
        emit("error", message=f"File not found: {filepath}")
        sys.exit(1)

    try:
        emit("progress", percent=10, message="Parsing incumbency certificate...")
        result = parse_incumbency(filepath)

        emit("progress", percent=100, message="Complete!")
        emit("result", success=True, **result)

    except Exception as e:
        emit("error", message=str(e))
        sys.exit(1)


if __name__ == "__main__":
    main()
