#!/usr/bin/env python3
"""
EmmaNeigh - Signature Packet Shell Generator
v5.1.1: Generate a combined signature packet with all pages requiring signatures.

Creates a single "shell" packet containing all unique signature pages from multiple documents,
with footers added to identify each document.
"""

import fitz
import os
import pandas as pd
import re
import sys
import json
import hashlib
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Import shared functions from signature_packets
try:
    from signature_packets import (
        emit, normalize_name, is_probable_person,
        detect_signature_page_extended, extract_footer_from_pdf_page,
        extract_footer_from_docx, extract_text_from_docx, extract_tables_from_docx
    )
except ImportError:
    # Fallback definitions if running standalone
    def emit(msg_type, **kwargs):
        print(json.dumps({"type": msg_type, **kwargs}), flush=True)


# ========== DOCUMENT TITLE EXTRACTION ==========

def extract_document_title_from_pdf(doc_path):
    """
    Extract document title from PDF first page.
    Looks for: centered text, larger font, or known document types.
    """
    try:
        doc = fitz.open(doc_path)
        if doc.page_count == 0:
            doc.close()
            return os.path.splitext(os.path.basename(doc_path))[0]

        first_page = doc[0]
        blocks = first_page.get_text("dict")["blocks"]

        # Find text blocks and sort by font size (larger = more likely title)
        text_spans = []
        for block in blocks:
            if "lines" in block:
                for line in block["lines"]:
                    for span in line["spans"]:
                        if span.get("text", "").strip():
                            text_spans.append({
                                "text": span["text"].strip(),
                                "size": span.get("size", 12),
                                "y": line["bbox"][1]  # vertical position
                            })

        # Sort by size descending, then by position (top of page first)
        text_spans.sort(key=lambda x: (-x["size"], x["y"]))

        # Look for known document types in the largest text
        doc_types = [
            'CREDIT AGREEMENT', 'GUARANTEE', 'GUARANTY', 'PLEDGE AGREEMENT',
            'SECURITY AGREEMENT', 'LOAN AGREEMENT', 'NOTE PURCHASE AGREEMENT',
            'INDENTURE', 'AMENDMENT', 'CONSENT', 'JOINDER', 'ASSIGNMENT',
            'PROMISSORY NOTE', 'CERTIFICATE', 'INCUMBENCY', 'RESOLUTION',
            'OFFICER\'S CERTIFICATE', 'SECRETARY\'S CERTIFICATE'
        ]

        # Check top 5 largest text spans for document type keywords
        for span in text_spans[:10]:
            text_upper = span["text"].upper()
            for doc_type in doc_types:
                if doc_type in text_upper:
                    doc.close()
                    return span["text"]

        # Fallback: use largest text from top quarter of page
        page_height = first_page.rect.height
        top_quarter_spans = [s for s in text_spans if s["y"] < page_height / 4]
        if top_quarter_spans:
            # Get the largest text in top quarter
            largest = max(top_quarter_spans, key=lambda x: x["size"])
            if largest["size"] > 14:  # Likely a title if larger than body text
                doc.close()
                return largest["text"]

        doc.close()
    except Exception:
        pass

    # Fallback: use filename
    return os.path.splitext(os.path.basename(doc_path))[0]


def extract_document_title_from_docx(doc_path):
    """
    Extract document title from DOCX first paragraphs.
    Looks for: centered text, Title style, or known document types.
    """
    try:
        doc = Document(doc_path)

        # Check document core properties for title
        if doc.core_properties.title:
            return doc.core_properties.title

        # Look in first 10 paragraphs
        for i, para in enumerate(doc.paragraphs[:10]):
            text = para.text.strip()
            if not text:
                continue

            # Check if centered
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                if len(text) > 5:  # Not just a page number
                    return text

            # Check for Title style
            if para.style and 'Title' in para.style.name:
                return text

            # Check for document type keywords
            text_upper = text.upper()
            doc_types = [
                'CREDIT AGREEMENT', 'GUARANTEE', 'GUARANTY', 'PLEDGE AGREEMENT',
                'SECURITY AGREEMENT', 'LOAN AGREEMENT', 'NOTE PURCHASE AGREEMENT',
                'INDENTURE', 'AMENDMENT', 'CONSENT', 'JOINDER', 'ASSIGNMENT',
                'PROMISSORY NOTE', 'CERTIFICATE', 'INCUMBENCY', 'RESOLUTION'
            ]
            for doc_type in doc_types:
                if doc_type in text_upper:
                    return text

    except Exception:
        pass

    # Fallback: use filename
    return os.path.splitext(os.path.basename(doc_path))[0]


def extract_document_title(doc_path):
    """Extract document title based on file type."""
    if doc_path.lower().endswith('.pdf'):
        return extract_document_title_from_pdf(doc_path)
    elif doc_path.lower().endswith('.docx'):
        return extract_document_title_from_docx(doc_path)
    return os.path.splitext(os.path.basename(doc_path))[0]


# ========== PAGE CONTENT HASHING (for deduplication) ==========

def hash_page_content(text):
    """Create a hash of page content for deduplication."""
    # Normalize text: remove whitespace variations
    normalized = re.sub(r'\s+', ' ', text.strip().lower())
    return hashlib.md5(normalized.encode()).hexdigest()


# ========== FOOTER INSERTION ==========

def add_footer_to_pdf_page(page, footer_text):
    """
    Add footer text to a PDF page.
    Inserts at bottom center of the page.
    """
    rect = page.rect
    footer_rect = fitz.Rect(
        rect.width * 0.1,  # 10% from left
        rect.height - 40,  # 40 points from bottom
        rect.width * 0.9,  # 90% width
        rect.height - 20   # 20 points from bottom
    )

    # Insert footer text
    page.insert_textbox(
        footer_rect,
        footer_text,
        fontsize=9,
        fontname="helv",
        align=fitz.TEXT_ALIGN_CENTER
    )


def add_footer_to_docx_section(doc, footer_text):
    """
    Add footer to DOCX document section.
    """
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        # Clear existing footer
        for para in footer.paragraphs:
            para.clear()

        # Add new footer
        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()

        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(footer_text)
        run.font.size = Pt(9)


# ========== SIGNATURE PAGE COLLECTION ==========

def collect_signature_pages_from_pdf(doc_path, doc_title):
    """
    Collect all signature pages from a PDF document.

    Returns:
        List of dicts: [{
            'page_num': int,
            'signers': set,
            'footer': str,
            'content_hash': str,
            'has_footer': bool,
            'doc_title': str,
            'doc_path': str
        }]
    """
    pages = []

    try:
        doc = fitz.open(doc_path)

        for page_num, page in enumerate(doc, start=1):
            text = page.get_text()

            # Get tables
            tables_data = []
            try:
                tables = page.find_tables()
                for table in tables.tables:
                    data = table.extract()
                    if data:
                        tables_data.append(data)
            except Exception:
                pass

            # Detect signature page
            is_sig_page, signers, method = detect_signature_page_extended(text, tables_data)

            if is_sig_page:
                footer = extract_footer_from_pdf_page(page)
                has_footer = bool(footer and 'SIGNATURE PAGE' in footer.upper())

                pages.append({
                    'page_num': page_num,
                    'signers': signers,
                    'footer': footer,
                    'content_hash': hash_page_content(text),
                    'has_footer': has_footer,
                    'doc_title': doc_title,
                    'doc_path': doc_path,
                    'format': 'pdf'
                })

        doc.close()
    except Exception as e:
        emit("progress", percent=0, message=f"Warning: {doc_path} - {str(e)}")

    return pages


def collect_signature_pages_from_docx(doc_path, doc_title):
    """
    Collect signature page info from a DOCX document.
    DOCX doesn't have pages, so we treat the whole document as one "page" if it has signatures.
    """
    pages = []

    try:
        text = extract_text_from_docx(doc_path)
        tables_data = extract_tables_from_docx(doc_path)

        is_sig_page, signers, method = detect_signature_page_extended(text, tables_data)

        if is_sig_page:
            footer = extract_footer_from_docx(doc_path)
            has_footer = bool(footer and 'SIGNATURE PAGE' in footer.upper())

            pages.append({
                'page_num': 1,
                'signers': signers,
                'footer': footer,
                'content_hash': hash_page_content(text),
                'has_footer': has_footer,
                'doc_title': doc_title,
                'doc_path': doc_path,
                'format': 'docx'
            })

    except Exception as e:
        emit("progress", percent=0, message=f"Warning: {doc_path} - {str(e)}")

    return pages


# ========== PACKET SHELL CREATION ==========

def create_pdf_shell(pages, output_path, add_missing_footers=True):
    """
    Create a combined PDF signature packet shell.

    Args:
        pages: List of page dicts from collect_signature_pages_*
        output_path: Where to save the output PDF
        add_missing_footers: Whether to add footers to pages missing them

    Returns:
        Path to created PDF or None
    """
    try:
        shell_doc = fitz.open()

        for page_info in pages:
            if page_info['format'] != 'pdf':
                continue

            src_doc = fitz.open(page_info['doc_path'])
            page_idx = page_info['page_num'] - 1

            if page_idx < src_doc.page_count:
                shell_doc.insert_pdf(src_doc, from_page=page_idx, to_page=page_idx)

                # Add footer if missing
                if add_missing_footers and not page_info['has_footer']:
                    # Get the newly added page (last page)
                    new_page = shell_doc[-1]
                    footer_text = f"Signature Page to {page_info['doc_title']}"
                    add_footer_to_pdf_page(new_page, footer_text)

            src_doc.close()

        if shell_doc.page_count > 0:
            shell_doc.save(output_path)
            page_count = shell_doc.page_count
            shell_doc.close()
            return output_path, page_count

        shell_doc.close()
    except Exception as e:
        emit("progress", percent=0, message=f"Error creating PDF shell: {str(e)}")

    return None, 0


def create_docx_shell(pages, output_path, add_missing_footers=True):
    """
    Create a combined DOCX signature packet shell.

    Args:
        pages: List of page dicts from collect_signature_pages_*
        output_path: Where to save the output DOCX
        add_missing_footers: Whether to add footers to pages missing them

    Returns:
        Path to created DOCX or None
    """
    try:
        shell_doc = Document()

        # Add title
        title_para = shell_doc.add_paragraph()
        title_run = title_para.add_run("SIGNATURE PACKET SHELL")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        shell_doc.add_paragraph()  # Blank line

        docs_added = 0

        for page_info in pages:
            if page_info['format'] != 'docx':
                continue

            try:
                src_doc = Document(page_info['doc_path'])

                # Add document separator
                sep_para = shell_doc.add_paragraph()
                sep_para.add_run("─" * 50)

                # Add document title
                doc_title_para = shell_doc.add_paragraph()
                doc_title_run = doc_title_para.add_run(f"Document: {page_info['doc_title']}")
                doc_title_run.bold = True

                shell_doc.add_paragraph()

                # Copy paragraphs
                for para in src_doc.paragraphs:
                    new_para = shell_doc.add_paragraph()
                    try:
                        new_para.style = para.style
                    except Exception:
                        pass

                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        if run.font.size:
                            new_run.font.size = run.font.size

                # Copy tables
                for table in src_doc.tables:
                    new_table = shell_doc.add_table(
                        rows=len(table.rows),
                        cols=len(table.columns)
                    )
                    new_table.style = 'Table Grid'

                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                                new_table.rows[i].cells[j].text = cell.text

                # Add footer if missing
                if add_missing_footers and not page_info['has_footer']:
                    footer_para = shell_doc.add_paragraph()
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    footer_run = footer_para.add_run(
                        f"Signature Page to {page_info['doc_title']}"
                    )
                    footer_run.font.size = Pt(9)

                docs_added += 1
                shell_doc.add_page_break()

            except Exception as e:
                continue

        if docs_added > 0:
            # Add footer to document sections
            if add_missing_footers:
                add_footer_to_docx_section(shell_doc, "SIGNATURE PACKET SHELL")

            shell_doc.save(output_path)
            return output_path, docs_added

    except Exception as e:
        emit("progress", percent=0, message=f"Error creating DOCX shell: {str(e)}")

    return None, 0


# ========== MAIN ==========

def main():
    if len(sys.argv) < 2:
        emit("error", message="No input provided.")
        sys.exit(1)

    # Parse arguments
    config_path = None
    output_format = 'both'  # 'pdf', 'docx', or 'both'

    if sys.argv[1] == '--config':
        if len(sys.argv) < 3:
            emit("error", message="No config file provided.")
            sys.exit(1)
        config_path = sys.argv[2]
    else:
        emit("error", message="Usage: packet_shell_generator.py --config <config.json>")
        sys.exit(1)

    # Load config
    try:
        with open(config_path, 'r') as f:
            config = json.load(f)
        file_paths = config.get('files', [])
        output_format = config.get('output_format', 'both')
        output_dir = config.get('output_dir', None)

        if not file_paths:
            emit("error", message="No files in config.")
            sys.exit(1)

        if not output_dir:
            import tempfile
            output_dir = tempfile.mkdtemp(prefix='emmaneigh_shell_')

    except Exception as e:
        emit("error", message=f"Failed to read config: {str(e)}")
        sys.exit(1)

    # Set up output directory
    os.makedirs(output_dir, exist_ok=True)

    # Filter to valid files
    document_files = [
        f for f in file_paths
        if os.path.isfile(f) and f.lower().endswith(('.pdf', '.docx'))
    ]

    if not document_files:
        emit("error", message="No PDF or Word files found.")
        sys.exit(1)

    total = len(document_files)
    emit("progress", percent=0, message=f"Found {total} documents")

    # Collect all signature pages
    all_pages = []
    seen_hashes = set()  # For deduplication

    for idx, doc_path in enumerate(document_files):
        percent = int((idx / total) * 60)
        filename = os.path.basename(doc_path)
        emit("progress", percent=percent, message=f"Scanning {filename}")

        # Extract document title
        doc_title = extract_document_title(doc_path)

        # Collect signature pages
        if doc_path.lower().endswith('.pdf'):
            pages = collect_signature_pages_from_pdf(doc_path, doc_title)
        else:
            pages = collect_signature_pages_from_docx(doc_path, doc_title)

        # Add to collection with deduplication
        for page in pages:
            if page['content_hash'] not in seen_hashes:
                seen_hashes.add(page['content_hash'])
                all_pages.append(page)

    if not all_pages:
        emit("error", message="No signature pages detected in any documents.")
        sys.exit(1)

    emit("progress", percent=65, message=f"Found {len(all_pages)} unique signature pages")

    # Create output packets
    results = []

    # Separate PDF and DOCX pages
    pdf_pages = [p for p in all_pages if p['format'] == 'pdf']
    docx_pages = [p for p in all_pages if p['format'] == 'docx']

    # Create PDF shell
    if output_format in ('pdf', 'both') and pdf_pages:
        emit("progress", percent=75, message="Creating PDF shell...")
        pdf_path = os.path.join(output_dir, "SIGNATURE_PACKET_SHELL.pdf")
        result_path, page_count = create_pdf_shell(pdf_pages, pdf_path)
        if result_path:
            results.append({
                'format': 'pdf',
                'path': result_path,
                'pages': page_count
            })

    # Create DOCX shell
    if output_format in ('docx', 'both') and docx_pages:
        emit("progress", percent=85, message="Creating DOCX shell...")
        docx_path = os.path.join(output_dir, "SIGNATURE_PACKET_SHELL.docx")
        result_path, doc_count = create_docx_shell(docx_pages, docx_path)
        if result_path:
            results.append({
                'format': 'docx',
                'path': result_path,
                'documents': doc_count
            })

    # Create summary index
    emit("progress", percent=95, message="Creating index...")
    index_rows = []
    for page in all_pages:
        for signer in page['signers']:
            index_rows.append({
                'Document': page['doc_title'],
                'Page': page['page_num'],
                'Signer': signer,
                'Footer': page['footer'],
                'Has Footer': 'Yes' if page['has_footer'] else 'No'
            })

    if index_rows:
        df = pd.DataFrame(index_rows)
        df = df.sort_values(['Document', 'Page', 'Signer'])
        index_path = os.path.join(output_dir, "SIGNATURE_SHELL_INDEX.xlsx")
        df.to_excel(index_path, index=False)
        results.append({
            'format': 'xlsx',
            'path': index_path,
            'rows': len(index_rows)
        })

    emit("progress", percent=100, message="Complete!")
    emit("result",
         success=True,
         outputPath=output_dir,
         totalPages=len(all_pages),
         results=results)


if __name__ == "__main__":
    main()
