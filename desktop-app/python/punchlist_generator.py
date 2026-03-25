#!/usr/bin/env python3
"""
Punchlist Generator - Generate daily punchlist from transaction checklist

This module parses a transaction checklist (Word document) and generates
a formatted punchlist showing only open/pending items organized by status.

A punchlist is a daily working document that shows:
- What items are still pending/open
- What needs attention today
- Grouped by status (pending draft, with counsel, awaiting signature, etc.)

Usage:
    python punchlist_generator.py <checklist_path> <output_folder> [status_filters_json] <api_key> [provider] [model] [provider_base_url]
"""

import sys
import os
import json
import re
import socket
import time
import urllib.error
import urllib.request
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

DEFAULT_CLAUDE_MODEL = "claude-sonnet-4-20250514"
DEFAULT_OPENAI_MODEL = "gpt-4.1-mini"
DEFAULT_HARVEY_BASE_URL = "https://api.harvey.ai"
DEFAULT_LOCALAI_BASE_URL = "http://127.0.0.1:11435"
DEFAULT_OLLAMA_BASE_URL = "http://127.0.0.1:11434"
DEFAULT_LMSTUDIO_BASE_URL = "http://127.0.0.1:1234"
DEFAULT_LOCALAI_MODEL = "qwen2.5:1.5b"
DEFAULT_OLLAMA_MODEL = "llama3.1:8b"
DEFAULT_LMSTUDIO_MODEL = "local-model"
MAX_HTTP_ATTEMPTS = 3
RETRYABLE_HTTP_STATUS = {408, 425, 429, 500, 502, 503, 504}
RETRYABLE_ERROR_MARKERS = (
    "timeout",
    "temporarily unavailable",
    "connection reset",
    "connection aborted",
    "network is unreachable",
    "name or service not known",
    "temporary failure in name resolution",
)


def parse_json_response_text(response_text):
    """
    Parse JSON from plain text or markdown fenced code block output.
    """
    text = (response_text or "").strip()
    if not text:
        raise ValueError("LLM returned empty response")

    if text.startswith("```"):
        lines = text.split("\n")
        json_lines = []
        in_block = False
        for line in lines:
            marker = line.strip()
            if marker.startswith("```"):
                if not in_block:
                    in_block = True
                    continue
                break
            if in_block:
                json_lines.append(line)
        if json_lines:
            text = "\n".join(json_lines).strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        fragment = extract_json_object_fragment(text)
        if fragment:
            return json.loads(fragment)
        raise ValueError("LLM response was not valid JSON.")


def extract_json_object_fragment(text):
    """Extract the first complete JSON object from a larger text blob."""
    source = str(text or "")
    start = source.find("{")
    if start == -1:
        return None

    depth = 0
    in_string = False
    escaped = False

    for idx in range(start, len(source)):
        ch = source[idx]

        if in_string:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == '"':
                in_string = False
            continue

        if ch == '"':
            in_string = True
            continue
        if ch == "{":
            depth += 1
            continue
        if ch == "}":
            depth -= 1
            if depth == 0:
                return source[start:idx + 1]

    return None


def canonical_doc_key(value):
    """Normalize document names for key matching."""
    text = re.sub(r'\s+', ' ', str(value or '')).strip().lower()
    return text


def normalize_api_key(api_key):
    return str(api_key or "").strip()


def normalize_provider(provider):
    value = str(provider or "").strip().lower()
    if value == "claude":
        return "anthropic"
    if value in ("managed", "localai", "anthropic", "openai", "harvey", "ollama", "lmstudio"):
        return value
    if value in ("managed ai", "managed-ai", "managed remote", "managed-remote", "firm", "firm-managed", "firmmanaged"):
        return "managed"
    if value in ("local", "local ai", "local-ai", "local pack", "localpack", "built-in", "builtin"):
        return "localai"
    if value in ("lm studio", "lm-studio"):
        return "lmstudio"
    return "anthropic"


def infer_provider_from_api_key(api_key):
    value = normalize_api_key(api_key).lower()
    if not value:
        return None
    if value.startswith("sk-ant-"):
        return "anthropic"
    if value.startswith("harvey_") or value.startswith("hv_"):
        return "harvey"
    if value.startswith("sk-proj-") or value.startswith("sk-") or value.startswith("sess-"):
        return "openai"
    return None


def resolve_provider(provider, api_key):
    preferred = normalize_provider(provider)
    if preferred in ("managed", "localai", "ollama", "lmstudio"):
        return preferred
    inferred = infer_provider_from_api_key(api_key)
    return inferred or preferred


def provider_requires_api_key(provider):
    return normalize_provider(provider) in ("anthropic", "openai", "harvey")


def extract_openai_text(content):
    if isinstance(content, str):
        return content.strip()
    if not isinstance(content, list):
        return ""
    parts = []
    for item in content:
        if isinstance(item, str):
            parts.append(item)
            continue
        if isinstance(item, dict):
            text = item.get("text")
            if isinstance(text, str):
                parts.append(text)
    return "\n".join(parts).strip()


def parse_error_detail(raw_text, network_error=None):
    if network_error:
        return str(network_error)

    try:
        payload = json.loads(raw_text or "{}")
    except Exception:
        payload = {}

    if isinstance(payload, dict):
        err = payload.get("error")
        if isinstance(err, dict) and err.get("message"):
            return str(err.get("message"))
        if payload.get("message"):
            return str(payload.get("message"))

    return (raw_text or "").strip()[:300] or "Unknown error"


def should_retry_request(status_code, detail):
    detail_text = str(detail or "").lower()
    if status_code in RETRYABLE_HTTP_STATUS:
        return True
    return any(marker in detail_text for marker in RETRYABLE_ERROR_MARKERS)


def backoff_sleep(attempt_idx):
    # Short incremental backoff for transient API/network failures.
    time.sleep(min(0.5 * (attempt_idx + 1), 2.0))


def format_provider_error(provider_label, status_code, detail):
    if status_code in (401, 403):
        return (
            f"{provider_label} authentication failed ({status_code}). "
            f"Check API key and account/model access. Details: {detail}"
        )
    if status_code == 429:
        return f"{provider_label} rate limit reached (429). Please retry. Details: {detail}"
    if status_code == 0:
        return f"{provider_label} network error: {detail}"
    return f"{provider_label} error ({status_code}): {detail}"


def is_likely_model_error(status_code, detail):
    detail_text = str(detail or "").lower()
    if status_code == 404:
        return True
    if status_code == 400 and "model" in detail_text:
        return True
    return "model" in detail_text and any(
        part in detail_text
        for part in ("not found", "invalid", "unsupported", "available", "access")
    )


def perform_http_request(url, headers, body_bytes, timeout=90):
    req = urllib.request.Request(url, data=body_bytes, headers=headers, method="POST")
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            status_code = resp.getcode()
            raw_text = resp.read().decode("utf-8", errors="replace")
            return status_code, raw_text, None
    except urllib.error.HTTPError as e:
        raw_text = e.read().decode("utf-8", errors="replace")
        return e.code, raw_text, None
    except urllib.error.URLError as e:
        reason = getattr(e, "reason", e)
        return 0, "", f"Network error: {reason}"
    except (socket.timeout, TimeoutError):
        return 0, "", "Network error: request timed out"
    except Exception as e:
        return 0, "", f"Network error: {e}"


def perform_json_request(url, headers, payload, timeout=90):
    request_headers = dict(headers or {})
    request_headers["Content-Type"] = "application/json"
    body_bytes = json.dumps(payload).encode("utf-8")
    status_code, raw_text, network_error = perform_http_request(
        url, request_headers, body_bytes, timeout=timeout
    )
    try:
        parsed = json.loads(raw_text) if raw_text else {}
    except Exception:
        parsed = {}
    return status_code, parsed, raw_text, network_error


def build_multipart_form_data(fields):
    boundary = "----EmmaNeighBoundary" + uuid.uuid4().hex
    chunks = []
    for key, value in fields.items():
        chunks.append(f"--{boundary}".encode("utf-8"))
        chunks.append(
            f'Content-Disposition: form-data; name="{key}"'.encode("utf-8")
        )
        chunks.append(b"")
        chunks.append(str(value).encode("utf-8"))
    chunks.append(f"--{boundary}--".encode("utf-8"))
    body = b"\r\n".join(chunks) + b"\r\n"
    return boundary, body


def call_anthropic_prompt(prompt, api_key, model_name):
    api_key = normalize_api_key(api_key)
    model_candidates = []
    for candidate in (
        model_name,
        os.environ.get("CLAUDE_MODEL"),
        DEFAULT_CLAUDE_MODEL,
        "claude-3-5-sonnet-20241022",
        "claude-3-5-haiku-20241022",
    ):
        candidate_value = str(candidate or "").strip()
        if candidate_value and candidate_value not in model_candidates:
            model_candidates.append(candidate_value)

    last_detail = "Unknown error"
    for candidate in model_candidates:
        model_unavailable = False
        for attempt_idx in range(MAX_HTTP_ATTEMPTS):
            payload = {
                "model": candidate,
                "max_tokens": 2048,
                "temperature": 0,
                "messages": [{"role": "user", "content": prompt}],
            }
            status_code, parsed, raw_text, network_error = perform_json_request(
                "https://api.anthropic.com/v1/messages",
                {
                    "x-api-key": api_key,
                    "anthropic-version": "2023-06-01",
                },
                payload,
            )

            if status_code != 200:
                detail = parse_error_detail(raw_text, network_error=network_error)
                last_detail = detail
                if should_retry_request(status_code, detail) and attempt_idx < (MAX_HTTP_ATTEMPTS - 1):
                    backoff_sleep(attempt_idx)
                    continue
                if is_likely_model_error(status_code, detail):
                    model_unavailable = True
                    break
                raise RuntimeError(format_provider_error("Anthropic", status_code, detail))

            content = parsed.get("content") if isinstance(parsed, dict) else None
            if isinstance(content, list) and content:
                first = content[0]
                if isinstance(first, dict) and first.get("text"):
                    return str(first.get("text")).strip(), candidate
            raise RuntimeError("Anthropic response was missing message content.")

        if model_unavailable:
            continue

    raise RuntimeError(
        f"No supported Claude model is available for this API key. Last error: {last_detail}"
    )


def call_openai_prompt(prompt, api_key, model_name):
    api_key = normalize_api_key(api_key)
    model_candidates = []
    for candidate in (
        model_name,
        os.environ.get("OPENAI_MODEL"),
        DEFAULT_OPENAI_MODEL,
        "gpt-4.1-mini",
        "gpt-4o-mini",
        "gpt-4.1",
        "gpt-4o",
    ):
        candidate_value = str(candidate or "").strip()
        if candidate_value and candidate_value not in model_candidates:
            model_candidates.append(candidate_value)

    last_detail = "Unknown error"
    for candidate in model_candidates:
        model_unavailable = False
        for attempt_idx in range(MAX_HTTP_ATTEMPTS):
            payload = {
                "model": candidate,
                "max_tokens": 2048,
                "temperature": 0,
                "messages": [{"role": "user", "content": prompt}],
            }
            status_code, parsed, raw_text, network_error = perform_json_request(
                "https://api.openai.com/v1/chat/completions",
                {"Authorization": f"Bearer {api_key}"},
                payload,
            )

            if status_code != 200:
                detail = parse_error_detail(raw_text, network_error=network_error)
                last_detail = detail
                if should_retry_request(status_code, detail) and attempt_idx < (MAX_HTTP_ATTEMPTS - 1):
                    backoff_sleep(attempt_idx)
                    continue
                if is_likely_model_error(status_code, detail):
                    model_unavailable = True
                    break
                raise RuntimeError(format_provider_error("OpenAI", status_code, detail))

            choices = parsed.get("choices") if isinstance(parsed, dict) else None
            message = choices[0].get("message") if isinstance(choices, list) and choices else {}
            content = extract_openai_text(message.get("content") if isinstance(message, dict) else "")
            if content:
                return content, candidate
            raise RuntimeError("OpenAI response was missing message content.")

        if model_unavailable:
            continue

    raise RuntimeError(
        f"No supported OpenAI model is available for this API key. Last error: {last_detail}"
    )


def call_openai_compatible_prompt(prompt, api_key, model_candidates, base_url, provider_label):
    api_key = normalize_api_key(api_key)
    candidates = []
    for candidate in model_candidates:
        candidate_value = str(candidate or "").strip()
        if candidate_value and candidate_value not in candidates:
            candidates.append(candidate_value)

    last_detail = "Unknown error"
    for candidate in candidates:
        model_unavailable = False
        for attempt_idx in range(MAX_HTTP_ATTEMPTS):
            payload = {
                "model": candidate,
                "max_tokens": 2048,
                "temperature": 0,
                "messages": [{"role": "user", "content": prompt}],
            }
            headers = {}
            if api_key:
                headers["Authorization"] = f"Bearer {api_key}"
            status_code, parsed, raw_text, network_error = perform_json_request(
                f"{base_url}/v1/chat/completions",
                headers,
                payload,
            )

            if status_code != 200:
                detail = parse_error_detail(raw_text, network_error=network_error)
                last_detail = detail
                if should_retry_request(status_code, detail) and attempt_idx < (MAX_HTTP_ATTEMPTS - 1):
                    backoff_sleep(attempt_idx)
                    continue
                if is_likely_model_error(status_code, detail):
                    model_unavailable = True
                    break
                raise RuntimeError(format_provider_error(provider_label, status_code, detail))

            choices = parsed.get("choices") if isinstance(parsed, dict) else None
            message = choices[0].get("message") if isinstance(choices, list) and choices else {}
            content = extract_openai_text(message.get("content") if isinstance(message, dict) else "")
            if content:
                return content, candidate
            raise RuntimeError(f"{provider_label} response was missing message content.")

        if model_unavailable:
            continue

    raise RuntimeError(
        f"No supported {provider_label} model is available. Last error: {last_detail}"
    )


def call_localai_prompt(prompt, api_key, model_name, provider_base_url):
    base_url = str(provider_base_url or os.environ.get("LOCALAI_BASE_URL") or DEFAULT_LOCALAI_BASE_URL).strip()
    if not base_url:
        base_url = DEFAULT_LOCALAI_BASE_URL
    if base_url.endswith("/"):
        base_url = base_url[:-1]
    return call_openai_compatible_prompt(
        prompt,
        api_key,
        [
            model_name,
            os.environ.get("LOCALAI_MODEL"),
            DEFAULT_LOCALAI_MODEL,
            "qwen2.5:1.5b",
            "qwen2.5:7b",
            "gemma3:1b",
            "gemma3",
        ],
        base_url,
        "Local AI",
    )


def call_ollama_prompt(prompt, api_key, model_name, provider_base_url):
    base_url = str(provider_base_url or os.environ.get("OLLAMA_BASE_URL") or DEFAULT_OLLAMA_BASE_URL).strip()
    if not base_url:
        base_url = DEFAULT_OLLAMA_BASE_URL
    if base_url.endswith("/"):
        base_url = base_url[:-1]
    return call_openai_compatible_prompt(
        prompt,
        api_key,
        [
            model_name,
            os.environ.get("OLLAMA_MODEL"),
            DEFAULT_OLLAMA_MODEL,
            "llama3.1:8b",
            "qwen2.5:7b",
        ],
        base_url,
        "Ollama",
    )


def call_lmstudio_prompt(prompt, api_key, model_name, provider_base_url):
    base_url = str(provider_base_url or os.environ.get("LMSTUDIO_BASE_URL") or DEFAULT_LMSTUDIO_BASE_URL).strip()
    if not base_url:
        base_url = DEFAULT_LMSTUDIO_BASE_URL
    if base_url.endswith("/"):
        base_url = base_url[:-1]
    return call_openai_compatible_prompt(
        prompt,
        api_key,
        [
            model_name,
            os.environ.get("LMSTUDIO_MODEL"),
            DEFAULT_LMSTUDIO_MODEL,
        ],
        base_url,
        "LM Studio",
    )


def call_managed_prompt(prompt, api_key, model_name, provider_base_url):
    endpoint = str(provider_base_url or os.environ.get("MANAGED_AI_PROMPT_URL") or "").strip()
    if not endpoint:
        raise RuntimeError("Managed AI prompt endpoint is not configured.")

    headers = {}
    token = normalize_api_key(api_key)
    if token:
        headers["Authorization"] = f"Bearer {token}"

    payload = {
        "prompt": prompt,
        "max_tokens": 2048,
    }
    if model_name:
        payload["model"] = str(model_name).strip()

    last_detail = "Unknown error"
    for attempt_idx in range(MAX_HTTP_ATTEMPTS):
        status_code, parsed, raw_text, network_error = perform_json_request(
            endpoint,
            headers,
            payload,
        )

        if status_code != 200:
            detail = parse_error_detail(raw_text, network_error=network_error)
            last_detail = detail
            if should_retry_request(status_code, detail) and attempt_idx < (MAX_HTTP_ATTEMPTS - 1):
                backoff_sleep(attempt_idx)
                continue
            raise RuntimeError(format_provider_error("Managed AI", status_code, detail))

        text = ""
        if isinstance(parsed, dict):
            text = str(parsed.get("text") or parsed.get("message") or "").strip()
        if text:
            return text, str(parsed.get("model") or model_name or "managed-ai").strip()
        raise RuntimeError("Managed AI response was missing message content.")

    raise RuntimeError(f"Managed AI error: {last_detail}")


def call_harvey_prompt(prompt, api_key, max_tokens, provider_base_url):
    api_key = normalize_api_key(api_key)
    base_url = str(
        provider_base_url or os.environ.get("HARVEY_BASE_URL") or DEFAULT_HARVEY_BASE_URL
    ).strip()
    if not base_url:
        base_url = DEFAULT_HARVEY_BASE_URL
    if base_url.endswith("/"):
        base_url = base_url[:-1]

    boundary, body = build_multipart_form_data(
        {
            "prompt": prompt,
            "mode": "assist",
            "stream": "false",
            "max_tokens": str(max_tokens),
        }
    )

    parsed = {}
    for attempt_idx in range(MAX_HTTP_ATTEMPTS):
        status_code, raw_text, network_error = perform_http_request(
            f"{base_url}/api/v2/completion?include_citations=false",
            {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": f"multipart/form-data; boundary={boundary}",
            },
            body,
        )
        detail = parse_error_detail(raw_text, network_error=network_error)
        if status_code != 200:
            if should_retry_request(status_code, detail) and attempt_idx < (MAX_HTTP_ATTEMPTS - 1):
                backoff_sleep(attempt_idx)
                continue
            raise RuntimeError(format_provider_error("Harvey", status_code, detail))

        try:
            parsed = json.loads(raw_text) if raw_text else {}
        except Exception:
            parsed = {}
        break

    text = ""
    if isinstance(parsed, dict):
        text = parsed.get("response") or parsed.get("text") or ""
    if not text:
        raise RuntimeError("Harvey response was missing message content.")

    return str(text).strip(), "harvey-assist"


def call_provider_prompt_json(prompt, api_key, provider, model_name=None, provider_base_url=None):
    provider_name = resolve_provider(provider, api_key)
    if provider_name == "managed":
        response_text, _ = call_managed_prompt(prompt, api_key, model_name, provider_base_url)
    elif provider_name == "localai":
        response_text, _ = call_localai_prompt(prompt, api_key, model_name, provider_base_url)
    elif provider_name == "openai":
        response_text, _ = call_openai_prompt(prompt, api_key, model_name)
    elif provider_name == "harvey":
        response_text, _ = call_harvey_prompt(prompt, api_key, 2048, provider_base_url)
    elif provider_name == "ollama":
        response_text, _ = call_ollama_prompt(prompt, api_key, model_name, provider_base_url)
    elif provider_name == "lmstudio":
        response_text, _ = call_lmstudio_prompt(prompt, api_key, model_name, provider_base_url)
    else:
        response_text, _ = call_anthropic_prompt(prompt, api_key, model_name)

    parsed = parse_json_response_text(response_text)
    if not isinstance(parsed, dict):
        raise ValueError("LLM response was not a JSON object.")
    return parsed


# Status categories for grouping (order determines display order)
STATUS_CATEGORIES = {
    'pending': {
        'title': 'PENDING DRAFTING',
        'patterns': [
            r'^pending',
            r'^to\s*be\s*drafted',
            r'^needs?\s*draft',
            r'^not\s*started',
            r'^tbd',
            r'^$',  # Empty status = pending
        ],
        'description': 'Documents that need initial drafts'
    },
    'review': {
        'title': 'UNDER REVIEW / WITH OPPOSING COUNSEL',
        'patterns': [
            r'with\s*(?:opposing\s*)?counsel',
            r'under\s*review',
            r'(?:awaiting|pending)\s*(?:comments?|review|feedback)',
            r'sent\s*(?:to|for)',
            r'circulated',
            r'draft\s*(?:circulated|sent)',
        ],
        'description': 'Documents out for review or with counterparty'
    },
    'signature': {
        'title': 'AWAITING SIGNATURE',
        'patterns': [
            r'execution\s*version',
            r'(?:ready|awaiting|pending)\s*(?:for\s*)?(?:signature|execution)',
            r'agreed\s*form',
            r'final\s*(?:form|version)',
            r'for\s*signature',
        ],
        'description': 'Documents ready for signature'
    },
    'executed': {
        'title': 'EXECUTED / COMPLETE',
        'patterns': [
            r'^executed',
            r'^signed',
            r'^complete[d]?',
            r'^done',
            r'fully\s*executed',
        ],
        'description': 'Completed documents (typically excluded from punchlist)'
    },
}

# Column name patterns
DOCUMENT_COLUMN_PATTERNS = [
    r'^document',
    r'^doc\s*name',
    r'^agreement',
    r'^instrument',
    r'^deliverable',
    r'^item',
    r'^description',
    r'^closing\s*document',
    r'^#',  # Sometimes just numbered
]

STATUS_COLUMN_PATTERNS = [
    r'^status',
    r'^state',
    r'^progress',
    r'^current\s*status',
]

PARTY_COLUMN_PATTERNS = [
    r'^part(?:y|ies)',
    r'^signator',
    r'^signer',
    r'^responsible',
    r'^who',
]

NOTES_COLUMN_PATTERNS = [
    r'^notes?',
    r'^comments?',
    r'^remarks?',
]


def find_column_index(headers, patterns):
    """Find column index matching any of the patterns."""
    for i, header in enumerate(headers):
        header_lower = header.lower().strip()
        for pattern in patterns:
            if re.match(pattern, header_lower):
                return i
    return -1


def normalize_cell_text(value):
    return re.sub(r'\s+', ' ', str(value or '')).strip()


def infer_document_column(headers, rows):
    if not headers:
        return -1

    max_cols = max(len(headers), max((len(r) for r in rows), default=0))
    best_col = -1
    best_score = -1

    for col_idx in range(max_cols):
        header_value = normalize_cell_text(headers[col_idx] if col_idx < len(headers) else '')
        header_lower = header_value.lower()
        if any(term in header_lower for term in ('status', 'state', 'progress', 'date')):
            continue

        non_empty = 0
        rich_text = 0
        for row in rows[:150]:
            cell = normalize_cell_text(row[col_idx] if col_idx < len(row) else '')
            if not cell:
                continue
            non_empty += 1
            if re.search(r'[a-zA-Z]', cell):
                rich_text += 1

        score = (non_empty * 2) + (rich_text * 3)
        if score > best_score:
            best_score = score
            best_col = col_idx

    return best_col


def score_header_candidate(all_rows, header_idx):
    headers = all_rows[header_idx]
    non_empty_headers = [normalize_cell_text(h) for h in headers if normalize_cell_text(h)]
    if len(non_empty_headers) < 2:
        return None

    doc_col = find_column_index(headers, DOCUMENT_COLUMN_PATTERNS)
    status_col = find_column_index(headers, STATUS_COLUMN_PATTERNS)
    party_col = find_column_index(headers, PARTY_COLUMN_PATTERNS)
    notes_col = find_column_index(headers, NOTES_COLUMN_PATTERNS)

    avg_header_len = (
        sum(len(h) for h in non_empty_headers) / len(non_empty_headers)
        if non_empty_headers else 0
    )
    looks_like_header = 1 if avg_header_len <= 45 else -5

    data_rows = all_rows[header_idx + 1: header_idx + 11]
    non_empty_data_rows = sum(1 for row in data_rows if any(normalize_cell_text(c) for c in row))

    score = len(non_empty_headers) + (non_empty_data_rows * 3) + looks_like_header
    if doc_col != -1:
        score += 35
    if status_col != -1:
        score += 15
    if party_col != -1:
        score += 6
    if notes_col != -1:
        score += 4

    return {
        'score': score,
        'header_idx': header_idx,
        'doc_col': doc_col,
        'status_col': status_col,
        'party_col': party_col,
        'notes_col': notes_col,
    }


def categorize_status(status_text):
    """
    Determine which category a status belongs to.

    Returns category key (pending, review, signature, executed) or 'pending' as default
    """
    if not status_text:
        return 'pending'

    status_lower = status_text.lower().strip()

    for category, config in STATUS_CATEGORIES.items():
        for pattern in config['patterns']:
            if re.search(pattern, status_lower):
                return category

    # Default to pending if no match
    return 'pending'


def categorize_items_with_llm(
    items,
    api_key,
    provider="anthropic",
    model_name=None,
    provider_base_url=None,
):
    """Use the selected LLM provider to categorize all checklist rows."""
    if provider_requires_api_key(provider) and not api_key:
        raise ValueError("API key is required for punchlist LLM categorization.")

    def normalize_response(payload):
        normalized = {"by_row": {}, "by_doc": {}}
        if not isinstance(payload, dict):
            return normalized

        # New format: {"categories":[...]}
        categories = payload.get("categories")
        if isinstance(categories, list):
            for item in categories:
                if not isinstance(item, dict):
                    continue
                category = str(item.get("category", "")).strip().lower()
                if category not in STATUS_CATEGORIES:
                    continue
                doc_name = str(item.get("document_name", "")).strip()
                row_id = item.get("row_id")
                try:
                    row_id = int(str(row_id).strip())
                except Exception:
                    row_id = None
                payload_item = {
                    "row_id": row_id,
                    "document_name": doc_name,
                    "category": category,
                    "confidence": item.get("confidence"),
                    "reasoning": str(item.get("reasoning", "")).strip(),
                }
                if row_id is not None:
                    normalized["by_row"][row_id] = payload_item
                if doc_name:
                    normalized["by_doc"][canonical_doc_key(doc_name)] = payload_item

        # Legacy format: {"Document Name":"pending"}
        for key, value in payload.items():
            if key == "categories":
                continue
            category = str(value).strip().lower() if isinstance(value, str) else ''
            if category not in STATUS_CATEGORIES:
                continue
            normalized["by_doc"][canonical_doc_key(key)] = {
                "row_id": None,
                "document_name": str(key),
                "category": category,
                "confidence": None,
                "reasoning": "",
            }

        return normalized

    # Prepare items for the prompt.
    items_for_prompt = [
        {
            "row_id": item.get('row_id'),
            "document_name": item.get('document_name', ''),
            "status": item.get('status', ''),
            "party": item.get('party', ''),
            "notes": item.get('notes', ''),
            "row_context": item.get('row_context', ''),
        }
        for item in items[:100]  # Limit to 100 items
    ]

    prompt = f"""You are categorizing legal transaction checklist rows for a daily punchlist.

For each row, classify into one of these categories:
- "pending": Needs drafting, not started, to be drafted, TBD
- "review": Under review, with counsel, sent to counterparty, awaiting comments, circulated
- "signature": Execution version, agreed form, ready for signature, final form
- "executed": Fully executed, signed, complete, done

Rows to categorize:
{json.dumps(items_for_prompt, indent=2)}

Guidance:
- Use document name + status + party + notes together; do not rely on exact header names.
- If information is ambiguous but appears open, prefer "pending" over "executed".
- If it is clearly signed/complete, use "executed".

Return JSON only in this exact shape:
{{
  "categories": [
    {{
      "row_id": 1,
      "document_name": "Credit Agreement",
      "category": "review",
      "confidence": 0.81,
      "reasoning": "Notes indicate it is with opposing counsel."
    }}
  ]
}}

Include EVERY row_id from the input exactly once.
ONLY return JSON, no extra text."""

    parsed = call_provider_prompt_json(
        prompt,
        api_key,
        provider,
        model_name=model_name,
        provider_base_url=provider_base_url,
    )
    return normalize_response(parsed)


def parse_checklist_for_punchlist(doc):
    """
    Parse Word document checklist table.

    Returns dict with columns and rows data
    """
    if not doc.tables:
        return None

    best_candidate = None

    for table in doc.tables:
        all_rows = []
        for row in table.rows:
            all_rows.append([normalize_cell_text(cell.text) for cell in row.cells])
        if not all_rows:
            continue

        max_probe = min(6, len(all_rows) - 1)
        for header_idx in range(max_probe + 1):
            candidate = score_header_candidate(all_rows, header_idx)
            if not candidate:
                continue
            candidate['table'] = table
            candidate['all_rows'] = all_rows
            if best_candidate is None or candidate['score'] > best_candidate['score']:
                best_candidate = candidate

    if not best_candidate:
        return None

    table = best_candidate['table']
    all_rows = best_candidate['all_rows']
    header_idx = best_candidate['header_idx']
    headers = all_rows[header_idx]
    rows = [row for row in all_rows[header_idx + 1:] if any(normalize_cell_text(c) for c in row)]

    doc_col = best_candidate['doc_col']
    if doc_col == -1:
        doc_col = infer_document_column(headers, rows)
    status_col = best_candidate['status_col']
    party_col = best_candidate['party_col']
    notes_col = best_candidate['notes_col']

    return {
        'headers': headers,
        'rows': rows,
        'table': table,
        'doc_col': doc_col,
        'status_col': status_col,
        'party_col': party_col,
        'notes_col': notes_col,
        'data_row_start_idx': header_idx + 1
    }


def extract_transaction_name(checklist_path, doc):
    """
    Try to extract transaction name from document.

    Looks in:
    1. Document title/heading
    2. Filename
    """
    # Try to find title in document
    for para in doc.paragraphs[:5]:  # Check first 5 paragraphs
        text = para.text.strip()
        if text and len(text) > 5 and len(text) < 100:
            # Skip if it looks like a column header
            if not any(h in text.lower() for h in ['document', 'status', 'party', 'item']):
                return text

    # Fall back to filename
    base_name = os.path.splitext(os.path.basename(checklist_path))[0]
    # Clean up common suffixes
    for suffix in ['_checklist', '_closing', '_documents', 'checklist', 'closing']:
        base_name = re.sub(f'{suffix}$', '', base_name, flags=re.IGNORECASE)

    return base_name.replace('_', ' ').replace('-', ' ').strip()


def generate_punchlist(
    checklist_path,
    output_folder,
    status_filters=None,
    api_key=None,
    provider="anthropic",
    model_name=None,
    provider_base_url=None,
):
    """
    Generate punchlist document from checklist.

    Args:
        checklist_path: Path to Word document with checklist
        output_folder: Folder to save punchlist
        status_filters: List of status categories to include ['pending', 'review', 'signature']
                       If None, includes all except 'executed'
        api_key: Provider API key for LLM-based categorization
        provider: LLM provider (anthropic/openai/harvey)
        model_name: Optional provider model override
        provider_base_url: Optional provider API base URL

    Returns:
        dict with: success, output_path, item_count, categories
    """
    result = {
        'success': False,
        'output_path': None,
        'item_count': 0,
        'categories': {},
        'error': None
    }

    # Default: include all except executed
    if status_filters is None:
        status_filters = ['pending', 'review', 'signature']

    try:
        api_key = normalize_api_key(api_key)
        provider = resolve_provider(provider, api_key)

        if provider_requires_api_key(provider) and not api_key:
            result['error'] = 'LLM API key is required to generate punchlist.'
            return result

        # Open checklist
        doc = Document(checklist_path)

        # Parse table
        parsed = parse_checklist_for_punchlist(doc)
        if not parsed:
            result['error'] = 'No table found in checklist document'
            return result

        # Get transaction name
        transaction_name = extract_transaction_name(checklist_path, doc)

        # Collect all items first
        all_items = []
        for row_idx, row in enumerate(parsed['rows']):
            # Get document name
            doc_name = ''
            if parsed['doc_col'] >= 0 and parsed['doc_col'] < len(row):
                doc_name = normalize_cell_text(row[parsed['doc_col']])

            if not doc_name:
                continue

            # Get status
            status = ''
            if parsed['status_col'] >= 0 and parsed['status_col'] < len(row):
                status = normalize_cell_text(row[parsed['status_col']])

            # Get party/responsible
            party = ''
            if parsed['party_col'] >= 0 and parsed['party_col'] < len(row):
                party = normalize_cell_text(row[parsed['party_col']])

            # Get notes
            notes = ''
            if parsed['notes_col'] >= 0 and parsed['notes_col'] < len(row):
                notes = normalize_cell_text(row[parsed['notes_col']])

            # Additional context from non-primary columns.
            context_parts = []
            for col_idx, cell_value in enumerate(row):
                if col_idx in (parsed['doc_col'], parsed['status_col'], parsed['party_col'], parsed['notes_col']):
                    continue
                cleaned = normalize_cell_text(cell_value)
                if not cleaned:
                    continue
                header_label = normalize_cell_text(parsed['headers'][col_idx] if col_idx < len(parsed['headers']) else f"Column {col_idx + 1}")
                if header_label:
                    context_parts.append(f"{header_label}: {cleaned}")
                else:
                    context_parts.append(cleaned)

            all_items.append({
                'row_id': parsed.get('data_row_start_idx', 1) + row_idx,
                'document_name': doc_name,
                'status': status,
                'party': party,
                'notes': notes,
                'row_context': ' | '.join(context_parts[:8]),
            })

        if not all_items:
            result['error'] = 'No checklist items found to categorize.'
            return result

        # LLM categorization is mandatory for this workflow.
        try:
            llm_categories = categorize_items_with_llm(
                all_items,
                api_key,
                provider=provider,
                model_name=model_name,
                provider_base_url=provider_base_url,
            )
        except Exception as e:
            result['error'] = f'LLM punchlist categorization failed: {e}'
            return result

        # Categorize all items
        categorized_items = {cat: [] for cat in STATUS_CATEGORIES.keys()}
        llm_row_matches = llm_categories.get('by_row', {}) if isinstance(llm_categories, dict) else {}
        llm_doc_matches = llm_categories.get('by_doc', {}) if isinstance(llm_categories, dict) else {}
        llm_assigned_count = 0
        fallback_assigned_count = 0

        for item in all_items:
            doc_name = item['document_name']
            row_id = item.get('row_id')
            llm_match = llm_row_matches.get(row_id) if isinstance(llm_row_matches, dict) else None
            if not llm_match and isinstance(llm_doc_matches, dict):
                llm_match = llm_doc_matches.get(canonical_doc_key(doc_name))

            category = ''
            if isinstance(llm_match, dict):
                category = str(llm_match.get('category', '')).strip().lower()
            elif isinstance(llm_match, str):
                category = llm_match.strip().lower()

            if category in STATUS_CATEGORIES:
                llm_assigned_count += 1
            else:
                # If the LLM misses a row, preserve continuity with a deterministic fallback.
                category = categorize_status(" ".join([
                    item.get('status', ''),
                    item.get('notes', ''),
                    item.get('row_context', ''),
                ]))
                fallback_assigned_count += 1

            categorized_items[category].append({
                'document': doc_name,
                'status': item['status'],
                'party': item['party'],
                'notes': item['notes']
            })

        # Create punchlist document
        punchlist_doc = Document()

        # Set narrow margins
        for section in punchlist_doc.sections:
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)

        # Title
        title = punchlist_doc.add_paragraph()
        title_run = title.add_run(f'DAILY PUNCHLIST')
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Transaction name
        if transaction_name:
            txn_para = punchlist_doc.add_paragraph()
            txn_run = txn_para.add_run(transaction_name)
            txn_run.font.size = Pt(14)
            txn_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        date_para = punchlist_doc.add_paragraph()
        date_run = date_para.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}')
        date_run.font.size = Pt(11)
        date_run.font.color.rgb = RGBColor(100, 100, 100)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        punchlist_doc.add_paragraph()  # Spacer

        # Count total open items
        total_open = sum(
            len(items) for cat, items in categorized_items.items()
            if cat in status_filters
        )

        # Summary line
        summary = punchlist_doc.add_paragraph()
        summary_run = summary.add_run(f'OPEN ITEMS: {total_open}')
        summary_run.bold = True
        summary_run.font.size = Pt(12)

        punchlist_doc.add_paragraph()  # Spacer

        # Add each category
        category_counts = {}
        for category in ['pending', 'review', 'signature', 'executed']:
            if category not in status_filters:
                continue

            items = categorized_items[category]
            if not items:
                continue

            config = STATUS_CATEGORIES[category]
            category_counts[category] = len(items)

            # Category header
            cat_header = punchlist_doc.add_paragraph()
            cat_run = cat_header.add_run(f'{config["title"]} ({len(items)})')
            cat_run.bold = True
            cat_run.font.size = Pt(11)

            # Items
            for item in items:
                item_para = punchlist_doc.add_paragraph()
                item_para.paragraph_format.left_indent = Inches(0.25)

                # Checkbox character
                checkbox = item_para.add_run('☐ ')
                checkbox.font.size = Pt(11)

                # Document name
                doc_run = item_para.add_run(item['document'])
                doc_run.font.size = Pt(11)

                # Status/notes in lighter text if available
                extra_info = []
                if item['status'] and item['status'].lower() not in ['pending', 'tbd', '']:
                    extra_info.append(item['status'])
                if item['party']:
                    extra_info.append(f"({item['party']})")

                if extra_info:
                    info_run = item_para.add_run(f" - {', '.join(extra_info)}")
                    info_run.font.size = Pt(10)
                    info_run.font.color.rgb = RGBColor(100, 100, 100)

            punchlist_doc.add_paragraph()  # Spacer between categories

        # Footer
        footer = punchlist_doc.add_paragraph()
        footer_run = footer.add_run(f'Generated by EmmaNeigh on {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(150, 150, 150)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save punchlist
        os.makedirs(output_folder, exist_ok=True)

        timestamp = datetime.now().strftime('%Y%m%d')
        output_filename = f'Punchlist_{timestamp}.docx'
        output_path = os.path.join(output_folder, output_filename)

        punchlist_doc.save(output_path)

        result['success'] = True
        result['output_path'] = output_path
        result['item_count'] = total_open
        result['categories'] = category_counts
        result['llm_assigned_count'] = llm_assigned_count
        result['fallback_assigned_count'] = fallback_assigned_count

        return result

    except Exception as e:
        result['error'] = str(e)
        return result


def main():
    if len(sys.argv) < 3:
        print(json.dumps({
            'success': False,
            'error': 'Usage: punchlist_generator.py <checklist_path> <output_folder> [status_filters_json] <api_key> [provider] [model] [provider_base_url]'
        }))
        sys.exit(1)

    checklist_path = sys.argv[1]
    output_folder = sys.argv[2]

    # Parse status filters if provided
    status_filters = None
    if len(sys.argv) > 3:
        try:
            status_filters = json.loads(sys.argv[3])
        except json.JSONDecodeError:
            pass

    # API key is required for LLM-driven punchlist categorization.
    api_key = sys.argv[4] if len(sys.argv) > 4 else (
        os.environ.get('ANTHROPIC_API_KEY')
        or os.environ.get('OPENAI_API_KEY')
        or os.environ.get('HARVEY_API_KEY')
    )
    provider = resolve_provider(
        sys.argv[5] if len(sys.argv) > 5 else os.environ.get('AI_PROVIDER', 'anthropic'),
        api_key,
    )
    model_name = sys.argv[6] if len(sys.argv) > 6 else (
        os.environ.get('OPENAI_MODEL') if provider == 'openai'
        else os.environ.get('LOCALAI_MODEL') if provider == 'localai'
        else os.environ.get('OLLAMA_MODEL') if provider == 'ollama'
        else os.environ.get('LMSTUDIO_MODEL') if provider == 'lmstudio'
        else os.environ.get('CLAUDE_MODEL')
    )
    provider_base_url = sys.argv[7] if len(sys.argv) > 7 else (
        os.environ.get('HARVEY_BASE_URL', DEFAULT_HARVEY_BASE_URL) if provider == 'harvey'
        else os.environ.get('LOCALAI_BASE_URL', DEFAULT_LOCALAI_BASE_URL) if provider == 'localai'
        else os.environ.get('OLLAMA_BASE_URL', DEFAULT_OLLAMA_BASE_URL) if provider == 'ollama'
        else os.environ.get('LMSTUDIO_BASE_URL', DEFAULT_LMSTUDIO_BASE_URL) if provider == 'lmstudio'
        else None
    )

    result = generate_punchlist(
        checklist_path,
        output_folder,
        status_filters,
        api_key,
        provider=provider,
        model_name=model_name,
        provider_base_url=provider_base_url,
    )
    print(json.dumps(result))


if __name__ == '__main__':
    main()
