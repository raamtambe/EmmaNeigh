#!/usr/bin/env python3
"""
EmmaNeigh - Signature Packet Processor
v5.1.6:
- Fixed folder processing (now properly handles folder input vs file list)
- Fixed footer detection (only captures text in actual footer region near bottom of page)
- Added Document ID extraction from bottom-left of footer
"""

import fitz
import os
import pandas as pd
import re
import sys
import json
import shutil

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None
    Pt = None
    Inches = None
    WD_ALIGN_PARAGRAPH = None


# Column header patterns for signature table detection
NAME_HEADERS = ["NAME", "PRINTED NAME", "SIGNATORY", "SIGNER", "PRINT NAME", "TITLE"]
SIGNATURE_HEADERS = ["SIGNATURE", "SIGN", "BY", "SIGN HERE"]

# Trigger phrases that indicate signature pages without BY:
SIGNATURE_TRIGGER_PHRASES = [
    "AGREED", "ACKNOWLEDGED", "UNDERSIGNED", "WITNESS", "ATTEST",
    "IN WITNESS WHEREOF", "EXECUTED", "CERTIFIED", "AUTHORIZED",
    "THE PARTIES HERETO", "DULY AUTHORIZED"
]

ENTITY_TERMS = ["LLC", "INC", "CORP", "CORPORATION", "LP", "LLP", "TRUST", "COMPANY", "LTD", "LIMITED", "HOLDINGS", "GROUP", "CO."]
SIGNATURE_LABEL_PREFIXES = ["NAME", "PRINTED NAME", "PRINT NAME", "SIGNATORY", "SIGNER"]
IGNORE_SIGNER_VALUES = {
    "BY", "NAME", "PRINTED NAME", "PRINT NAME", "SIGNATORY", "SIGNER",
    "TITLE", "DATE", "SIGNATURE", "ITS", "WITNESS", "ATTEST", "EXECUTED"
}
UNASSIGNED_SIGNER_BUCKET = "UNASSIGNED SIGNATURE PAGES - REVIEW REQUIRED"

ANNEX_KEYWORDS = ["SCHEDULE", "EXHIBIT", "ANNEX", "APPENDIX"]
LONG_ANNEX_PAGE_THRESHOLD = 100


def ensure_docx_support():
    if Document is None:
        raise ImportError("python-docx is required for Word signature packet support")


def emit(msg_type, **kwargs):
    """Output JSON message to stdout for the Electron app."""
    print(json.dumps({"type": msg_type, **kwargs}), flush=True)


def normalize_name(name):
    """Normalize signer name: uppercase, remove punctuation, collapse spaces."""
    if not name:
        return ""
    name = name.upper()
    name = re.sub(r"[.,]", "", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name


def is_probable_person(name):
    """Check if name is likely a person (not an entity)."""
    if not name:
        return False
    name_upper = name.upper()
    if any(term in name_upper for term in ENTITY_TERMS):
        return False
    # Check for reasonable word count (2-4 words typical for person names)
    word_count = len(name.split())
    if word_count < 1 or word_count > 5:
        return False
    # Check that it's not just numbers or special characters
    if not re.search(r'[A-Za-z]{2,}', name):
        return False
    return True


def is_probable_entity_signer(name):
    """Check if text looks like an entity-style signatory fallback."""
    if not name:
        return False
    cleaned = re.sub(r"\s+", " ", str(name).strip())
    upper = cleaned.upper()
    if upper in IGNORE_SIGNER_VALUES:
        return False
    if len(cleaned.split()) > 10:
        return False
    if any(term in upper for term in ENTITY_TERMS):
        return True
    words = [word for word in cleaned.split() if re.search(r"[A-Z]", word)]
    if len(words) >= 2 and cleaned == cleaned.upper():
        return True
    return False


def clean_signer_candidate(raw_text):
    """Normalize signer text extracted from signature blocks."""
    text = str(raw_text or "").strip()
    if not text:
        return ""
    text = re.sub(r'(?i)^/s/\s*', '', text)
    text = re.sub(r'(?i)^(?:by|name|printed\s+name|print\s+name|signatory|signer|title)\s*:?', '', text)
    text = re.sub(r'\[[^\]]*\]', ' ', text)
    text = re.sub(r'_{2,}', ' ', text)
    text = re.sub(r'\.{2,}', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip(" :;/,-")
    if text.upper() in IGNORE_SIGNER_VALUES:
        return ""
    return text


def normalize_signer_candidate(raw_text):
    cleaned = clean_signer_candidate(raw_text)
    return normalize_name(cleaned) if cleaned else ""


# ========== FOOTER EXTRACTION ==========

def extract_footer_from_pdf_page(page):
    """
    Extract footer text from the bottom of a PDF page.
    v5.1.6: Only captures text that is actually in the footer region (bottom 10% of page).
    Returns "N/A" if no footer text is found.
    """
    page_height = page.rect.height
    page_width = page.rect.width

    # Footer region: bottom 10% of page (typical footer area)
    footer_threshold = page_height * 0.90  # Text must be below this Y coordinate

    # Get text blocks with position information
    blocks = page.get_text("dict")["blocks"]

    footer_texts = []

    for block in blocks:
        if "lines" not in block:
            continue

        for line in block["lines"]:
            for span in line["spans"]:
                # Check if text is in footer region (below 90% of page height)
                y_pos = span["bbox"][1]  # Top Y position of text
                if y_pos >= footer_threshold:
                    text = span["text"].strip()
                    if text and len(text) > 1:
                        # Skip if it's just a page number
                        if not re.match(r'^[\d\s\-\.]+$', text):
                            footer_texts.append({
                                "text": text,
                                "x": span["bbox"][0],
                                "y": y_pos
                            })

    if not footer_texts:
        return "N/A"

    # Sort by Y position (top to bottom), then combine texts on same line
    footer_texts.sort(key=lambda t: (t["y"], t["x"]))

    # Combine texts that are on approximately the same Y position
    combined_lines = []
    current_line = []
    current_y = None

    for ft in footer_texts:
        if current_y is None or abs(ft["y"] - current_y) < 5:  # Same line if Y within 5 pts
            current_line.append(ft["text"])
            current_y = ft["y"]
        else:
            if current_line:
                combined_lines.append(" ".join(current_line))
            current_line = [ft["text"]]
            current_y = ft["y"]

    if current_line:
        combined_lines.append(" ".join(current_line))

    # Priority 1: Look for "SIGNATURE PAGE TO X" pattern
    for line in combined_lines:
        if 'SIGNATURE PAGE' in line.upper():
            return line.strip()

    # Return the first meaningful footer line
    for line in combined_lines:
        if len(line) > 3 and not re.match(r'^[\d\s\-\.]+$', line):
            return line.strip()

    return "N/A"


def extract_document_id_from_pdf_page(page):
    """
    Extract document ID from bottom-left of footer area.
    v5.1.6: Document IDs typically appear in bottom-left corner.
    Returns "N/A" if no document ID found.
    """
    page_height = page.rect.height
    page_width = page.rect.width

    # Footer region: bottom 10% of page
    footer_threshold = page_height * 0.90
    # Left region: left 40% of page
    left_threshold = page_width * 0.40

    # Get text blocks with position information
    blocks = page.get_text("dict")["blocks"]

    bottom_left_texts = []

    for block in blocks:
        if "lines" not in block:
            continue

        for line in block["lines"]:
            for span in line["spans"]:
                x_pos = span["bbox"][0]  # Left X position
                y_pos = span["bbox"][1]  # Top Y position

                # Check if text is in bottom-left region
                if y_pos >= footer_threshold and x_pos <= left_threshold:
                    text = span["text"].strip()
                    if text:
                        bottom_left_texts.append({
                            "text": text,
                            "x": x_pos,
                            "y": y_pos
                        })

    if not bottom_left_texts:
        return "N/A"

    # Sort by Y (bottom to top) then X (left to right)
    bottom_left_texts.sort(key=lambda t: (t["y"], t["x"]))

    # Look for document ID patterns:
    # - Alphanumeric codes (e.g., "DOC-12345", "123456789", "ID: ABC-123")
    # - Often starts with numbers or has specific prefixes

    for item in bottom_left_texts:
        text = item["text"]

        # Pattern 1: ID with prefix like "Doc ID:", "ID:", "Ref:", etc.
        id_match = re.search(r'(?:ID|Doc|Ref|No\.?|#)\s*[:\-]?\s*([A-Z0-9\-\.]+)', text, re.IGNORECASE)
        if id_match:
            return id_match.group(1).strip()

        # Pattern 2: Standalone alphanumeric code (at least 4 chars, mix of letters/numbers)
        if re.match(r'^[A-Z0-9\-\.]{4,}$', text, re.IGNORECASE):
            # Avoid capturing simple page numbers
            if not re.match(r'^\d{1,3}$', text):
                return text

        # Pattern 3: Number sequence that looks like an ID (5+ digits)
        if re.match(r'^\d{5,}$', text):
            return text

    # If no clear ID pattern, return first bottom-left text as potential ID
    if bottom_left_texts:
        first_text = bottom_left_texts[0]["text"]
        # Only return if it looks like an identifier (short, no spaces in middle)
        if len(first_text) <= 30 and not re.match(r'^Page\s', first_text, re.IGNORECASE):
            return first_text

    return "N/A"


def extract_footer_from_docx(docx_path):
    """
    Extract footer from DOCX document sections.
    v5.1.6: Only returns actual footer content, not last paragraph.
    Returns "N/A" if no footer found.
    """
    try:
        ensure_docx_support()
        doc = Document(docx_path)

        # Try to get explicit footer from sections
        for section in doc.sections:
            try:
                footer = section.footer
                if footer and footer.paragraphs:
                    footer_text = ' '.join(
                        p.text.strip() for p in footer.paragraphs if p.text.strip()
                    )
                    if footer_text:
                        return footer_text
            except Exception:
                pass

        # No explicit footer found - return N/A (don't fall back to last paragraph)
        return "N/A"
    except Exception:
        pass

    return "N/A"


def extract_document_id_from_docx(docx_path):
    """
    Extract document ID from DOCX footer (bottom-left).
    v5.1.6: Looks for ID patterns in footer sections.
    Returns "N/A" if no document ID found.
    """
    try:
        ensure_docx_support()
        doc = Document(docx_path)

        # Check footer sections for document IDs
        for section in doc.sections:
            try:
                footer = section.footer
                if footer and footer.paragraphs:
                    for para in footer.paragraphs:
                        text = para.text.strip()
                        if not text:
                            continue

                        # Pattern 1: ID with prefix
                        id_match = re.search(r'(?:ID|Doc|Ref|No\.?|#)\s*[:\-]?\s*([A-Z0-9\-\.]+)', text, re.IGNORECASE)
                        if id_match:
                            return id_match.group(1).strip()

                        # Pattern 2: Standalone alphanumeric code
                        if re.match(r'^[A-Z0-9\-\.]{4,}$', text, re.IGNORECASE):
                            if not re.match(r'^\d{1,3}$', text):
                                return text

                        # Pattern 3: Number sequence (5+ digits)
                        if re.match(r'^\d{5,}$', text):
                            return text
            except Exception:
                pass

        return "N/A"
    except Exception:
        pass

    return "N/A"


# ========== EXTENDED SIGNATURE DETECTION ==========

def extract_signers_underscore_name(text):
    """
    Pattern: Name followed by underscore line (resolution style)
    Example: John Smith ________________________
    """
    signers = set()

    # Pattern: Name followed by 4+ underscores
    pattern = r'([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){0,3})\s*_{4,}'
    matches = re.findall(pattern, text)

    for name in matches:
        name = name.strip()
        if is_probable_person(name):
            signers.add(normalize_name(name))

    return signers


def extract_signers_underscore_label(text):
    """
    Pattern: Underscore line followed by Name:/Title: label
    Example:
    _______________________________
    Name: John Smith
    Title: President
    """
    signers = set()
    lines = text.splitlines()

    for i, line in enumerate(lines):
        # Look for line with 10+ underscores
        if re.search(r'_{10,}', line):
            # Check next 3 lines for Name: label
            for j in range(1, 4):
                if i + j < len(lines):
                    next_line = lines[i + j]
                    name_match = re.search(r'Name:\s*(.+)', next_line, re.IGNORECASE)
                    if name_match:
                        name = name_match.group(1).strip()
                        # Clean up the name (remove trailing underscores, etc.)
                        name = re.sub(r'_{2,}.*$', '', name).strip()
                        if name and is_probable_person(name):
                            signers.add(normalize_name(name))
                        break

    return signers


def extract_signers_trigger_phrase(text):
    """
    Pattern: Trigger phrases followed by names in subsequent lines
    Example:
    THE UNDERSIGNED HEREBY AGREE:

    Person A
    Person B
    Person C
    """
    signers = set()
    text_upper = text.upper()

    for phrase in SIGNATURE_TRIGGER_PHRASES:
        if phrase in text_upper:
            # Find where the phrase occurs
            phrase_idx = text_upper.find(phrase)
            subsequent_text = text[phrase_idx:]
            lines = subsequent_text.split('\n')[1:15]  # Check next 15 lines

            for line in lines:
                line = line.strip()
                # Skip empty lines, short lines, and lines that are just underscores
                if not line or len(line) < 3 or re.match(r'^[_\-\s]+$', line):
                    continue
                # Skip lines that look like instructions
                if any(word in line.upper() for word in ['PLEASE', 'SIGN', 'DATE', 'PRINT', 'BELOW']):
                    continue

                # Check if line looks like a name (2-4 words, starts with capital)
                words = line.split()
                if 1 <= len(words) <= 5:
                    # Remove trailing underscores or colons
                    candidate = normalize_signer_candidate(re.sub(r'[_:]+$', '', ' '.join(words)).strip())
                    if is_probable_person(candidate):
                        signers.add(candidate)

    return signers


def extract_signers_horizontal_table(table_data):
    """
    Pattern: Horizontal signature tables (common in incumbency certs)
    Example:
    | Name        | Title       | Signature    |
    |-------------|-------------|--------------|
    | John Smith  | CEO         | ____________ |
    """
    if not table_data or len(table_data) < 2:
        return set()

    headers = table_data[0]
    headers_upper = [(h.upper().strip() if h else "") for h in headers]

    # Check if this looks like a signature/incumbency table
    has_name = any(any(nh in h for nh in NAME_HEADERS) for h in headers_upper)
    has_title = any('TITLE' in h for h in headers_upper)
    has_sig_or_empty = any(
        any(sh in h for sh in SIGNATURE_HEADERS) or h == "" or '___' in h
        for h in headers_upper
    )

    if not (has_name or has_title):
        return set()

    signers = set()
    name_col_idx = find_name_column(headers)

    # If no explicit name column, try first column
    if name_col_idx is None:
        name_col_idx = 0

    for row in table_data[1:]:
        if name_col_idx < len(row) and row[name_col_idx]:
            cell_text = row[name_col_idx]
            if isinstance(cell_text, str):
                # Handle multi-line cells
                for line in cell_text.split('\n'):
                    name = normalize_name(line.strip())
                    if name and is_probable_person(name):
                        signers.add(name)

    return signers


def extract_signers_name_title_pattern(text):
    """
    v5.1.4: New pattern - NAME: followed by TITLE: on subsequent lines.
    This is a common signature block format in legal documents.

    Example:
    NAME: John Smith
    TITLE: President

    Returns set of normalized signer names.
    """
    signers = set()
    lines = text.splitlines()

    for i, line in enumerate(lines):
        # Look for NAME: label
        name_match = re.search(r'\bNAME\s*:\s*(.+)', line, re.IGNORECASE)
        if name_match:
            # Check if TITLE: appears within next 3 lines (confirms this is a signature block)
            has_title_nearby = False
            for j in range(1, 4):
                if i + j < len(lines):
                    if re.search(r'\bTITLE\s*:', lines[i + j], re.IGNORECASE):
                        has_title_nearby = True
                        break

            if has_title_nearby:
                name = normalize_signer_candidate(name_match.group(1))
                if name and is_probable_person(name):
                    signers.add(name)

    return signers


def extract_signers_from_labeled_lines(text):
    """Extract signer labels from explicit Name/Printed Name/Signatory lines."""
    signers = set()
    entity_fallbacks = set()
    lines = text.splitlines()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        match = re.search(r'\b(?:NAME|PRINTED NAME|PRINT NAME|SIGNATORY|SIGNER)\s*:\s*(.+)', stripped, re.IGNORECASE)
        if not match:
            continue
        candidate = normalize_signer_candidate(match.group(1))
        if not candidate:
            continue
        if is_probable_person(candidate):
            signers.add(candidate)
        elif is_probable_entity_signer(candidate):
            entity_fallbacks.add(candidate)

    return signers, entity_fallbacks


def extract_entities_near_by_blocks(lines, by_index):
    """Look above a BY block for an entity signer fallback."""
    entities = set()
    for offset in range(1, 4):
        prior_index = by_index - offset
        if prior_index < 0:
            break
        candidate = normalize_signer_candidate(lines[prior_index])
        if not candidate:
            continue
        if candidate in {"AND", "ITS", "BY"}:
            continue
        if is_probable_entity_signer(candidate):
            entities.add(candidate)
            break
        # Stop once we hit unrelated prose instead of climbing too far.
        if len(candidate.split()) > 6:
            break
    return entities


def collect_signature_cues(text, tables=None):
    """Score how much a page looks like a signature page even if no signer was extracted."""
    hits = []
    upper = str(text or "").upper()
    if re.search(r'\bBY\s*:', upper):
        hits.append("BY")
    if re.search(r'\b(?:NAME|PRINTED NAME|PRINT NAME|SIGNATORY|SIGNER)\s*:', upper):
        hits.append("NAME_LABEL")
    if re.search(r'\bTITLE\s*:', upper):
        hits.append("TITLE_LABEL")
    if re.search(r'\bDATE\s*:', upper):
        hits.append("DATE_LABEL")
    if re.search(r'_{6,}', text or ""):
        hits.append("UNDERSCORE")
    if any(phrase in upper for phrase in SIGNATURE_TRIGGER_PHRASES):
        hits.append("TRIGGER_PHRASE")
    if tables:
        for table_data in tables:
            if table_data and len(table_data) > 0 and is_signature_table(table_data[0]):
                hits.append("SIGNATURE_TABLE")
                break
    score = 0
    for hit in hits:
        if hit in {"BY", "NAME_LABEL", "SIGNATURE_TABLE"}:
            score += 3
        elif hit in {"TITLE_LABEL", "DATE_LABEL", "TRIGGER_PHRASE"}:
            score += 2
        else:
            score += 1
    return score, hits


def analyze_signature_page_text(text, tables=None):
    """Second-pass signature-page analysis with named-signer extraction and review flags."""
    all_person_signers = set()
    entity_fallbacks = set()
    methods_used = []

    by_signers = extract_signers_from_by_blocks(text)
    if by_signers:
        all_person_signers.update(by_signers)
        methods_used.append("BY_BLOCK")

    labeled_signers, labeled_entities = extract_signers_from_labeled_lines(text)
    if labeled_signers:
        all_person_signers.update(labeled_signers)
        methods_used.append("LABEL")
    if labeled_entities:
        entity_fallbacks.update(labeled_entities)
        methods_used.append("LABEL_ENTITY")

    name_title_signers = extract_signers_name_title_pattern(text)
    if name_title_signers:
        all_person_signers.update(name_title_signers)
        methods_used.append("NAME_TITLE")

    trigger_signers = extract_signers_trigger_phrase(text)
    if trigger_signers:
        all_person_signers.update(trigger_signers)
        methods_used.append("TRIGGER")

    if tables:
        for table_data in tables:
            if table_data and len(table_data) > 0 and is_signature_table(table_data[0]):
                table_signers = extract_signers_from_table(table_data)
                if table_signers:
                    all_person_signers.update(table_signers)
                    methods_used.append("TABLE")

    lines = [line.strip() for line in str(text or "").splitlines() if line.strip()]
    if not all_person_signers:
        for i, line in enumerate(lines):
            if re.search(r'\bBY\s*:', line, re.IGNORECASE):
                entity_fallbacks.update(extract_entities_near_by_blocks(lines, i))
        if entity_fallbacks:
            methods_used.append("ENTITY_FALLBACK")

    cue_score, cue_hits = collect_signature_cues(text, tables)
    detected_signers = set(all_person_signers) if all_person_signers else set(entity_fallbacks)
    needs_review = False

    if detected_signers and not all_person_signers:
        needs_review = True

    is_signature_page = bool(detected_signers)
    if not is_signature_page and cue_score >= 5:
        is_signature_page = True
        needs_review = True
        methods_used.append("SIGNATURE_CUE_REVIEW")

    return {
        "is_signature_page": is_signature_page,
        "signers": detected_signers,
        "method": ",".join(dict.fromkeys(methods_used)),
        "cue_score": cue_score,
        "cue_hits": cue_hits,
        "needs_review": needs_review
    }


def detect_signature_page_extended(text, tables=None):
    analysis = analyze_signature_page_text(text, tables)
    return (analysis["is_signature_page"], analysis["signers"], analysis["method"] or None)


def extract_signers_from_by_blocks(text):
    """Extract signer names from traditional BY:/Name: blocks."""
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    signers = set()

    for i, line in enumerate(lines):
        if re.search(r'\bBY\s*:', line, re.IGNORECASE):
            inline_match = re.search(r'\bBY\s*:\s*(.+)$', line, re.IGNORECASE)
            if inline_match:
                inline_candidate = normalize_signer_candidate(inline_match.group(1))
                if inline_candidate and is_probable_person(inline_candidate):
                    signers.add(inline_candidate)
                    continue

            # Tier 1: Prefer explicit Name: field
            for j in range(1, 7):
                if i + j >= len(lines):
                    break
                cand = lines[i + j]
                name_match = re.search(r'\b(?:NAME|PRINTED NAME|PRINT NAME|SIGNATORY|SIGNER)\s*:\s*(.+)', cand, re.IGNORECASE)
                if name_match:
                    candidate = normalize_signer_candidate(name_match.group(1))
                    if candidate and is_probable_person(candidate):
                        signers.add(candidate)
                    break
            else:
                # Tier 2: Look for probable person name nearby
                for j in range(1, 7):
                    if i + j >= len(lines):
                        break
                    cand = normalize_signer_candidate(lines[i + j])
                    if is_probable_person(cand):
                        signers.add(cand)
                        break
    return signers


def find_name_column(headers):
    """Find the index of the name column in table headers."""
    if not headers:
        return None
    headers_upper = [(h.upper().strip() if h else "") for h in headers]
    for i, h in enumerate(headers_upper):
        for name_header in NAME_HEADERS:
            if name_header in h:
                return i
    return None


def is_signature_table(headers):
    """
    Check if table headers indicate a signature table.
    v5.1.3: Removed empty column acceptance to prevent false positives.
    """
    if not headers:
        return False

    headers_upper = [(h.upper().strip() if h else "") for h in headers]

    # Must have a name-like column
    has_name = any(
        any(nh in h for nh in NAME_HEADERS)
        for h in headers_upper
    )

    # v5.1.3: STRICTER - Must have EXPLICIT signature-like column header
    # Removed: "or h == ''" which accepted any table with empty columns
    has_sig = any(
        any(sh in h for sh in SIGNATURE_HEADERS)
        for h in headers_upper
        if h  # Skip empty headers
    )

    return has_name and has_sig


def extract_signers_from_table(table_data):
    """Extract signer names from a signature table."""
    if not table_data or len(table_data) < 2:
        return set()

    headers = table_data[0]
    name_col_idx = find_name_column(headers)

    if name_col_idx is None:
        return set()

    signers = set()
    for row in table_data[1:]:  # Skip header row
        if name_col_idx < len(row) and row[name_col_idx]:
            # Handle multi-line cell content
            cell_text = row[name_col_idx]
            if isinstance(cell_text, str):
                cell_text = " ".join(cell_text.split())  # Normalize whitespace
            name = normalize_signer_candidate(cell_text)
            if name and (is_probable_person(name) or is_probable_entity_signer(name)):
                signers.add(name)

    return signers


def extract_person_signers(page):
    """
    Extract signer names from PDF page using extended detection.
    Returns a tuple: (signers: set, detection_method: str)
    """
    analysis = analyze_pdf_signature_page(page)
    return analysis["signers"], analysis["method"] if analysis["method"] else ""


def analyze_pdf_signature_page(page):
    text = page.get_text()

    tables_data = []
    try:
        tables = page.find_tables()
        for table in tables.tables:
            data = table.extract()
            if data:
                tables_data.append(data)
    except Exception:
        pass

    return analyze_signature_page_text(text, tables_data)


# ========== DOCX SUPPORT ==========

def extract_text_from_docx(docx_path):
    """Extract all text from DOCX document (headers, body)."""
    ensure_docx_support()
    doc = Document(docx_path)
    text_parts = []

    # Get header text
    for section in doc.sections:
        try:
            for para in section.header.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
        except Exception:
            pass

    # Get body text
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    return '\n'.join(text_parts)


def extract_tables_from_docx(docx_path):
    """Extract all tables from DOCX as list of rows."""
    ensure_docx_support()
    doc = Document(docx_path)
    tables_list = []

    for table in doc.tables:
        rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(row_data)
        if rows:
            tables_list.append(rows)

    return tables_list


def extract_signers_from_docx(docx_path):
    """
    Extract signers from DOCX file using extended detection.
    Returns a tuple: (signers: set, detection_method: str)
    """
    analysis = analyze_docx_signature_page(docx_path)
    return analysis["signers"], analysis["method"] if analysis["method"] else ""


def analyze_docx_signature_page(docx_path):
    text = extract_text_from_docx(docx_path)
    tables_data = extract_tables_from_docx(docx_path)
    return analyze_signature_page_text(text, tables_data)


def find_first_signature_page(page_iterable):
    """
    Return the first 1-based page number that contains a signature block.
    Returns None if no signature page is found.
    """
    for page_num, page in enumerate(page_iterable, start=1):
        signers, _ = extract_person_signers(page)
        if signers:
            return page_num
    return None


def has_annex_keywords_after_signature(doc, first_signature_page):
    """
    Check a few pages after the first signature block for common annex markers.
    This is advisory only; the page-count threshold controls the actual warning.
    """
    if not first_signature_page:
        return False

    start_index = first_signature_page
    end_index = min(doc.page_count, first_signature_page + 5)

    for page_index in range(start_index, end_index):
        try:
            text = doc.load_page(page_index).get_text().upper()
        except Exception:
            continue
        if any(keyword in text for keyword in ANNEX_KEYWORDS):
            return True

    return False


def detect_long_annex_documents(document_files, page_threshold=LONG_ANNEX_PAGE_THRESHOLD):
    """
    Quickly scan PDFs for large annexes after the first signature block.
    Returns a list of warning dictionaries for the UI prompt.
    """
    warnings = []
    total = max(len(document_files), 1)

    for idx, (filename, filepath) in enumerate(document_files):
        percent = min(4, 1 + int(((idx + 1) / total) * 3))
        emit("progress", percent=percent, message=f"Checking annex length in {filename}")

        if not filename.lower().endswith('.pdf'):
            continue

        try:
            doc = fitz.open(filepath)
            first_signature_page = find_first_signature_page(doc)

            if first_signature_page:
                pages_after_signature = max(0, doc.page_count - first_signature_page)
                if pages_after_signature > page_threshold:
                    warnings.append({
                        "document": filename,
                        "total_pages": doc.page_count,
                        "first_signature_page": first_signature_page,
                        "pages_after_signature_block": pages_after_signature,
                        "annex_keywords_detected": has_annex_keywords_after_signature(doc, first_signature_page)
                    })

            doc.close()
        except Exception:
            continue

    return warnings


def apply_matching_paragraph_style(target_doc, target_para, source_para):
    """Apply the source paragraph style only if the target document exposes it."""
    try:
        style_name = source_para.style.name if source_para.style else ""
    except Exception:
        style_name = ""

    if not style_name:
        return

    try:
        target_para.style = target_doc.styles[style_name]
    except Exception:
        try:
            target_para.style = style_name
        except Exception:
            pass


def copy_run_formatting(source_run, target_run):
    """Copy a small, safe subset of run formatting across documents."""
    for attr in ("bold", "italic", "underline"):
        try:
            setattr(target_run, attr, getattr(source_run, attr))
        except Exception:
            pass

    try:
        if source_run.font.size:
            target_run.font.size = source_run.font.size
    except Exception:
        pass


def create_docx_packet(signer_name, docs_for_signer, output_folder):
    """
    Create a DOCX signature packet for a signer from DOCX source documents.

    Args:
        signer_name: Name of the signer
        docs_for_signer: List of (filename, filepath) tuples for this signer
        output_folder: Where to save the packet

    Returns:
        Path to created packet, or None if failed
    """
    try:
        ensure_docx_support()
        # Create new document for the packet
        packet_doc = Document()

        # Add header
        header_para = packet_doc.add_paragraph()
        header_run = header_para.add_run(f"SIGNATURE PACKET FOR {signer_name}")
        header_run.bold = True
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        packet_doc.add_paragraph()  # Blank line

        docs_added = 0

        for filename, filepath in docs_for_signer:
            if not filepath:
                emit("progress", percent=0, message=f"Warning: Missing file path for Word document {filename}")
                continue

            if not filepath.lower().endswith('.docx'):
                emit("progress", percent=0, message=f"Warning: Unsupported Word format for {filename}. Only .docx is supported in signature packets.")
                continue

            if not os.path.isfile(filepath):
                emit("progress", percent=0, message=f"Warning: Word document not found: {filepath}")
                continue

            try:
                # Open source document
                source_doc = Document(filepath)

                # Add document separator
                sep_para = packet_doc.add_paragraph()
                sep_run = sep_para.add_run(f"─" * 50)
                packet_doc.add_paragraph()

                doc_title = packet_doc.add_paragraph()
                title_run = doc_title.add_run(f"Document: {filename}")
                title_run.bold = True

                packet_doc.add_paragraph()  # Blank line

                # Copy content from source document
                for para in source_doc.paragraphs:
                    new_para = packet_doc.add_paragraph()
                    apply_matching_paragraph_style(packet_doc, new_para, para)

                    try:
                        new_para.alignment = para.alignment
                    except Exception:
                        pass

                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        copy_run_formatting(run, new_run)

                # Copy tables
                for table in source_doc.tables:
                    # Create table with same dimensions
                    new_table = packet_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    new_table.style = 'Table Grid'

                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            new_table.rows[i].cells[j].text = cell.text

                docs_added += 1

                # Add page break between documents
                packet_doc.add_page_break()

            except Exception as e:
                emit("progress", percent=0, message=f"Warning: Could not add Word document {filename}: {str(e)}")
                continue

        if docs_added > 0:
            packet_path = os.path.join(output_folder, f"signature_packet - {signer_name}.docx")
            packet_doc.save(packet_path)
            return packet_path

    except Exception as e:
        emit("progress", percent=0, message=f"Warning: Failed to assemble Word signature packet for {signer_name}: {str(e)}")

    return None


def create_pdf_packet(signer_name, docs_for_signer, output_folder, filepath_lookup):
    """
    Create a PDF signature packet for a signer from PDF source documents.

    Args:
        signer_name: Name of the signer
        docs_for_signer: DataFrame rows for this signer
        output_folder: Where to save the packet
        filepath_lookup: Dict mapping filename -> filepath

    Returns:
        Tuple of (path, page_count) or (None, 0) if failed
    """
    try:
        packet = fitz.open()

        for _, r in docs_for_signer.iterrows():
            if r["Document"].lower().endswith('.pdf'):
                try:
                    doc_path = filepath_lookup.get(r["Document"], r["Document"])
                    src = fitz.open(doc_path)
                    packet.insert_pdf(src, from_page=r["Page"] - 1, to_page=r["Page"] - 1)
                    src.close()
                except Exception:
                    pass

        if packet.page_count > 0:
            pdf_path = os.path.join(output_folder, f"signature_packet - {signer_name}.pdf")
            packet.save(pdf_path)
            page_count = packet.page_count
            packet.close()
            return pdf_path, page_count

        packet.close()
    except Exception:
        pass

    return None, 0


# ========== FORMAT CONVERSION ==========

def convert_pdf_to_docx(pdf_path, docx_path):
    """
    Convert PDF to DOCX using pdf2docx for high fidelity conversion.
    Returns True on success, False on failure.
    """
    try:
        from pdf2docx import Converter
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        return True
    except ImportError:
        emit("progress", percent=0, message="Warning: pdf2docx not installed, skipping conversion")
        return False
    except Exception as e:
        emit("progress", percent=0, message=f"Warning: PDF to DOCX conversion failed: {str(e)}")
        return False


def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Convert DOCX to PDF.
    Uses python-docx to read and PyMuPDF to create PDF.
    Note: This is a basic conversion - complex formatting may not be preserved perfectly.
    """
    try:
        ensure_docx_support()
        # Read DOCX content
        doc = Document(docx_path)

        # Create new PDF
        pdf_doc = fitz.open()
        page = pdf_doc.new_page()

        # Calculate page dimensions
        rect = page.rect
        margin = 72  # 1 inch margins
        y_position = margin

        # Process paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                y_position += 12  # Blank line
                continue

            # Determine font size based on style
            font_size = 11
            if para.style and 'Heading' in para.style.name:
                font_size = 14
            elif para.style and 'Title' in para.style.name:
                font_size = 16

            # Check if we need a new page
            if y_position > rect.height - margin:
                page = pdf_doc.new_page()
                y_position = margin

            # Insert text
            text_rect = fitz.Rect(margin, y_position, rect.width - margin, y_position + font_size + 4)
            page.insert_textbox(text_rect, text, fontsize=font_size)
            y_position += font_size + 6

        # Process tables
        for table in doc.tables:
            # Simple table representation
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if y_position > rect.height - margin:
                    page = pdf_doc.new_page()
                    y_position = margin

                text_rect = fitz.Rect(margin, y_position, rect.width - margin, y_position + 14)
                page.insert_textbox(text_rect, row_text, fontsize=10)
                y_position += 16

        if pdf_doc.page_count > 0:
            pdf_doc.save(pdf_path)
            pdf_doc.close()
            return True

        pdf_doc.close()
    except Exception as e:
        emit("progress", percent=0, message=f"Warning: DOCX to PDF conversion failed: {str(e)}")

    return False


def count_docx_packet_markers(docx_path):
    """Count the inserted document markers inside a generated DOCX packet."""
    ensure_docx_support()
    doc = Document(docx_path)
    return sum(1 for para in doc.paragraphs if para.text.strip().startswith("Document: "))


def build_signature_packet_verification(df, packets_created):
    """Verify packet outputs and flag ambiguous signature pages for review."""
    warnings = []
    packet_checks = []
    status = "passed"

    if df is None or df.empty:
        return {
            "status": "review_required",
            "warnings": ["No signature rows were captured for verification."],
            "packet_checks": [],
            "review_required_count": 0,
            "unassigned_signature_count": 0
        }

    review_required_count = 0
    unassigned_signature_count = 0
    review_documents = []

    if "Review Required" in df.columns:
        review_required_count = int(df["Review Required"].fillna(False).astype(bool).sum())

    if "Signer Name" in df.columns:
        unassigned_df = df[df["Signer Name"] == UNASSIGNED_SIGNER_BUCKET]
        unassigned_signature_count = int(len(unassigned_df))
        if not unassigned_df.empty and "Document" in unassigned_df.columns:
            review_documents = sorted(set(unassigned_df["Document"].astype(str).tolist()))

    if review_required_count > 0:
        status = "review_required"
        warnings.append(
            f"{review_required_count} signature assignment(s) need review because EmmaNeigh found signature cues without a confident signer match."
        )

    if unassigned_signature_count > 0:
        status = "review_required"
        warnings.append(
            f"{unassigned_signature_count} signature page(s) were routed into '{UNASSIGNED_SIGNER_BUCKET}' instead of being omitted."
        )

    for packet in packets_created or []:
        packet_path = str(packet.get("path") or "").strip()
        packet_format = str(packet.get("format") or "").strip().lower()
        expected_units = int(packet.get("pages") or 0)
        packet_check = {
            "path": packet_path,
            "format": packet_format,
            "expected_units": expected_units,
            "exists": os.path.isfile(packet_path)
        }

        if not packet_check["exists"]:
            packet_check["status"] = "missing"
            status = "review_required"
            warnings.append(f"Generated packet file is missing: {packet_path}")
            packet_checks.append(packet_check)
            continue

        try:
            if packet_format == "pdf":
                doc = fitz.open(packet_path)
                actual_units = doc.page_count
                doc.close()
                packet_check["actual_units"] = actual_units
                packet_check["status"] = "verified" if actual_units == expected_units else "count_mismatch"
            elif packet_format == "docx" and Document is not None:
                actual_units = count_docx_packet_markers(packet_path)
                packet_check["actual_units"] = actual_units
                packet_check["status"] = "verified" if actual_units == expected_units else "count_mismatch"
            else:
                packet_check["status"] = "unchecked"
        except Exception as exc:
            packet_check["status"] = "verification_error"
            packet_check["error"] = str(exc)

        if packet_check["status"] in {"count_mismatch", "verification_error"}:
            status = "review_required"
            warnings.append(f"Packet verification issue for {os.path.basename(packet_path)} ({packet_check['status']}).")

        packet_checks.append(packet_check)

    return {
        "status": status,
        "warnings": warnings,
        "packet_checks": packet_checks,
        "review_required_count": review_required_count,
        "unassigned_signature_count": unassigned_signature_count,
        "review_documents": review_documents
    }


def create_packet_with_format(signer_name, docs_for_signer, output_folder, filepath_lookup, output_format='preserve'):
    """
    Create signature packet(s) with specified output format.

    Args:
        signer_name: Name of the signer
        docs_for_signer: DataFrame rows for this signer
        output_folder: Where to save packets
        filepath_lookup: Dict mapping filename -> filepath
        output_format: 'preserve', 'pdf', 'docx', or 'both'

    Returns:
        List of created packet info dicts
    """
    packets = []

    # Separate by source format
    pdf_docs = docs_for_signer[docs_for_signer["Document"].str.lower().str.endswith('.pdf')]
    docx_docs = docs_for_signer[docs_for_signer["Document"].str.lower().str.endswith('.docx')]

    if output_format == 'preserve':
        # Original behavior: output matches input
        if len(pdf_docs) > 0:
            pdf_path, page_count = create_pdf_packet(signer_name, pdf_docs, output_folder, filepath_lookup)
            if pdf_path:
                packets.append({"name": signer_name, "pages": page_count, "format": "pdf", "path": pdf_path})

        if len(docx_docs) > 0:
            docx_files = [(r["Document"], filepath_lookup.get(r["Document"])) for _, r in docx_docs.iterrows()]
            docx_path = create_docx_packet(signer_name, docx_files, output_folder)
            if docx_path:
                packets.append({"name": signer_name, "pages": len(docx_files), "format": "docx", "path": docx_path})

    elif output_format == 'pdf':
        # Convert everything to PDF
        if len(pdf_docs) > 0:
            pdf_path, page_count = create_pdf_packet(signer_name, pdf_docs, output_folder, filepath_lookup)
            if pdf_path:
                packets.append({"name": signer_name, "pages": page_count, "format": "pdf", "path": pdf_path})

        if len(docx_docs) > 0:
            # First create DOCX packet, then convert to PDF
            docx_files = [(r["Document"], filepath_lookup.get(r["Document"])) for _, r in docx_docs.iterrows()]
            temp_docx_path = create_docx_packet(signer_name + "_temp", docx_files, output_folder)
            if temp_docx_path:
                pdf_path = os.path.join(output_folder, f"signature_packet - {signer_name} (from docx).pdf")
                if convert_docx_to_pdf(temp_docx_path, pdf_path):
                    packets.append({"name": signer_name, "pages": len(docx_files), "format": "pdf", "path": pdf_path})
                # Clean up temp file
                try:
                    os.remove(temp_docx_path)
                except:
                    pass

    elif output_format == 'docx':
        # Convert everything to DOCX
        if len(docx_docs) > 0:
            docx_files = [(r["Document"], filepath_lookup.get(r["Document"])) for _, r in docx_docs.iterrows()]
            docx_path = create_docx_packet(signer_name, docx_files, output_folder)
            if docx_path:
                packets.append({"name": signer_name, "pages": len(docx_files), "format": "docx", "path": docx_path})

        if len(pdf_docs) > 0:
            # First create PDF packet, then convert to DOCX
            pdf_path, page_count = create_pdf_packet(signer_name + "_temp", pdf_docs, output_folder, filepath_lookup)
            if pdf_path:
                docx_path = os.path.join(output_folder, f"signature_packet - {signer_name} (from pdf).docx")
                if convert_pdf_to_docx(pdf_path, docx_path):
                    packets.append({"name": signer_name, "pages": page_count, "format": "docx", "path": docx_path})
                # Clean up temp file
                try:
                    os.remove(pdf_path)
                except:
                    pass

    elif output_format == 'both':
        # Create both formats
        # First, create native format packets
        if len(pdf_docs) > 0:
            pdf_path, page_count = create_pdf_packet(signer_name, pdf_docs, output_folder, filepath_lookup)
            if pdf_path:
                packets.append({"name": signer_name, "pages": page_count, "format": "pdf", "path": pdf_path})
                # Also create DOCX version
                docx_path = os.path.join(output_folder, f"signature_packet - {signer_name} (from pdf).docx")
                if convert_pdf_to_docx(pdf_path, docx_path):
                    packets.append({"name": signer_name, "pages": page_count, "format": "docx", "path": docx_path})

        if len(docx_docs) > 0:
            docx_files = [(r["Document"], filepath_lookup.get(r["Document"])) for _, r in docx_docs.iterrows()]
            docx_path = create_docx_packet(signer_name, docx_files, output_folder)
            if docx_path:
                packets.append({"name": signer_name, "pages": len(docx_files), "format": "docx", "path": docx_path})
                # Also create PDF version
                pdf_path = os.path.join(output_folder, f"signature_packet - {signer_name} (from docx).pdf")
                if convert_docx_to_pdf(docx_path, pdf_path):
                    packets.append({"name": signer_name, "pages": len(docx_files), "format": "pdf", "path": pdf_path})

    return packets


# ========== MAIN ==========

def main():
    if len(sys.argv) < 2:
        emit("error", message="No input provided.")
        sys.exit(1)

    input_dir = None
    file_paths = None
    output_format = 'preserve'  # Default: output format matches input format
    preflight_long_annex_check = False
    long_annex_threshold = LONG_ANNEX_PAGE_THRESHOLD

    # Check if we have --config argument (for file list or folder config) or direct folder path
    if sys.argv[1] == '--config':
        if len(sys.argv) < 3:
            emit("error", message="No config file provided.")
            sys.exit(1)
        config_path = sys.argv[2]
        try:
            with open(config_path, 'r') as f:
                config = json.load(f)
            file_paths = config.get('files', [])
            output_format = config.get('output_format', 'preserve')
            preflight_long_annex_check = bool(config.get('preflight_long_annex_check', False))
            long_annex_threshold = int(config.get('long_annex_threshold', LONG_ANNEX_PAGE_THRESHOLD))

            # v5.1.6: Handle both 'files' list and 'folder' path from config
            if file_paths:
                # File list provided - create temp output folder only when doing the full run
                input_dir = None
            elif config.get('folder'):
                # Folder path provided in config
                input_dir = config.get('folder')
                if not os.path.isdir(input_dir):
                    emit("error", message=f"Invalid folder in config: {input_dir}")
                    sys.exit(1)
            else:
                emit("error", message="No files or folder in config.")
                sys.exit(1)
        except Exception as e:
            emit("error", message=f"Failed to read config: {str(e)}")
            sys.exit(1)
    else:
        input_dir = sys.argv[1]
        if not os.path.isdir(input_dir):
            emit("error", message=f"Invalid folder: {input_dir}")
            sys.exit(1)

    # Get document files - either from file_paths list or scan directory (including subdirectories)
    if file_paths:
        # Filter to valid PDF/DOCX files
        document_files = [(os.path.basename(f), f) for f in file_paths
                          if os.path.isfile(f) and f.lower().endswith((".pdf", ".docx"))]
    else:
        # v5.1.6: Scan directory recursively for all PDF/DOCX files
        document_files = []
        for root, dirs, files in os.walk(input_dir):
            # Skip output directories
            if 'signature_packets_output' in root:
                continue
            for f in files:
                if f.lower().endswith((".pdf", ".docx")):
                    filepath = os.path.join(root, f)
                    # Use relative path as display name if in subdirectory
                    rel_path = os.path.relpath(filepath, input_dir)
                    document_files.append((rel_path, filepath))

    if not document_files:
        emit("error", message="No PDF or Word files found.")
        sys.exit(1)

    total = len(document_files)
    emit("progress", percent=8, message=f"Found {total} documents")

    if preflight_long_annex_check:
        warnings = detect_long_annex_documents(document_files, long_annex_threshold)
        emit(
            "result",
            success=True,
            documentsChecked=total,
            longAnnexThreshold=long_annex_threshold,
            longAnnexWarnings=warnings
        )
        return

    if file_paths and not input_dir:
        import tempfile
        input_dir = tempfile.mkdtemp(prefix='emmaneigh_packets_')

    # Set up output directories
    output_base = os.path.join(input_dir, "signature_packets_output")
    output_pdf_dir = os.path.join(output_base, "packets")
    output_table_dir = os.path.join(output_base, "tables")

    os.makedirs(output_pdf_dir, exist_ok=True)
    os.makedirs(output_table_dir, exist_ok=True)

    # Scan all documents for signature pages
    rows = []
    # Build filepath lookup for later use
    filepath_lookup = {filename: filepath for filename, filepath in document_files}

    for idx, (filename, filepath) in enumerate(document_files):
        percent = 10 + int(((idx + 1) / total) * 40)
        emit("progress", percent=percent, message=f"Scanning {filename}")

        try:
            if filename.lower().endswith('.pdf'):
                # PDF processing
                doc = fitz.open(filepath)
                for page_num, page in enumerate(doc, start=1):
                    analysis = analyze_pdf_signature_page(page)
                    if analysis["is_signature_page"]:
                        signers = analysis["signers"] or {UNASSIGNED_SIGNER_BUCKET}
                        footer = extract_footer_from_pdf_page(page)
                        doc_id = extract_document_id_from_pdf_page(page)
                        cue_hits = ",".join(analysis["cue_hits"])
                        for signer in signers:
                            rows.append({
                                "Signer Name": signer,
                                "Document": filename,
                                "Page": page_num,
                                "Document ID": doc_id,
                                "Footer": footer,
                                "Detection Method": analysis["method"],
                                "Review Required": bool(analysis["needs_review"] or signer == UNASSIGNED_SIGNER_BUCKET),
                                "Signature Cue Score": analysis["cue_score"],
                                "Signature Cues": cue_hits
                            })
                doc.close()
            elif filename.lower().endswith('.docx'):
                # DOCX processing
                analysis = analyze_docx_signature_page(filepath)
                if analysis["is_signature_page"]:
                    signers = analysis["signers"] or {UNASSIGNED_SIGNER_BUCKET}
                    footer = extract_footer_from_docx(filepath)
                    doc_id = extract_document_id_from_docx(filepath)
                    cue_hits = ",".join(analysis["cue_hits"])
                    for signer in signers:
                        rows.append({
                            "Signer Name": signer,
                            "Document": filename,
                            "Page": 1,  # DOCX doesn't have pages
                            "Document ID": doc_id,
                            "Footer": footer,
                            "Detection Method": analysis["method"],
                            "Review Required": bool(analysis["needs_review"] or signer == UNASSIGNED_SIGNER_BUCKET),
                            "Signature Cue Score": analysis["cue_score"],
                            "Signature Cues": cue_hits
                        })
        except Exception as e:
            emit("progress", percent=percent, message=f"Warning: {filename} - {str(e)}")

    if not rows:
        emit("error", message="No signature pages detected in any documents.")
        sys.exit(1)

    # Create DataFrame and sort
    # Columns: Signer Name, Document, Page, Document ID, Footer, Detection Method
    df = pd.DataFrame(rows)
    # Reorder columns for cleaner output
    column_order = [
        "Signer Name", "Document", "Page", "Document ID", "Footer", "Detection Method",
        "Review Required", "Signature Cue Score", "Signature Cues"
    ]
    df = df[[col for col in column_order if col in df.columns]]
    df = df.sort_values(["Signer Name", "Document", "Page"])

    # Save master index
    emit("progress", percent=60, message="Creating master index...")
    df.to_excel(os.path.join(output_table_dir, "MASTER_SIGNATURE_INDEX.xlsx"), index=False)

    # Create individual packets with specified output format
    signers = df.groupby("Signer Name")
    total_signers = len(signers)
    packets_created = []

    emit("progress", percent=65, message=f"Creating packets (format: {output_format})...")

    for idx, (signer, group) in enumerate(signers):
        percent = 65 + int(((idx + 1) / total_signers) * 30)
        emit("progress", percent=percent, message=f"Creating packet for {signer}")

        # Save signer's Excel file
        group.to_excel(
            os.path.join(output_table_dir, f"signature_packet - {signer}.xlsx"),
            index=False
        )

        # Create packets using the new format-aware function
        signer_packets = create_packet_with_format(
            signer, group, output_pdf_dir, filepath_lookup, output_format
        )
        packets_created.extend(signer_packets)

    verification = build_signature_packet_verification(df, packets_created)
    with open(os.path.join(output_table_dir, "SIGNATURE_PACKET_VERIFICATION.json"), "w", encoding="utf-8") as handle:
        json.dump(verification, handle, indent=2)

    emit("progress", percent=96, message="Preparing download...")
    emit("result",
         success=True,
         outputPath=output_base,
         packetsCreated=len(packets_created),
         packets=packets_created,
         verification=verification)


if __name__ == "__main__":
    main()
