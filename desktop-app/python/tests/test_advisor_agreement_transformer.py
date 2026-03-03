import json
import tempfile
import unittest
import zipfile
from pathlib import Path
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from advisor_agreement_transformer import transform_agreement

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def _add_clause_with_blue(paragraph, prefix: str, placeholder: str, suffix: str):
    r1 = paragraph.add_run("[")
    r1.font.highlight_color = WD_COLOR_INDEX.YELLOW

    r2 = paragraph.add_run(prefix)
    r2.font.highlight_color = WD_COLOR_INDEX.YELLOW

    r3 = paragraph.add_run("[" + placeholder + "]")
    r3.font.highlight_color = WD_COLOR_INDEX.BLUE

    r4 = paragraph.add_run(suffix)
    r4.font.highlight_color = WD_COLOR_INDEX.YELLOW

    r5 = paragraph.add_run("]")
    r5.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _add_simple_yellow_clause(paragraph, text_inside: str):
    r = paragraph.add_run("[" + text_inside + "]")
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _attach_footnote_ref(paragraph, footnote_id: int):
    run = paragraph.add_run("")
    footnote_ref = OxmlElement("w:footnoteReference")
    footnote_ref.set(qn("w:id"), str(footnote_id))
    run._r.append(footnote_ref)


def _inject_footnotes_xml(docx_path: Path, footnotes: dict[int, str]):
    with zipfile.ZipFile(docx_path, "a") as zf:
        if "word/footnotes.xml" in zf.namelist():
            zf.writestr("word/footnotes.xml", _build_footnotes_xml(footnotes))
            return

        zf.writestr("word/footnotes.xml", _build_footnotes_xml(footnotes))

        # Add content type override.
        ct = ET.fromstring(zf.read("[Content_Types].xml"))
        has_override = any(
            node.attrib.get("PartName") == "/word/footnotes.xml"
            for node in ct.findall("{*}Override")
        )
        if not has_override:
            override = ET.Element(
                "Override",
                {
                    "PartName": "/word/footnotes.xml",
                    "ContentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
                },
            )
            ct.append(override)
            zf.writestr("[Content_Types].xml", ET.tostring(ct, encoding="utf-8", xml_declaration=True))

        # Add relationship from document.xml to footnotes.xml.
        rels_path = "word/_rels/document.xml.rels"
        rels = ET.fromstring(zf.read(rels_path))
        exists = any(
            node.attrib.get("Type", "").endswith("/footnotes")
            for node in rels.findall("{*}Relationship")
        )
        if not exists:
            rel_ids = []
            for node in rels.findall("{*}Relationship"):
                rid = node.attrib.get("Id", "")
                if rid.startswith("rId") and rid[3:].isdigit():
                    rel_ids.append(int(rid[3:]))
            next_id = (max(rel_ids) + 1) if rel_ids else 1
            rel = ET.Element(
                "Relationship",
                {
                    "Id": f"rId{next_id}",
                    "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
                    "Target": "footnotes.xml",
                },
            )
            rels.append(rel)
            zf.writestr(rels_path, ET.tostring(rels, encoding="utf-8", xml_declaration=True))


def _build_footnotes_xml(footnotes: dict[int, str]) -> bytes:
    root = ET.Element(f"{{{W_NS}}}footnotes")
    ET.SubElement(root, f"{{{W_NS}}}footnote", {f"{{{W_NS}}}id": "0"})
    ET.SubElement(root, f"{{{W_NS}}}footnote", {f"{{{W_NS}}}id": "-1"})

    for fid, text in footnotes.items():
        fn = ET.SubElement(root, f"{{{W_NS}}}footnote", {f"{{{W_NS}}}id": str(fid)})
        p = ET.SubElement(fn, f"{{{W_NS}}}p")
        r = ET.SubElement(p, f"{{{W_NS}}}r")
        t = ET.SubElement(r, f"{{{W_NS}}}t")
        t.text = text

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _all_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    return "\n".join(p.text for p in doc.paragraphs)


class TestDeterministicAgreementTransformer(unittest.TestCase):
    def test_nested_yellow_blue_clause(self):
        with tempfile.TemporaryDirectory() as td:
            base = Path(td)
            template = base / "template.docx"
            output = base / "output.docx"

            doc = Document()
            p1 = doc.add_paragraph()
            _add_simple_yellow_clause(p1, "DELETE_ME")

            p2 = doc.add_paragraph()
            _add_clause_with_blue(p2, "Advisor ", "ADVISOR_NAME", " engaged")
            doc.save(template)

            log = transform_agreement(
                template,
                {"advisor_name": "Jane Roe"},
                output,
            )

            text = _all_text(output)
            self.assertIn("Advisor Jane Roe engaged", text)
            self.assertNotIn("[", text)
            self.assertTrue(log["deleted_first_yellow_block"])

    def test_conditional_clause_deletion(self):
        with tempfile.TemporaryDirectory() as td:
            base = Path(td)
            template = base / "template.docx"
            output = base / "output.docx"

            doc = Document()
            p1 = doc.add_paragraph()
            _add_simple_yellow_clause(p1, "DELETE_ME")

            p2 = doc.add_paragraph()
            _add_clause_with_blue(p2, "Tax clause ", "ADVISOR_NAME", " end")
            _attach_footnote_ref(p2, 1)
            doc.save(template)
            _inject_footnotes_xml(template, {1: "Only include if advisor is a US taxpayer."})

            log = transform_agreement(
                template,
                {"advisor_name": "Alex", "is_us_taxpayer": "no"},
                output,
            )

            text = _all_text(output)
            self.assertNotIn("Tax clause", text)
            self.assertGreaterEqual(len(log["clauses_deleted"]), 2)  # first yellow + conditional

    def test_conditional_clause_retention(self):
        with tempfile.TemporaryDirectory() as td:
            base = Path(td)
            template = base / "template.docx"
            output = base / "output.docx"

            doc = Document()
            p1 = doc.add_paragraph()
            _add_simple_yellow_clause(p1, "DELETE_ME")

            p2 = doc.add_paragraph()
            _add_clause_with_blue(p2, "Tax clause ", "ADVISOR_NAME", " end")
            _attach_footnote_ref(p2, 1)
            doc.save(template)
            _inject_footnotes_xml(template, {1: "Only include if advisor is a US taxpayer."})

            transform_agreement(
                template,
                {"advisor_name": "Alex", "is_us_taxpayer": "yes"},
                output,
            )

            text = _all_text(output)
            self.assertIn("Tax clause Alex end", text)
            self.assertNotIn("[", text)

    def test_always_delete_first_clause(self):
        with tempfile.TemporaryDirectory() as td:
            base = Path(td)
            template = base / "template.docx"
            output = base / "output.docx"

            doc = Document()
            p1 = doc.add_paragraph()
            _add_simple_yellow_clause(p1, "FIRST")
            p2 = doc.add_paragraph()
            _add_clause_with_blue(p2, "Second ", "ADVISOR_NAME", " clause")
            doc.save(template)

            log = transform_agreement(template, {"advisor_name": "A"}, output)
            text = _all_text(output)
            self.assertNotIn("FIRST", text)
            self.assertTrue(log["deleted_first_yellow_block"])

    def test_missing_data_preservation(self):
        with tempfile.TemporaryDirectory() as td:
            base = Path(td)
            template = base / "template.docx"
            output = base / "output.docx"

            doc = Document()
            p1 = doc.add_paragraph()
            _add_simple_yellow_clause(p1, "DELETE_ME")

            p2 = doc.add_paragraph()
            _add_clause_with_blue(p2, "Keep unresolved ", "ADVISOR_NAME", " end")
            _attach_footnote_ref(p2, 1)
            doc.save(template)
            _inject_footnotes_xml(template, {1: "Only include if advisor is a US taxpayer."})

            log = transform_agreement(
                template,
                {"is_us_taxpayer": "yes"},
                output,
            )

            text = _all_text(output)
            self.assertIn("[Keep unresolved [ADVISOR_NAME] end]", text)
            self.assertIn("ADVISOR_NAME", log["blue_placeholders_missing"])

    def test_footnotes_removed(self):
        with tempfile.TemporaryDirectory() as td:
            base = Path(td)
            template = base / "template.docx"
            output = base / "output.docx"

            doc = Document()
            p1 = doc.add_paragraph()
            _add_simple_yellow_clause(p1, "DELETE_ME")
            p2 = doc.add_paragraph()
            _add_clause_with_blue(p2, "Tax ", "ADVISOR_NAME", " x")
            _attach_footnote_ref(p2, 1)
            doc.save(template)
            _inject_footnotes_xml(template, {1: "Only include if advisor is a US taxpayer."})

            transform_agreement(
                template,
                {"advisor_name": "A", "is_us_taxpayer": "yes"},
                output,
            )

            # Ensure no footnoteReference remains in document xml.
            with zipfile.ZipFile(output, "r") as zf:
                doc_xml = zf.read("word/document.xml").decode("utf-8")
                self.assertNotIn("footnoteReference", doc_xml)

                if "word/footnotes.xml" in zf.namelist():
                    fn_xml = zf.read("word/footnotes.xml").decode("utf-8")
                    self.assertNotIn('w:id="1"', fn_xml)

    def test_currency_global_selection(self):
        with tempfile.TemporaryDirectory() as td:
            base = Path(td)
            template = base / "template.docx"
            output = base / "output.docx"

            doc = Document()
            p1 = doc.add_paragraph()
            _add_simple_yellow_clause(p1, "DELETE_ME")

            p2 = doc.add_paragraph()
            _add_clause_with_blue(p2, "Pay in ", "$/EUROS/GBP", " only")
            p3 = doc.add_paragraph()
            _add_clause_with_blue(p3, "Again ", "$/EUROS/GBP", " here")
            doc.save(template)

            log = transform_agreement(
                template,
                {"currency": "EUR"},
                output,
            )

            text = _all_text(output)
            self.assertIn("Pay in EUROS only", text)
            self.assertIn("Again EUROS here", text)
            self.assertEqual(log["global_variables"].get("currency_selection"), "EUROS")


if __name__ == "__main__":
    unittest.main()
