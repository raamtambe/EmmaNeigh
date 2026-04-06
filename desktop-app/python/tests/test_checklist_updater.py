import unittest
from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

try:
    from docx import Document
except Exception:  # pragma: no cover - environment dependent
    Document = None

try:
    from checklist_updater import (
        build_email_threads,
        build_row_thread_candidates,
        classify_checklist_row,
        infer_email_direction,
        normalize_email_record,
        refresh_document_draft_stamp,
    )
except Exception as exc:  # pragma: no cover - environment dependent
    build_email_threads = None
    build_row_thread_candidates = None
    classify_checklist_row = None
    infer_email_direction = None
    normalize_email_record = None
    refresh_document_draft_stamp = None
    IMPORT_ERROR = exc
else:  # pragma: no cover - environment dependent
    IMPORT_ERROR = None


@unittest.skipUnless(build_email_threads is not None, f"checklist updater import failed: {IMPORT_ERROR}")
class ChecklistUpdaterHeuristicsTests(unittest.TestCase):
    def test_normalize_email_record_uses_current_message_body(self):
        email = normalize_email_record({
            "subject": "Purchase Agreement comments",
            "body": "Please see attached markup.\n\n-----Original Message-----\nOlder email about disclosure schedules only.",
        })

        self.assertIn("Please see attached markup.", email["body_current"])
        self.assertNotIn("Older email", email["body_current"])
        self.assertNotIn("Older email", email["searchable"])

    def test_infer_email_direction_uses_sent_folder_flag(self):
        email = normalize_email_record({
            "subject": "Purchase Agreement markup",
            "from": "",
            "to": "Buyer Counsel",
            "date_sent": "2026-03-12T11:00:00",
            "folder": r"\\Mailbox - Team\\Deal Files\\Acme Nova",
            "root_folder": "Deal Files",
            "store": "Mailbox - Team",
            "is_sent_folder": True,
        })

        self.assertEqual(infer_email_direction(email), "sent")

    def test_build_email_threads_groups_replies_together(self):
        emails = [
            normalize_email_record({
                "subject": "Purchase Agreement comments",
                "conversation_topic": "Purchase Agreement comments",
                "from": "Seller Counsel",
                "to": "Me",
                "date_received": "2026-03-10T09:00:00",
                "body": "Attached is the revised draft.",
                "attachments": ["Purchase Agreement.docx"],
                "folder": r"\\Inbox\\Acme Nova",
                "root_folder": "Inbox",
            }),
            normalize_email_record({
                "subject": "RE: Purchase Agreement comments",
                "conversation_topic": "Purchase Agreement comments",
                "from": "Me",
                "to": "Buyer Counsel",
                "date_sent": "2026-03-11T10:00:00",
                "body": "Please see our markup.",
                "attachments": ["Purchase Agreement markup.docx"],
                "folder": r"\\Sent Items",
                "root_folder": "Sent Items",
            }),
            normalize_email_record({
                "subject": "Disclosure Schedule",
                "from": "Paralegal",
                "to": "Me",
                "date_received": "2026-03-09T08:00:00",
                "body": "Administrative follow-up only.",
                "folder": r"\\Inbox\\Acme Nova",
                "root_folder": "Inbox",
            }),
        ]

        threads = build_email_threads(emails)

        self.assertEqual(len(threads), 2)
        purchase_thread = next(thread for thread in threads if "purchase agreement" in thread["normalized_subject"])
        self.assertEqual(purchase_thread["message_count"], 2)
        self.assertIn("Seller Counsel", purchase_thread["participants"])
        self.assertIn("Buyer Counsel", purchase_thread["participants"])

    def test_row_thread_candidates_capture_draft_flow_and_issues(self):
        checklist_items = [{
            "row_id": 4,
            "document_name": "Purchase Agreement",
            "current_status": "",
            "existing_notes": "",
            "row_context": "Buyer: Acme | Seller: Nova | Responsible: Buyer Counsel",
        }]
        emails = [
            normalize_email_record({
                "subject": "Purchase Agreement - revised draft",
                "conversation_id": "deal-thread-1",
                "from": "Seller Counsel",
                "to": "Me",
                "date_received": "2026-03-10T09:00:00",
                "body": "Attached is the revised Purchase Agreement. Open issue: working capital adjustment remains bracketed.",
                "attachments": ["Acme Nova Purchase Agreement redline.docx"],
                "folder": r"\\Inbox\\Acme Nova",
                "root_folder": "Inbox",
            }),
            normalize_email_record({
                "subject": "RE: Purchase Agreement - revised draft",
                "conversation_id": "deal-thread-1",
                "from": "Me",
                "to": "Buyer Counsel",
                "date_sent": "2026-03-12T11:00:00",
                "body": "Please see our markup. We still need to resolve the disclosure schedules.",
                "attachments": ["Purchase Agreement markup.docx"],
                "folder": r"\\Sent Items",
                "root_folder": "Sent Items",
            }),
        ]

        candidate_map, _, deal_profile = build_row_thread_candidates(
            "/tmp/Acme Nova Closing Checklist.docx",
            checklist_items,
            emails,
        )

        self.assertIn("acme", deal_profile["anchor_tokens"])
        self.assertIn("nova", deal_profile["anchor_tokens"])
        self.assertIn(4, candidate_map)
        self.assertTrue(candidate_map[4])

        top_candidate = candidate_map[4][0]
        self.assertEqual(top_candidate["status_signal"], "With Opposing Counsel")
        self.assertIn("Received draft from Seller Counsel on 03/10/2026.", top_candidate["events"])
        self.assertIn("Sent draft to Buyer Counsel on 03/12/2026.", top_candidate["events"])
        self.assertTrue(any("working capital adjustment" in issue.lower() for issue in top_candidate["issues"]))
        self.assertIn("Issues flagged", top_candidate["suggested_comment"])

    def test_right_deal_thread_ranks_above_wrong_deal_thread(self):
        checklist_items = [{
            "row_id": 7,
            "document_name": "Purchase Agreement",
            "current_status": "",
            "existing_notes": "",
            "row_context": "Buyer: Acme | Seller: Nova",
        }]
        emails = [
            normalize_email_record({
                "subject": "Purchase Agreement draft",
                "conversation_id": "correct-deal",
                "from": "Seller Counsel",
                "to": "Me",
                "date_received": "2026-03-10T09:00:00",
                "body": "Acme and Nova comments are attached.",
                "attachments": ["Acme Nova Purchase Agreement.docx"],
                "folder": r"\\Inbox\\Acme Nova",
                "root_folder": "Inbox",
            }),
            normalize_email_record({
                "subject": "Purchase Agreement draft",
                "conversation_id": "wrong-deal",
                "from": "Other Counsel",
                "to": "Me",
                "date_received": "2026-03-10T12:00:00",
                "body": "Orion and Zenith comments are attached.",
                "attachments": ["Orion Zenith Purchase Agreement.docx"],
                "folder": r"\\Inbox\\Orion Zenith",
                "root_folder": "Inbox",
            }),
        ]

        candidate_map, _, _ = build_row_thread_candidates(
            "/tmp/Acme Nova Closing Checklist.docx",
            checklist_items,
            emails,
        )

        self.assertGreaterEqual(len(candidate_map[7]), 2)
        self.assertGreater(candidate_map[7][0]["score"], candidate_map[7][1]["score"])
        self.assertIn("acme", candidate_map[7][0]["deal_anchor_hits"])

    def test_classify_checklist_row_skips_subrows_and_section_headers(self):
        header_row = classify_checklist_row(
            "PURCHASE AGREEMENT COMMENTS:",
            ["PURCHASE AGREEMENT COMMENTS:", "", ""],
            ["Document", "Status", "Comments"],
            0,
            1,
            2,
            previous_document_name="Purchase Agreement",
        )
        subrow = classify_checklist_row(
            "1. Representations and Warranties",
            ["1. Representations and Warranties", "", ""],
            ["Document", "Status", "Comments"],
            0,
            1,
            2,
            previous_document_name="Purchase Agreement",
        )
        document_row = classify_checklist_row(
            "Purchase Agreement",
            ["Purchase Agreement", "", ""],
            ["Document", "Status", "Comments"],
            0,
            1,
            2,
            previous_document_name="",
        )

        self.assertEqual(header_row["row_type"], "section_header")
        self.assertFalse(header_row["should_update"])
        self.assertEqual(subrow["row_type"], "subrow")
        self.assertFalse(subrow["should_update"])
        self.assertEqual(document_row["row_type"], "document")
        self.assertTrue(document_row["should_update"])

    @unittest.skipUnless(Document is not None and refresh_document_draft_stamp is not None, "python-docx is required")
    def test_refresh_document_draft_stamp_updates_header_date(self):
        doc = Document()
        section = doc.sections[0]
        section.header.paragraphs[0].text = "DRAFT 03/10/2026"

        refresh_document_draft_stamp(doc)

        self.assertIn("DRAFT", section.header.paragraphs[0].text.upper())
        self.assertNotIn("03/10/2026", section.header.paragraphs[0].text)


if __name__ == "__main__":
    unittest.main()
