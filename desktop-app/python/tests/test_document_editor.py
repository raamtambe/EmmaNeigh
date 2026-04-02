import tempfile
import unittest
from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

try:
    import fitz
except ImportError:  # pragma: no cover - environment dependent
    fitz = None

try:
    from docx import Document
except ImportError:  # pragma: no cover - environment dependent
    Document = None

from document_editor import (
    apply_docx_replacements,
    apply_pdf_replacements,
    extract_text_from_docx,
    extract_text_from_pdf,
)


@unittest.skipUnless(fitz is not None and Document is not None, "document editor tests require PyMuPDF and python-docx")
class DocumentEditorTests(unittest.TestCase):
    def test_docx_replacements_can_span_runs(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            base = Path(temp_dir)
            input_path = base / "agreement.docx"
            output_path = base / "agreement_edited.docx"

            doc = Document()
            paragraph = doc.add_paragraph()
            paragraph.add_run("The Closing Date is January ")
            paragraph.add_run("1, 2026")
            doc.add_paragraph("Borrower: Old Parent, LLC")
            doc.save(input_path)

            result = apply_docx_replacements(
                str(input_path),
                str(output_path),
                [
                    {"find_text": "January 1, 2026", "replace_text": "June 30, 2026", "reason": "date"},
                    {"find_text": "Old Parent, LLC", "replace_text": "New Parent, LLC", "reason": "party"},
                ],
            )

            self.assertTrue(result["success"])
            self.assertEqual(result["total_replacements"], 2)

            text = extract_text_from_docx(str(output_path))
            self.assertIn("June 30, 2026", text)
            self.assertIn("New Parent, LLC", text)
            self.assertNotIn("January 1, 2026", text)
            self.assertNotIn("Old Parent, LLC", text)

    def test_pdf_replacements_update_visible_text(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            base = Path(temp_dir)
            input_path = base / "notice.pdf"
            output_path = base / "notice_edited.pdf"

            pdf = fitz.open()
            page = pdf.new_page()
            page.insert_text((72, 72), "Execution Date: January 1, 2026", fontsize=12)
            page.insert_text((72, 96), "Seller: Old Parent, LLC", fontsize=12)
            pdf.save(input_path)
            pdf.close()

            result = apply_pdf_replacements(
                str(input_path),
                str(output_path),
                [
                    {"find_text": "January 1, 2026", "replace_text": "June 30, 2026", "reason": "date"},
                    {"find_text": "Old Parent, LLC", "replace_text": "New Parent, LLC", "reason": "party"},
                ],
            )

            self.assertTrue(result["success"])
            self.assertEqual(result["total_replacements"], 2)

            text = extract_text_from_pdf(str(output_path))
            self.assertIn("June 30, 2026", text)
            self.assertIn("New Parent, LLC", text)


if __name__ == "__main__":
    unittest.main()
