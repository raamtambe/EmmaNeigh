import os
import sys
import tempfile
import unittest
from pathlib import Path

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
except Exception:
    Document = None
    WD_STYLE_TYPE = None

TESTS_DIR = Path(__file__).resolve().parent
PYTHON_DIR = TESTS_DIR.parent
if str(PYTHON_DIR) not in sys.path:
    sys.path.insert(0, str(PYTHON_DIR))

try:
    from signature_packets import (
        UNASSIGNED_SIGNER_BUCKET,
        analyze_signature_page_text,
        build_signature_packet_verification,
        create_docx_packet,
    )
except Exception:
    UNASSIGNED_SIGNER_BUCKET = None
    analyze_signature_page_text = None
    build_signature_packet_verification = None
    create_docx_packet = None


@unittest.skipUnless(analyze_signature_page_text is not None and build_signature_packet_verification is not None, "signature_packets helpers are required")
class SignaturePacketsDetectionTests(unittest.TestCase):
    def test_detects_inline_by_signer(self):
        analysis = analyze_signature_page_text("ACME HOLDINGS LLC\nBy: /s/ Jane Doe\nName: Jane Doe\nTitle: CEO")

        self.assertTrue(analysis["is_signature_page"])
        self.assertIn("JANE DOE", analysis["signers"])
        self.assertFalse(analysis["needs_review"])

    def test_routes_signature_cues_without_signer_to_review(self):
        analysis = analyze_signature_page_text("ACME HOLDINGS LLC\nBy: __________________\nName: __________________\nTitle: __________________")

        self.assertTrue(analysis["is_signature_page"])
        self.assertTrue(analysis["needs_review"])
        self.assertIn("ACME HOLDINGS LLC", analysis["signers"])

    def test_verification_flags_unassigned_review_rows(self):
        import pandas as pd

        df = pd.DataFrame([
            {
                "Signer Name": UNASSIGNED_SIGNER_BUCKET,
                "Document": "Purchase Agreement.pdf",
                "Page": 12,
                "Review Required": True,
                "Detection Method": "SIGNATURE_CUE_REVIEW"
            }
        ])

        verification = build_signature_packet_verification(df, [])

        self.assertEqual("review_required", verification["status"])
        self.assertEqual(1, verification["review_required_count"])
        self.assertEqual(1, verification["unassigned_signature_count"])


@unittest.skipUnless(Document is not None and WD_STYLE_TYPE is not None and create_docx_packet is not None, "python-docx and signature_packets are required")
class SignaturePacketsDocxTests(unittest.TestCase):
    def test_create_docx_packet_handles_custom_source_styles(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            source_path = tmp_path / "Purchase Agreement.docx"
            output_dir = tmp_path / "output"
            output_dir.mkdir()

            source_doc = Document()
            custom_style = source_doc.styles.add_style("ClosingSignatureStyle", WD_STYLE_TYPE.PARAGRAPH)
            custom_style.font.bold = True

            source_doc.add_paragraph("Execution Version")
            source_doc.add_paragraph("By: John Smith", style="ClosingSignatureStyle")
            source_doc.add_paragraph("Name: John Smith")
            source_doc.save(source_path)

            packet_path = create_docx_packet("John Smith", [(source_path.name, str(source_path))], str(output_dir))

            self.assertTrue(packet_path)
            self.assertTrue(os.path.exists(packet_path))

            packet_doc = Document(packet_path)
            packet_text = "\n".join(paragraph.text for paragraph in packet_doc.paragraphs)

            self.assertIn("Document: Purchase Agreement.docx", packet_text)
            self.assertIn("By: John Smith", packet_text)


if __name__ == "__main__":
    unittest.main()
